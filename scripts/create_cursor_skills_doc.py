#!/usr/bin/env python3
"""
Generate Word document listing Cursor skills applicable to the Cruise Employee
English Assessment Platform. Saves to docs/Cursor_Skills_For_Project.docx.
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# Content: categories and skills (name, short description/use case)
SKILLS_BY_CATEGORY = [
    ("一、后端与 API", [
        ("fastapi-pro", "FastAPI 异步 API、SQLAlchemy 2、Pydantic V2、微服务与 WebSocket"),
        ("fastapi-router-py", "新路由、CRUD、认证依赖、响应模型"),
        ("fastapi-templates", "新 FastAPI 项目脚手架、async 与依赖注入"),
        ("api-design-principles", "REST/API 设计、版本、分页"),
        ("api-security-best-practices", "认证、鉴权、输入校验、限流、安全设计"),
        ("backend-security-coder", "输入校验、认证、API 安全编码"),
    ]),
    ("二、数据库与存储", [
        ("database", "SQL、NoSQL、迁移、优化、数据工程"),
        ("postgresql", "Postgres 建模、索引、约束、性能"),
        ("postgres-best-practices", "Postgres 性能与最佳实践（Supabase 等）"),
        ("sql-optimization-patterns", "慢查询、EXPLAIN、索引策略"),
        ("database-migration", "跨 ORM/平台迁移、零停机、回滚"),
    ]),
    ("三、AI / 语音 / 评分", [
        ("ai-ml", "LLM 应用、RAG、Agent、ML 流程"),
        ("voice-ai-development", "语音 AI（OpenAI Realtime、Deepgram、TTS 等）"),
        ("llm-app-patterns", "生产级 LLM 应用、RAG、Agent、LLMOps"),
        ("prompt-engineering", "提示词优化、调试 Agent 行为"),
        ("llm-evaluation", "LLM 应用评估、自动化指标、人工反馈"),
    ]),
    ("四、测试与质量", [
        ("python-testing-patterns", "pytest、fixtures、mock、TDD"),
        ("e2e-testing", "Playwright、浏览器自动化、视觉回归"),
        ("webapp-testing", "本地 Web 应用测试、Playwright、截图"),
        ("testing-qa", "单元/集成/E2E、质量保证流程"),
        ("verification-before-completion", "完成前跑验证、避免误报「已完成」"),
    ]),
    ("五、认证与安全", [
        ("auth-implementation-patterns", "JWT、OAuth2、Session、RBAC"),
        ("api-security-testing", "REST/GraphQL 安全测试、认证与限流"),
        ("security-audit", "Web/API 安全审计、渗透测试思路"),
    ]),
    ("六、文档与协作", [
        ("documentation", "API 文档、架构说明、README、技术写作"),
        ("readme", "创建/更新项目 README"),
        ("api-documentation", "OpenAPI、开发者文档、文档工作流"),
    ]),
    ("七、Python 与工程实践", [
        ("python-pro", "Python 3.12+、async、性能、uv/ruff/pydantic"),
        ("async-python-patterns", "asyncio、并发、async/await"),
        ("pydantic-models-py", "Base/Create/Update/Response 等多模型模式"),
        ("clean-code", "可读性、命名、函数设计、可维护性"),
        ("code-review-checklist", "功能、安全、性能、可维护性检查"),
    ]),
    ("八、前端与界面", [
        ("web-design-guidelines", "UI 规范、可访问性、设计审查"),
        ("fixing-accessibility", "ARIA、键盘导航、对比度、表单错误"),
        ("wcag-audit-patterns", "WCAG 2.2 审计与修复"),
    ]),
    ("九、部署与运维", [
        ("docker-expert", "多阶段构建、镜像优化、Compose、安全"),
        ("github-actions-templates", "CI/CD、测试、构建、部署工作流"),
        ("observability-engineer", "监控、日志、追踪、SRE"),
    ]),
    ("十、项目管理与规范", [
        ("plan-writing", "任务拆分、依赖、验收标准"),
        ("writing-plans", "有 spec/需求时先写计划再写代码"),
        ("concise-planning", "清晰、可执行的编码任务清单"),
        ("tdd-workflow", "红-绿-重构、先写测试"),
    ]),
]

PRIORITY_10 = [
    "fastapi-pro",
    "python-testing-patterns",
    "postgresql",
    "auth-implementation-patterns",
    "api-security-best-practices",
    "voice-ai-development",
    "documentation",
    "clean-code",
    "verification-before-completion",
    "plan-writing",
]


def build_document():
    doc = Document()
    doc.add_heading("Cursor Skills 适用清单", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        f"项目：Cruise Employee English Assessment Platform"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        f"生成日期：{datetime.now().strftime('%Y-%m-%d %H:%M')}"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    doc.add_heading("建议优先启用的 10 个 Skill", level=1)
    for i, name in enumerate(PRIORITY_10, 1):
        doc.add_paragraph(f"{i}. {name}", style="List Number")
    doc.add_paragraph()

    doc.add_heading("按类别完整清单", level=1)
    for category_title, skills in SKILLS_BY_CATEGORY:
        doc.add_heading(category_title, level=2)
        for skill_name, desc in skills:
            p = doc.add_paragraph()
            p.add_run(skill_name).bold = True
            p.add_run(f" — {desc}")
    doc.add_paragraph()
    doc.add_paragraph("以上 skill 名称对应 Cursor 中 .cursor/skills 下的技能目录名，可在 Cursor 设置中按名称启用。")
    return doc


def main():
    project_root = Path(__file__).resolve().parent.parent
    docs_dir = project_root / "docs"
    docs_dir.mkdir(exist_ok=True)
    out_path = docs_dir / "Cursor_Skills_For_Project.docx"
    doc = build_document()
    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    main()
