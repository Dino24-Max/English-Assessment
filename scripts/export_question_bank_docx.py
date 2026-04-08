"""
Export all questions and correct answers to a Word document (excluding speaking module).

Run from project root:
  python scripts/export_question_bank_docx.py

Output (default): output/question_bank_answers_no_speaking.docx
Requires: python-docx, local SQLite DB at data/assessment.db (same as app default).
"""
from __future__ import annotations

import json
import os
import sqlite3
import sys
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt

project_root = Path(__file__).resolve().parent.parent

MODULE_LABELS = {
    "listening": "听力 Listening",
    "time_numbers": "时间与数字 Time & Numbers",
    "grammar": "语法 Grammar",
    "vocabulary": "词汇 Vocabulary",
    "reading": "阅读 Reading",
}

DIVISION_LABELS = {
    "hotel": "酒店 Hotel",
    "marine": "海事 Marine",
    "casino": "赌场 Casino",
}

QUESTION_TYPE_LABELS = {
    "multiple_choice": "单选题",
    "fill_blank": "填空",
    "category_match": "分类匹配",
    "title_selection": "标题选择",
    "speaking_response": "口语",
}


def _db_path() -> Path:
    url = os.getenv("DATABASE_URL", "sqlite+aiosqlite:///./data/assessment.db")
    if "sqlite" not in url:
        raise SystemExit(
            "This exporter only supports SQLite. Set DATABASE_URL to a sqlite URL or use the default."
        )
    path_part = url.split("///", 1)[-1].strip()
    p = Path(path_part)
    if not p.is_absolute():
        p = project_root / p
    return p


def _format_options(raw) -> str:
    if raw is None or raw == "":
        return ""
    data = raw
    if isinstance(data, str):
        try:
            data = json.loads(data)
        except json.JSONDecodeError:
            return data
    if isinstance(data, list):
        lines = []
        for i, opt in enumerate(data):
            letter = chr(ord("A") + i) if i < 26 else str(i + 1)
            lines.append(f"{letter}. {opt}")
        return "\n".join(lines)
    if isinstance(data, dict):
        return "\n".join(f"{k}: {v}" for k, v in data.items())
    return str(data)


def _short_meta(raw) -> str:
    if raw is None or raw == "":
        return ""
    if isinstance(raw, str):
        try:
            data = json.loads(raw)
        except json.JSONDecodeError:
            return raw[:500]
    else:
        data = raw
    if not data:
        return ""
    text = json.dumps(data, ensure_ascii=False, indent=2)
    if len(text) > 1200:
        return text[:1200] + "\n…"
    return text


def main() -> None:
    db_file = _db_path()
    if not db_file.exists():
        raise SystemExit(f"Database not found: {db_file}\nStart the app once or run load_question_bank.py first.")

    out_dir = project_root / "output"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "question_bank_answers_no_speaking.docx"

    conn = sqlite3.connect(str(db_file))
    conn.row_factory = sqlite3.Row
    cur = conn.cursor()
    cur.execute(
        """
        SELECT
            id, module_type, division, question_type, question_text, options,
            correct_answer, audio_file_path, department, scenario_id,
            scenario_description, cefr_level, difficulty_level, points,
            question_metadata, is_safety_related
        FROM questions
        WHERE module_type != 'speaking'
        ORDER BY
            CASE module_type
                WHEN 'listening' THEN 1
                WHEN 'time_numbers' THEN 2
                WHEN 'grammar' THEN 3
                WHEN 'vocabulary' THEN 4
                WHEN 'reading' THEN 5
                ELSE 99
            END,
            CASE division
                WHEN 'hotel' THEN 1
                WHEN 'marine' THEN 2
                WHEN 'casino' THEN 3
                ELSE 99
            END,
            COALESCE(department, ''),
            id
        """
    )
    rows = cur.fetchall()
    conn.close()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(10)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    title = doc.add_heading("邮轮员工英语测评 — 题库与参考答案", level=0)
    title.runs[0].font.size = Pt(18)

    sub = doc.add_paragraph()
    sub.add_run("说明：本文件包含除「口语 Speaking」模块外的全部题目与系统记录的正确答案，便于审题与校对。")

    meta = doc.add_paragraph()
    meta.add_run(
        f"导出时间（UTC）：{datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S')}  |  "
        f"题目条数：{len(rows)}"
    )

    doc.add_paragraph()

    current_module = None
    current_division = None
    index_in_doc = 0

    for row in rows:
        r = dict(row)
        mod = r["module_type"]
        div = r["division"]
        if mod != current_module:
            current_module = mod
            current_division = None
            label = MODULE_LABELS.get(mod, mod)
            doc.add_heading(label, level=1)
        if div != current_division:
            current_division = div
            dlabel = DIVISION_LABELS.get(div, div or "")
            if dlabel:
                doc.add_heading(dlabel, level=2)

        index_in_doc += 1
        qt = QUESTION_TYPE_LABELS.get(r["question_type"], r["question_type"])

        p0 = doc.add_paragraph()
        p0.add_run(f"【{index_in_doc}】").bold = True
        p0.add_run(f" 数据库ID: {r['id']}  |  题型: {qt}  |  分值: {r['points']}")
        if r.get("cefr_level"):
            p0.add_run(f"  |  CEFR: {r['cefr_level']}")

        dept = r.get("department") or ""
        if dept:
            doc.add_paragraph(f"部门场景: {dept}" + (f" (#{r['scenario_id']})" if r.get("scenario_id") else ""))

        if r.get("scenario_description"):
            doc.add_paragraph(f"场景说明: {r['scenario_description']}")

        doc.add_paragraph(f"题干:\n{r['question_text']}")

        opts = _format_options(r["options"])
        if opts:
            doc.add_paragraph(f"选项:\n{opts}")

        if r.get("audio_file_path"):
            doc.add_paragraph(f"音频路径: {r['audio_file_path']}")

        ans = doc.add_paragraph()
        ans.add_run("正确答案: ").bold = True
        ans.add_run(r["correct_answer"] or "")

        extras = []
        if r.get("difficulty_level") is not None:
            extras.append(f"难度 {r['difficulty_level']}")
        if r.get("is_safety_related"):
            extras.append("安全相关题")
        if extras:
            doc.add_paragraph(" | ".join(extras))

        meta_txt = _short_meta(r.get("question_metadata"))
        if meta_txt:
            doc.add_paragraph(f"附加元数据:\n{meta_txt}")

        doc.add_paragraph("—" * 40)

    doc.save(str(out_path))
    print(f"Wrote {len(rows)} questions to:\n  {out_path}")


if __name__ == "__main__":
    try:
        main()
    except SystemExit as e:
        print(e, file=sys.stderr)
        raise
