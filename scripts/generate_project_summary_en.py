"""
Generate comprehensive project summary document with Snowflake recommendations (English Version)
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from datetime import datetime


def add_heading(doc, text, level=1):
    """Add formatted heading"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading


def add_paragraph(doc, text, bold=False, font_size=11):
    """Add formatted paragraph"""
    para = doc.add_paragraph(text)
    run = para.runs[0]
    run.font.size = Pt(font_size)
    run.font.name = 'Calibri'
    if bold:
        run.bold = True
    return para


def add_bullet_point(doc, text, level=0):
    """Add bullet point"""
    para = doc.add_paragraph(text, style='List Bullet')
    if level > 0:
        para.paragraph_format.left_indent = Inches(0.5 * level)
    return para


def add_table_with_header(doc, headers, rows):
    """Add formatted table"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'

    # Add headers
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)

    # Add data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)

    return table


def generate_project_summary():
    """Generate comprehensive project summary document"""

    doc = Document()

    # Set document properties
    doc.core_properties.title = "Cruise Employee English Assessment Platform - Project Summary & Recommendations"
    doc.core_properties.author = "Project Team"
    doc.core_properties.created = datetime.now()

    # ==================== TITLE PAGE ====================
    title = doc.add_heading('Cruise Employee English Assessment Platform', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading('AI-Powered English Proficiency Testing System', 2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_para.add_run('Project Summary & Snowflake Data Platform Integration Strategy\n').bold = True
    info_para.add_run(f'Generated: {datetime.now().strftime("%B %d, %Y")}\n')
    info_para.add_run('Version: 1.0')

    doc.add_page_break()

    # ==================== EXECUTIVE SUMMARY ====================
    add_heading(doc, '1. Executive Summary', 1)

    add_paragraph(doc, '1.1 Project Background', bold=True, font_size=12)
    add_paragraph(doc,
        'The Cruise Employee English Assessment Platform is a comprehensive AI-driven English proficiency '
        'evaluation system designed specifically for the cruise industry. It provides standardized, intelligent '
        'English proficiency testing across 16 specific positions in three major divisions: Hotel Operations, '
        'Marine Operations, and Casino Operations, ensuring cruise employees meet international maritime '
        'communication requirements.')

    add_paragraph(doc, '1.2 Core Features', bold=True, font_size=12)
    features = [
        '6 Assessment Modules: Listening, Time & Numbers, Grammar, Vocabulary, Reading, Speaking',
        '288 department-specific questions covering 16 positions across professional scenarios',
        'AI-powered intelligent scoring system with speech recognition and NLP',
        '100-point scoring system with mandatory pass requirements for safety questions',
        'Real-time performance analysis and detailed feedback reports',
        'Anti-cheating mechanisms: session management, IP tracking, time limits'
    ]
    for feature in features:
        add_bullet_point(doc, feature)

    add_paragraph(doc, '1.3 Project Statistics', bold=True, font_size=12)
    stats_data = [
        ['Total Assessment Questions', '288 (96 per division)'],
        ['Speaking Scenarios', '160 workplace situations'],
        ['Departments Covered', '16 (8 Hotel, 3 Marine, 3 Casino, 2 Safety)'],
        ['Assessment Modules', '6 core modules'],
        ['Maximum Score', '100 points (20 from speaking)'],
        ['Passing Criteria', '70% overall + all safety questions + 60% speaking minimum']
    ]
    add_table_with_header(doc, ['Metric', 'Value/Description'], stats_data)

    doc.add_page_break()

    # ==================== TECHNICAL ARCHITECTURE ====================
    add_heading(doc, '2. Technical Architecture', 1)

    add_paragraph(doc, '2.1 Current Technology Stack', bold=True, font_size=12)

    add_paragraph(doc, 'Backend Framework & Database:', bold=True)
    backend_tech = [
        'FastAPI 0.104.1 - High-performance async web framework',
        'Python 3.10+ - Primary development language',
        'SQLAlchemy 2.0.23 - Async ORM framework',
        'PostgreSQL - Relational database (current)',
        'asyncpg & psycopg2-binary - Database drivers',
        'Alembic 1.12.1 - Database migration tool'
    ]
    for tech in backend_tech:
        add_bullet_point(doc, tech)

    add_paragraph(doc, 'AI & Machine Learning:', bold=True)
    ai_tech = [
        'OpenAI API 1.3.7 - GPT model integration',
        'Anthropic Claude 0.7.7 - Claude AI integration',
        'librosa 0.10.1 - Audio processing',
        'scikit-learn 1.3.2 - Machine learning algorithms',
        'numpy & pandas - Data processing'
    ]
    for tech in ai_tech:
        add_bullet_point(doc, tech)

    add_paragraph(doc, 'Security & Session Management:', bold=True)
    security_tech = [
        'python-jose - JWT authentication',
        'passlib[bcrypt] - Password hashing',
        'Redis 5.0.1 - Session storage and caching',
        'Celery 5.3.4 - Background task queue'
    ]
    for tech in security_tech:
        add_bullet_point(doc, tech)

    add_paragraph(doc, '2.2 Current Data Models', bold=True, font_size=12)
    models = [
        'User - Candidate information (name, email, nationality, department)',
        'Question - Question bank (288 questions with module type, difficulty, correct answers)',
        'Assessment - Assessment session (status, scores, timestamps, anti-cheating info)',
        'AssessmentResponse - Answer records (user answers, scores, time spent, speech analysis)',
        'DivisionDepartment - Department mapping',
        'AssessmentConfig - System configuration'
    ]
    for model in models:
        add_bullet_point(doc, model)

    add_paragraph(doc, '2.3 Current System Architecture', bold=True, font_size=12)
    architecture = [
        'Three-tier architecture: API Layer → Business Logic Layer → Data Access Layer',
        'RESTful API design with CRUD operations',
        'Asynchronous processing for improved concurrency',
        'Jinja2 template engine for frontend rendering',
        'Docker containerization support',
        'Nginx reverse proxy'
    ]
    for item in architecture:
        add_bullet_point(doc, item)

    doc.add_page_break()

    # ==================== PROJECT STRUCTURE ====================
    add_heading(doc, '3. Project Structure Analysis', 1)

    add_paragraph(doc, '3.1 Core Modules', bold=True, font_size=12)

    modules_data = [
        ['Module Path', 'Function', 'Key Files'],
        ['api/routes/', 'API Routes', 'assessment.py, admin.py, analytics.py, ui.py'],
        ['core/', 'Core Logic', 'config.py, database.py, assessment_engine.py'],
        ['models/', 'Data Models', 'base.py, assessment.py'],
        ['services/', 'External Services', 'ai_service.py (AI integration)'],
        ['utils/', 'Utilities', 'scoring.py, anti_cheating.py'],
        ['middleware/', 'Middleware', 'security.py, session.py'],
        ['data/', 'Data Management', 'question_bank_loader.py, generate_question_bank.py'],
        ['training/', 'ML Training', 'speech_trainer.py'],
        ['evaluation/', 'Model Evaluation', 'model_evaluator.py'],
        ['inference/', 'Inference Services', 'speech_inference.py']
    ]
    add_table_with_header(doc, modules_data[0], modules_data[1:])

    add_paragraph(doc, '3.2 Assessment Workflow', bold=True, font_size=12)
    workflow = [
        '1. Candidate Registration → POST /api/v1/assessment/register',
        '2. Create Assessment Session → POST /api/v1/assessment/create',
        '3. Start Assessment → POST /api/v1/assessment/{id}/start',
        '4. Submit Answers → POST /api/v1/assessment/{id}/answer',
        '5. Submit Speaking Response → POST /api/v1/assessment/{id}/speaking',
        '6. Complete Assessment → POST /api/v1/assessment/{id}/complete',
        '7. Retrieve Results → GET /api/v1/assessment/{id}/status'
    ]
    for step in workflow:
        add_bullet_point(doc, step)

    doc.add_page_break()

    # ==================== SNOWFLAKE INTEGRATION ====================
    add_heading(doc, '4. Snowflake Data Platform Integration Strategy', 1)

    add_paragraph(doc, '4.1 Why Snowflake?', bold=True, font_size=12)
    reasons = [
        'Cloud-Native Architecture: Fully managed cloud data warehouse, no infrastructure maintenance',
        'Elastic Scalability: Scale compute and storage on-demand, support massive concurrent queries',
        'Data Sharing: Secure cross-organizational data sharing, ideal for multi-department collaboration',
        'High Performance: Automatic query optimization, supports large-scale data analytics',
        'Cost Optimization: Separate storage and compute, per-second billing, cost-effective',
        'Data Governance: Built-in data governance, security, and compliance features',
        'ML Integration: Supports Snowpark Python for in-database ML model execution'
    ]
    for reason in reasons:
        add_bullet_point(doc, reason)

    add_paragraph(doc, '4.2 Integration Architecture Design', bold=True, font_size=12)
    add_paragraph(doc, 'Recommended Hybrid Architecture: PostgreSQL (Transactional) + Snowflake (Analytical)')

    architecture_design = [
        'PostgreSQL Layer:',
        '  - Store real-time transactional data (user registration, assessment sessions, responses)',
        '  - Handle OLTP operations (Online Transaction Processing)',
        '  - Ensure data consistency and ACID properties',
        '',
        'Snowflake Layer:',
        '  - Store historical and aggregated data',
        '  - Support complex analytical queries (OLAP)',
        '  - Generate BI reports and data visualizations',
        '  - Machine learning model training and inference',
        '',
        'Data Synchronization:',
        '  - Real-time sync using CDC (Change Data Capture)',
        '  - Scheduled batch ETL jobs',
        '  - Snowpipe streaming data ingestion'
    ]
    for item in architecture_design:
        if item.startswith('  '):
            add_bullet_point(doc, item.strip(), level=1)
        elif item:
            add_paragraph(doc, item, bold=True)

    add_paragraph(doc, '4.3 Technical Implementation Plan', bold=True, font_size=12)

    add_paragraph(doc, 'Step 1: Install Snowflake Dependencies', bold=True)
    code_para = doc.add_paragraph('pip install snowflake-sqlalchemy snowflake-connector-python')
    code_para.runs[0].font.name = 'Courier New'
    code_para.runs[0].font.size = Pt(10)

    add_paragraph(doc, 'Step 2: Configure Connection Strings', bold=True)
    connection_config = [
        'Add Snowflake configuration to .env file:',
        'SNOWFLAKE_USER=your_username',
        'SNOWFLAKE_PASSWORD=your_password',
        'SNOWFLAKE_ACCOUNT=your_account_identifier',
        'SNOWFLAKE_WAREHOUSE=COMPUTE_WH',
        'SNOWFLAKE_DATABASE=CRUISE_ASSESSMENT',
        'SNOWFLAKE_SCHEMA=PUBLIC'
    ]
    for config in connection_config:
        config_para = doc.add_paragraph(config)
        config_para.runs[0].font.name = 'Courier New'
        config_para.runs[0].font.size = Pt(10)

    add_paragraph(doc, 'Step 3: Create Snowflake Database Manager', bold=True)
    add_paragraph(doc, 'Create connection manager in src/main/python/core/snowflake_db.py')

    add_paragraph(doc, 'Step 4: Design Snowflake Data Model', bold=True)
    snowflake_tables = [
        'FACT_ASSESSMENTS - Assessment fact table',
        'FACT_RESPONSES - Response fact table',
        'DIM_USERS - User dimension table',
        'DIM_QUESTIONS - Question dimension table',
        'DIM_DIVISIONS - Division dimension table',
        'DIM_DATE - Date dimension table',
        'AGG_DAILY_STATS - Daily statistics aggregate table',
        'AGG_DIVISION_PERFORMANCE - Division performance aggregate table'
    ]
    for table in snowflake_tables:
        add_bullet_point(doc, table)

    doc.add_page_break()

    # ==================== DEVELOPMENT RECOMMENDATIONS ====================
    add_heading(doc, '5. Development Recommendations', 1)

    add_paragraph(doc, '5.1 Data Architecture Optimization', bold=True, font_size=12)

    recommendations = [
        ['Priority', 'Recommendation', 'Description', 'Expected Benefits'],
        ['HIGH', 'Implement Snowflake Integration',
         'Migrate historical and analytical data to Snowflake, retain PostgreSQL for real-time transactions',
         '10x+ improvement in analytical query performance, reduced database load'],
        ['HIGH', 'Implement Data Warehouse Model',
         'Adopt star schema design with fact and dimension tables',
         'Optimize complex queries, enable multi-dimensional analysis'],
        ['MEDIUM', 'Establish CDC Data Pipeline',
         'Use Debezium or similar tools for real-time data synchronization',
         'Ensure data consistency, reduce sync latency'],
        ['MEDIUM', 'Implement Snowpipe Streaming Ingestion',
         'Automate data loading processes',
         'Reduce manual ETL workload, enable near real-time analytics'],
        ['LOW', 'Data Lake Integration',
         'Store raw audio files in S3/Azure Blob, access via Snowflake External Tables',
         'Lower storage costs, support unstructured data analytics']
    ]
    add_table_with_header(doc, recommendations[0], recommendations[1:])

    add_paragraph(doc, '5.2 Performance Optimization', bold=True, font_size=12)
    performance = [
        'Use Snowflake clustering keys to optimize large table queries',
        'Implement automatic clustering to maintain data order',
        'Use materialized views to cache complex query results',
        'Enable query result cache to reduce redundant computations',
        'Properly size virtual warehouses, use dynamic scaling',
        'Implement partitioning strategy, partition large tables by time or division'
    ]
    for item in performance:
        add_bullet_point(doc, item)

    add_paragraph(doc, '5.3 Security & Compliance', bold=True, font_size=12)
    security = [
        'Enable end-to-end encryption',
        'Implement role-based access control (RBAC)',
        'Use Snowflake dynamic data masking',
        'Configure network policies to restrict access sources',
        'Enable audit logging to track all data access',
        'Implement multi-factor authentication (MFA)',
        'Conduct regular security audits and vulnerability scans'
    ]
    for item in security:
        add_bullet_point(doc, item)

    add_paragraph(doc, '5.4 Machine Learning Enhancements', bold=True, font_size=12)
    ml_enhancements = [
        'Use Snowpark Python to train ML models within Snowflake',
        'Implement UDFs (User-Defined Functions) for real-time scoring',
        'Integrate Snowflake ML prediction features to predict candidate performance',
        'Use Snowflake Cortex for AI-driven data analytics',
        'Build feature engineering pipelines to extract assessment data features',
        'Implement A/B testing framework to optimize assessment processes'
    ]
    for item in ml_enhancements:
        add_bullet_point(doc, item)

    doc.add_page_break()

    # ==================== IMPLEMENTATION ROADMAP ====================
    add_heading(doc, '6. Implementation Roadmap', 1)

    phases = [
        ['Phase', 'Tasks', 'Duration', 'Key Deliverables'],
        ['Phase 1\nPreparation',
         '- Assess current data scale\n- Design Snowflake data model\n- Set up development environment\n- Create POC validation',
         '2-3 weeks',
         '- Technical proposal\n- POC demonstration\n- Cost estimation'],
        ['Phase 2\nDevelopment',
         '- Develop Snowflake connector\n- Implement data migration scripts\n- Create fact and dimension tables\n- Develop ETL pipeline',
         '4-6 weeks',
         '- Data pipeline code\n- Unit tests\n- Technical documentation'],
        ['Phase 3\nIntegration',
         '- Integrate analytics API\n- Develop BI reports\n- Implement data synchronization\n- Performance testing',
         '3-4 weeks',
         '- API endpoints\n- Report dashboards\n- Performance benchmarks'],
        ['Phase 4\nOptimization',
         '- Performance tuning\n- Security hardening\n- Documentation completion\n- Team training',
         '2-3 weeks',
         '- Optimization report\n- User manuals\n- Training materials'],
        ['Phase 5\nDeployment',
         '- Production environment deployment\n- Data migration\n- Monitoring configuration\n- Go-live support',
         '1-2 weeks',
         '- Production system\n- Monitoring dashboard\n- Operations manual']
    ]
    add_table_with_header(doc, phases[0], phases[1:])

    add_paragraph(doc, 'Estimated Total Duration: 12-18 weeks (approximately 3-4.5 months)', bold=True)

    doc.add_page_break()

    # ==================== COST ANALYSIS ====================
    add_heading(doc, '7. Cost Analysis', 1)

    add_paragraph(doc, '7.1 Snowflake Cost Components', bold=True, font_size=12)
    cost_components = [
        'Compute Costs (Virtual Warehouse)',
        '  - Per-second billing, 1-minute minimum',
        '  - X-Small: $2/hour, Small: $4/hour, Medium: $8/hour',
        '  - Recommendation: Use X-Small for dev, Small-Medium for production',
        '',
        'Storage Costs',
        '  - $23-40/TB/month (varies by region)',
        '  - Automatic compression typically saves 50-70% storage space',
        '',
        'Data Transfer Costs',
        '  - Free within region',
        '  - Cross-region transfers incur charges'
    ]
    for item in cost_components:
        if item.startswith('  '):
            add_bullet_point(doc, item.strip(), level=1)
        elif item:
            add_paragraph(doc, item, bold=True)

    add_paragraph(doc, '7.2 Estimated Monthly Costs (Medium Scale)', bold=True, font_size=12)
    cost_estimate = [
        ['Item', 'Scale', 'Estimated Cost (USD/month)'],
        ['Compute Resources', 'Small warehouse, 8 hours/day', '$960 (30 days x 8 hours x $4)'],
        ['Storage', '5TB data (after compression)', '$115-200'],
        ['Data Transfer', 'Within region', '$0'],
        ['Snowpipe', 'Continuous ingestion', '$50-100'],
        ['Total', '', 'Approximately $1,125-1,260/month']
    ]
    add_table_with_header(doc, cost_estimate[0], cost_estimate[1:])

    add_paragraph(doc, '7.3 Cost Optimization Strategies', bold=True, font_size=12)
    cost_optimization = [
        'Use auto-suspend feature to automatically shut down idle warehouses',
        'Properly size warehouses to avoid over-provisioning',
        'Use query tags to track cost attribution',
        'Enable resource monitors to control budgets',
        'Regularly clean up expired data and temporary tables',
        'Use Snowflake Cost Attribution to track costs'
    ]
    for item in cost_optimization:
        add_bullet_point(doc, item)

    doc.add_page_break()

    # ==================== RISKS AND MITIGATION ====================
    add_heading(doc, '8. Risks & Mitigation Strategies', 1)

    risks = [
        ['Risk Category', 'Specific Risk', 'Impact', 'Mitigation Strategy'],
        ['Technical', 'Data migration failure or data loss', 'High',
         '- Phased migration\n- Complete backups\n- Parallel validation'],
        ['Technical', 'Performance below expectations', 'Medium',
         '- POC performance testing\n- Use Snowflake optimization recommendations\n- Expert consultation'],
        ['Operational', 'Team lacks Snowflake experience', 'Medium',
         '- Official Snowflake training\n- Hire consultants\n- Gradual migration with learning'],
        ['Financial', 'Costs exceed budget', 'Medium',
         '- Set up resource monitors\n- Regular cost reviews\n- Optimize query efficiency'],
        ['Security', 'Data breach or unauthorized access', 'High',
         '- Implement RBAC\n- Enable encryption\n- Regular security audits'],
        ['Business', 'Impact on existing operations', 'High',
         '- Gradual rollout\n- Rollback plan\n- 24/7 monitoring']
    ]
    add_table_with_header(doc, risks[0], risks[1:])

    doc.add_page_break()

    # ==================== SUCCESS METRICS ====================
    add_heading(doc, '9. Success Metrics', 1)

    add_paragraph(doc, '9.1 Technical Metrics', bold=True, font_size=12)
    tech_metrics = [
        'Query Performance: Complex analytical queries respond in < 5 seconds',
        'Data Sync Latency: < 5 minutes',
        'System Availability: > 99.9%',
        'Data Consistency: 100% (via automated validation)',
        'API Response Time: P95 < 200ms',
        'Concurrent Support: > 100 concurrent queries'
    ]
    for metric in tech_metrics:
        add_bullet_point(doc, metric)

    add_paragraph(doc, '9.2 Business Metrics', bold=True, font_size=12)
    business_metrics = [
        'Report generation time reduced by 70%',
        'Data analyst efficiency improved by 50%',
        'Data storage costs reduced by 30%',
        'Concurrent assessment capacity increased 3x',
        'New analytics requirement delivery time reduced by 60%',
        'BI dashboard user satisfaction > 4.5/5'
    ]
    for metric in business_metrics:
        add_bullet_point(doc, metric)

    add_paragraph(doc, '9.3 User Experience Metrics', bold=True, font_size=12)
    ux_metrics = [
        'Admin report query satisfaction > 90%',
        'Data analyst training completion rate > 95%',
        'Report error rate < 1%',
        'Self-service analytics adoption rate > 70%'
    ]
    for metric in ux_metrics:
        add_bullet_point(doc, metric)

    doc.add_page_break()

    # ==================== NEXT STEPS ====================
    add_heading(doc, '10. Next Steps & Action Plan', 1)

    add_paragraph(doc, '10.1 Immediate Actions (Next 1 Month)', bold=True, font_size=12)
    immediate_actions = [
        'Form Snowflake integration project team',
        'Apply for Snowflake trial account',
        'Complete POC environment setup',
        'Assess current data scale and growth trends',
        'Design initial data warehouse model',
        'Develop detailed project plan and budget'
    ]
    for action in immediate_actions:
        add_bullet_point(doc, action)

    add_paragraph(doc, '10.2 Short-Term Actions (1-3 Months)', bold=True, font_size=12)
    short_term_actions = [
        'Complete Snowflake environment configuration',
        'Develop data synchronization pipeline',
        'Migrate historical data to Snowflake',
        'Develop analytics API and BI reports',
        'Conduct performance and stress testing',
        'Team training and knowledge transfer'
    ]
    for action in short_term_actions:
        add_bullet_point(doc, action)

    add_paragraph(doc, '10.3 Medium-Term Actions (3-6 Months)', bold=True, font_size=12)
    medium_term_actions = [
        'Production environment go-live',
        'Implement advanced analytics features',
        'Integrate machine learning models',
        'Optimize costs and performance',
        'Expand data source integrations',
        'Establish data governance framework'
    ]
    for action in medium_term_actions:
        add_bullet_point(doc, action)

    doc.add_page_break()

    # ==================== CONCLUSION ====================
    add_heading(doc, '11. Conclusion', 1)

    add_paragraph(doc,
        'The Cruise Employee English Assessment Platform has established a solid technical foundation. '
        'Through the integration of FastAPI, PostgreSQL, and AI services, it successfully delivers '
        'intelligent English proficiency assessment. The current system architecture is clear, code quality '
        'is good, and functional modules are complete.')

    doc.add_paragraph()

    add_paragraph(doc,
        'Integrating the Snowflake data platform will bring significant improvements:', bold=True)

    conclusion_benefits = [
        'Analytical Capabilities: Support complex multi-dimensional data analysis and BI reporting',
        'Performance Boost: Dramatically improve query performance, especially for aggregation scenarios',
        'Data Insights: Discover assessment trends and optimization opportunities through advanced analytics',
        'Enterprise-Grade: Enhanced data governance, security, and compliance capabilities',
        'Cost Optimization: Lower total cost through storage-compute separation and pay-per-use',
        'Future-Ready: Prepared for AI/ML model training and large-scale expansion'
    ]
    for benefit in conclusion_benefits:
        add_bullet_point(doc, benefit)

    doc.add_paragraph()

    add_paragraph(doc,
        'We recommend adopting a phased implementation strategy, starting with POC validation and gradually '
        'migrating data and functionality to ensure a smooth transition. Through proper cost control and '
        'performance optimization, Snowflake integration will bring long-term value to the Cruise Employee '
        'English Assessment Platform, supporting continuous business growth and data-driven decision-making.')

    doc.add_paragraph()
    doc.add_paragraph()

    final_para = doc.add_paragraph()
    final_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    final_para.add_run('--- End of Report ---\n').bold = True
    final_para.add_run(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')

    # ==================== APPENDIX ====================
    doc.add_page_break()
    add_heading(doc, 'Appendix A: Technical References', 1)

    references = [
        'Snowflake Documentation: https://docs.snowflake.com/',
        'Snowflake Python Connector: https://docs.snowflake.com/en/developer-guide/python-connector/',
        'Snowflake SQLAlchemy: https://docs.snowflake.com/en/developer-guide/python-connector/sqlalchemy',
        'FastAPI Documentation: https://fastapi.tiangolo.com/',
        'SQLAlchemy Async: https://docs.sqlalchemy.org/en/20/orm/extensions/asyncio.html',
        'Snowpark Python: https://docs.snowflake.com/en/developer-guide/snowpark/python/index',
        'Snowflake Cost Optimization: https://www.snowflake.com/en/data-cloud/workloads/cost-optimization/',
        'Snowflake Security Best Practices: https://docs.snowflake.com/en/user-guide/security-best-practices'
    ]
    for ref in references:
        add_bullet_point(doc, ref)

    doc.add_page_break()
    add_heading(doc, 'Appendix B: Project Files Inventory', 1)

    key_files = [
        'src/main/python/core/database.py - Database connection management',
        'src/main/python/models/assessment.py - Data model definitions',
        'src/main/python/api/routes/assessment.py - Assessment API routes',
        'src/main/python/api/routes/analytics.py - Analytics API routes',
        'src/main/python/services/ai_service.py - AI service integration',
        'src/main/python/utils/scoring.py - Scoring algorithms',
        'src/main/python/data/generate_question_bank.py - Question bank generation',
        'requirements.txt - Python dependencies',
        'README.md - Project documentation',
        'CLAUDE.md - Development guidelines'
    ]
    for file in key_files:
        add_bullet_point(doc, file)

    # Save document
    output_path = "C:\\Users\\szh2051\\OneDrive - Carnival Corporation\\Desktop\\Python\\Claude Demo\\output\\Cruise_Assessment_Snowflake_Integration_Report_EN.docx"
    doc.save(output_path)
    print(f"Project summary document generated successfully: {output_path}")
    return output_path


if __name__ == "__main__":
    generate_project_summary()
