"""
Generate comprehensive project summary document with Snowflake recommendations
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from datetime import datetime


def add_heading(doc, text, level=1):
    """Add formatted heading"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading


def add_paragraph(doc, text, bold=False, font_size=11):
    """Add formatted paragraph"""
    para = doc.add_paragraph(text)
    run = para.runs[0]
    run.font.size = Pt(font_size)
    run.font.name = 'Calibri'
    if bold:
        run.bold = True
    return para


def add_bullet_point(doc, text, level=0):
    """Add bullet point"""
    para = doc.add_paragraph(text, style='List Bullet')
    if level > 0:
        para.paragraph_format.left_indent = Inches(0.5 * level)
    return para


def add_table_with_header(doc, headers, rows):
    """Add formatted table"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'

    # Add headers
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)

    # Add data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)

    return table


def generate_project_summary():
    """Generate comprehensive project summary document"""

    doc = Document()

    # Set document properties
    doc.core_properties.title = "邮轮员工英语评估平台 - 项目总结与发展建议"
    doc.core_properties.author = "项目团队"
    doc.core_properties.created = datetime.now()

    # ==================== TITLE PAGE ====================
    title = doc.add_heading('邮轮员工英语评估平台', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading('Cruise Employee English Assessment Platform', 2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_para.add_run('项目总结与Snowflake数据平台集成建议\n').bold = True
    info_para.add_run(f'生成日期: {datetime.now().strftime("%Y年%m月%d日")}\n')
    info_para.add_run('版本: 1.0')

    doc.add_page_break()

    # ==================== EXECUTIVE SUMMARY ====================
    add_heading(doc, '一、项目概述 (Executive Summary)', 1)

    add_paragraph(doc, '1.1 项目背景', bold=True, font_size=12)
    add_paragraph(doc,
        '邮轮员工英语评估平台是一个专为邮轮行业设计的综合性AI驱动英语能力评估系统。'
        '该平台针对酒店运营、海事运营和赌场运营三大部门的16个具体岗位，提供标准化、'
        '智能化的英语水平测试服务，确保邮轮员工具备符合国际海事通信要求的英语能力。')

    add_paragraph(doc, '1.2 核心特性', bold=True, font_size=12)
    features = [
        '6大评估模块：听力、时间与数字、语法、词汇、阅读、口语',
        '288道分部门定制化题目，覆盖16个岗位的专业场景',
        'AI驱动的智能评分系统，包含语音识别和自然语言处理',
        '100分制评分体系，包含安全相关问题强制通过机制',
        '实时性能分析和详细反馈报告',
        '防作弊机制：会话管理、IP跟踪、时间限制'
    ]
    for feature in features:
        add_bullet_point(doc, feature)

    add_paragraph(doc, '1.3 项目统计', bold=True, font_size=12)
    stats_data = [
        ['评估题目总数', '288题 (每个部门96题)'],
        ['口语场景数', '160个工作场景'],
        ['覆盖部门', '16个 (酒店8个、海事3个、赌场3个、安全2个)'],
        ['评估模块', '6个核心模块'],
        ['总分', '100分 (口语模块20分)'],
        ['及格标准', '总分70% + 安全问题全对 + 口语60%以上']
    ]
    add_table_with_header(doc, ['项目指标', '数值/描述'], stats_data)

    doc.add_page_break()

    # ==================== TECHNICAL ARCHITECTURE ====================
    add_heading(doc, '二、技术架构 (Technical Architecture)', 1)

    add_paragraph(doc, '2.1 当前技术栈', bold=True, font_size=12)

    add_paragraph(doc, '后端框架与数据库：', bold=True)
    backend_tech = [
        'FastAPI 0.104.1 - 高性能异步Web框架',
        'Python 3.10+ - 主开发语言',
        'SQLAlchemy 2.0.23 - 异步ORM框架',
        'PostgreSQL - 关系型数据库 (当前)',
        'asyncpg & psycopg2-binary - 数据库驱动',
        'Alembic 1.12.1 - 数据库迁移工具'
    ]
    for tech in backend_tech:
        add_bullet_point(doc, tech)

    add_paragraph(doc, 'AI与机器学习：', bold=True)
    ai_tech = [
        'OpenAI API 1.3.7 - GPT模型集成',
        'Anthropic Claude 0.7.7 - Claude AI集成',
        'librosa 0.10.1 - 音频处理',
        'scikit-learn 1.3.2 - 机器学习算法',
        'numpy & pandas - 数据处理'
    ]
    for tech in ai_tech:
        add_bullet_point(doc, tech)

    add_paragraph(doc, '安全与会话管理：', bold=True)
    security_tech = [
        'python-jose - JWT认证',
        'passlib[bcrypt] - 密码哈希',
        'Redis 5.0.1 - 会话存储和缓存',
        'Celery 5.3.4 - 后台任务队列'
    ]
    for tech in security_tech:
        add_bullet_point(doc, tech)

    add_paragraph(doc, '2.2 当前数据模型', bold=True, font_size=12)
    models = [
        'User - 用户/候选人信息 (姓名、邮箱、国籍、部门)',
        'Question - 题库 (288道题目，包含模块类型、难度、正确答案)',
        'Assessment - 评估会话 (状态、分数、时间戳、防作弊信息)',
        'AssessmentResponse - 答题记录 (用户答案、得分、用时、语音分析)',
        'DivisionDepartment - 部门映射',
        'AssessmentConfig - 系统配置'
    ]
    for model in models:
        add_bullet_point(doc, model)

    add_paragraph(doc, '2.3 当前系统架构', bold=True, font_size=12)
    architecture = [
        '三层架构：API层 → 业务逻辑层 → 数据访问层',
        'RESTful API设计，支持CRUD操作',
        '异步处理机制，提高并发性能',
        'Jinja2模板引擎渲染前端页面',
        'Docker容器化部署支持',
        'Nginx反向代理'
    ]
    for item in architecture:
        add_bullet_point(doc, item)

    doc.add_page_break()

    # ==================== PROJECT STRUCTURE ====================
    add_heading(doc, '三、项目结构分析 (Project Structure)', 1)

    add_paragraph(doc, '3.1 核心模块', bold=True, font_size=12)

    modules_data = [
        ['api/routes/', 'API路由', 'assessment.py, admin.py, analytics.py, ui.py'],
        ['core/', '核心逻辑', 'config.py, database.py, assessment_engine.py'],
        ['models/', '数据模型', 'base.py, assessment.py'],
        ['services/', '外部服务', 'ai_service.py (AI集成)'],
        ['utils/', '工具函数', 'scoring.py, anti_cheating.py'],
        ['middleware/', '中间件', 'security.py, session.py'],
        ['data/', '数据管理', 'question_bank_loader.py, generate_question_bank.py'],
        ['training/', 'ML训练', 'speech_trainer.py'],
        ['evaluation/', '模型评估', 'model_evaluator.py'],
        ['inference/', '推理服务', 'speech_inference.py']
    ]
    add_table_with_header(doc, ['模块路径', '功能', '关键文件'], modules_data)

    add_paragraph(doc, '3.2 评估流程', bold=True, font_size=12)
    workflow = [
        '1. 候选人注册 → POST /api/v1/assessment/register',
        '2. 创建评估会话 → POST /api/v1/assessment/create',
        '3. 开始评估 → POST /api/v1/assessment/{id}/start',
        '4. 答题提交 → POST /api/v1/assessment/{id}/answer',
        '5. 口语答题 → POST /api/v1/assessment/{id}/speaking',
        '6. 完成评估 → POST /api/v1/assessment/{id}/complete',
        '7. 获取结果 → GET /api/v1/assessment/{id}/status'
    ]
    for step in workflow:
        add_bullet_point(doc, step)

    doc.add_page_break()

    # ==================== SNOWFLAKE INTEGRATION ====================
    add_heading(doc, '四、Snowflake数据平台集成方案 (Snowflake Integration)', 1)

    add_paragraph(doc, '4.1 为什么选择Snowflake？', bold=True, font_size=12)
    reasons = [
        '云原生架构：完全托管的云数据仓库，无需维护基础设施',
        '弹性扩展：按需扩展计算和存储资源，支持大规模并发查询',
        '数据共享：支持跨组织的安全数据共享，适合多部门协作',
        '高性能：自动优化查询，支持大规模数据分析',
        '成本优化：存储和计算分离，按秒计费，成本可控',
        '数据治理：内置数据治理、安全和合规功能',
        '机器学习集成：支持Snowpark Python，可在数据库内执行ML模型'
    ]
    for reason in reasons:
        add_bullet_point(doc, reason)

    add_paragraph(doc, '4.2 集成架构设计', bold=True, font_size=12)
    add_paragraph(doc, '建议采用混合架构：PostgreSQL (事务性数据) + Snowflake (分析性数据)')

    architecture_design = [
        'PostgreSQL层：',
        '  - 存储实时事务数据（用户注册、评估会话、答题记录）',
        '  - 处理OLTP操作（在线事务处理）',
        '  - 保证数据一致性和ACID特性',
        '',
        'Snowflake层：',
        '  - 存储历史数据和聚合数据',
        '  - 支持复杂分析查询（OLAP）',
        '  - 生成BI报表和数据可视化',
        '  - 机器学习模型训练和推理',
        '',
        '数据同步：',
        '  - 使用CDC (Change Data Capture) 实时同步',
        '  - 定期批量ETL任务',
        '  - Snowpipe流式数据摄取'
    ]
    for item in architecture_design:
        if item.startswith('  '):
            add_bullet_point(doc, item.strip(), level=1)
        elif item:
            add_paragraph(doc, item, bold=True)

    add_paragraph(doc, '4.3 技术实施方案', bold=True, font_size=12)

    add_paragraph(doc, 'Step 1: 安装Snowflake依赖', bold=True)
    code_para = doc.add_paragraph('pip install snowflake-sqlalchemy snowflake-connector-python')
    code_para.runs[0].font.name = 'Courier New'
    code_para.runs[0].font.size = Pt(10)

    add_paragraph(doc, 'Step 2: 配置连接字符串', bold=True)
    connection_config = [
        '在 .env 文件中添加Snowflake配置：',
        'SNOWFLAKE_USER=your_username',
        'SNOWFLAKE_PASSWORD=your_password',
        'SNOWFLAKE_ACCOUNT=your_account_identifier',
        'SNOWFLAKE_WAREHOUSE=COMPUTE_WH',
        'SNOWFLAKE_DATABASE=CRUISE_ASSESSMENT',
        'SNOWFLAKE_SCHEMA=PUBLIC'
    ]
    for config in connection_config:
        config_para = doc.add_paragraph(config)
        config_para.runs[0].font.name = 'Courier New'
        config_para.runs[0].font.size = Pt(10)

    add_paragraph(doc, 'Step 3: 创建Snowflake数据库管理器', bold=True)
    add_paragraph(doc, '在 src/main/python/core/snowflake_db.py 中创建连接管理器')

    add_paragraph(doc, 'Step 4: 设计Snowflake数据模型', bold=True)
    snowflake_tables = [
        'FACT_ASSESSMENTS - 评估事实表',
        'FACT_RESPONSES - 答题事实表',
        'DIM_USERS - 用户维度表',
        'DIM_QUESTIONS - 题目维度表',
        'DIM_DIVISIONS - 部门维度表',
        'DIM_DATE - 日期维度表',
        'AGG_DAILY_STATS - 每日统计聚合表',
        'AGG_DIVISION_PERFORMANCE - 部门绩效聚合表'
    ]
    for table in snowflake_tables:
        add_bullet_point(doc, table)

    doc.add_page_break()

    # ==================== DEVELOPMENT RECOMMENDATIONS ====================
    add_heading(doc, '五、开发建议 (Development Recommendations)', 1)

    add_paragraph(doc, '5.1 数据架构优化', bold=True, font_size=12)

    recommendations = [
        ['优先级', '建议项', '描述', '预期收益'],
        ['🔴 高', '实施Snowflake集成',
         '将历史数据和分析数据迁移到Snowflake，保留PostgreSQL处理实时事务',
         '提升分析查询性能10倍以上，降低数据库负载'],
        ['🔴 高', '实施数据仓库模型',
         '采用星型模式设计事实表和维度表',
         '优化复杂查询，支持多维度分析'],
        ['🟡 中', '建立CDC数据管道',
         '使用Debezium或类似工具实现实时数据同步',
         '确保数据一致性，减少同步延迟'],
        ['🟡 中', '实施Snowpipe流式摄取',
         '自动化数据加载流程',
         '减少手动ETL工作量，实现近实时分析'],
        ['🟢 低', '数据湖集成',
         '将原始音频文件存储到S3/Azure Blob，通过Snowflake External Tables访问',
         '降低存储成本，支持非结构化数据分析']
    ]
    add_table_with_header(doc, recommendations[0], recommendations[1:])

    add_paragraph(doc, '5.2 性能优化建议', bold=True, font_size=12)
    performance = [
        '使用Snowflake聚簇键 (Clustering Keys) 优化大表查询',
        '实施自动聚簇 (Automatic Clustering) 维护数据顺序',
        '使用物化视图 (Materialized Views) 缓存复杂查询结果',
        '启用查询结果缓存 (Result Cache) 减少重复计算',
        '合理设置虚拟仓库大小，使用动态扩展',
        '实施分区策略，按时间或部门分区大表'
    ]
    for item in performance:
        add_bullet_point(doc, item)

    add_paragraph(doc, '5.3 安全与合规', bold=True, font_size=12)
    security = [
        '启用端到端加密 (End-to-End Encryption)',
        '实施基于角色的访问控制 (RBAC)',
        '使用Snowflake动态数据脱敏 (Dynamic Data Masking)',
        '配置网络策略限制访问来源',
        '启用审计日志跟踪所有数据访问',
        '实施多因素认证 (MFA)',
        '定期进行安全审计和漏洞扫描'
    ]
    for item in security:
        add_bullet_point(doc, item)

    add_paragraph(doc, '5.4 机器学习增强', bold=True, font_size=12)
    ml_enhancements = [
        '使用Snowpark Python在Snowflake内训练ML模型',
        '实施UDF (User-Defined Functions) 进行实时评分',
        '集成Snowflake ML预测功能预测候选人表现',
        '使用Snowflake Cortex实现AI驱动的数据分析',
        '建立特征工程管道提取评估数据特征',
        '实施A/B测试框架优化评估流程'
    ]
    for item in ml_enhancements:
        add_bullet_point(doc, item)

    doc.add_page_break()

    # ==================== IMPLEMENTATION ROADMAP ====================
    add_heading(doc, '六、实施路线图 (Implementation Roadmap)', 1)

    phases = [
        ['阶段', '任务', '时间', '关键交付物'],
        ['第一阶段\n准备期',
         '• 评估当前数据规模\n• 设计Snowflake数据模型\n• 搭建开发环境\n• 创建POC验证',
         '2-3周',
         '• 技术方案文档\n• POC演示\n• 成本预估'],
        ['第二阶段\n开发期',
         '• 开发Snowflake连接器\n• 实施数据迁移脚本\n• 创建事实表和维度表\n• 开发ETL管道',
         '4-6周',
         '• 数据管道代码\n• 单元测试\n• 技术文档'],
        ['第三阶段\n集成期',
         '• 集成分析API\n• 开发BI报表\n• 实施数据同步\n• 性能测试',
         '3-4周',
         '• API端点\n• 报表仪表板\n• 性能基准'],
        ['第四阶段\n优化期',
         '• 性能调优\n• 安全加固\n• 文档完善\n• 团队培训',
         '2-3周',
         '• 优化报告\n• 用户手册\n• 培训材料'],
        ['第五阶段\n上线期',
         '• 生产环境部署\n• 数据迁移\n• 监控配置\n• 上线支持',
         '1-2周',
         '• 生产系统\n• 监控仪表板\n• 运维手册']
    ]
    add_table_with_header(doc, phases[0], phases[1:])

    add_paragraph(doc, '预计总时间：12-18周 (约3-4.5个月)', bold=True)

    doc.add_page_break()

    # ==================== COST ANALYSIS ====================
    add_heading(doc, '七、成本分析 (Cost Analysis)', 1)

    add_paragraph(doc, '7.1 Snowflake成本构成', bold=True, font_size=12)
    cost_components = [
        '计算成本 (Virtual Warehouse)',
        '  - 按秒计费，最小1分钟',
        '  - X-Small: $2/小时，Small: $4/小时，Medium: $8/小时',
        '  - 建议：开发环境使用X-Small，生产环境使用Small-Medium',
        '',
        '存储成本',
        '  - $23-40/TB/月 (按区域不同)',
        '  - 自动压缩，通常可节省50-70%存储空间',
        '',
        '数据传输成本',
        '  - 区域内免费',
        '  - 跨区域传输收费'
    ]
    for item in cost_components:
        if item.startswith('  '):
            add_bullet_point(doc, item.strip(), level=1)
        elif item:
            add_paragraph(doc, item, bold=True)

    add_paragraph(doc, '7.2 预估月度成本 (基于中等规模)', bold=True, font_size=12)
    cost_estimate = [
        ['项目', '规模', '成本估算 (USD/月)'],
        ['计算资源', 'Small warehouse, 8小时/天运行', '$960 (30天 × 8小时 × $4)'],
        ['存储', '5TB数据 (压缩后)', '$115-200'],
        ['数据传输', '区域内', '$0'],
        ['Snowpipe', '持续摄取', '$50-100'],
        ['总计', '', '约 $1,125-1,260/月']
    ]
    add_table_with_header(doc, cost_estimate[0], cost_estimate[1:])

    add_paragraph(doc, '7.3 成本优化策略', bold=True, font_size=12)
    cost_optimization = [
        '使用自动挂起 (Auto-Suspend) 功能，空闲时自动关闭仓库',
        '合理设置仓库大小，避免过度配置',
        '使用查询标签 (Query Tags) 跟踪成本归因',
        '启用资源监控器 (Resource Monitors) 控制预算',
        '定期清理过期数据和临时表',
        '使用Snowflake Cost Attribution跟踪成本'
    ]
    for item in cost_optimization:
        add_bullet_point(doc, item)

    doc.add_page_break()

    # ==================== RISKS AND MITIGATION ====================
    add_heading(doc, '八、风险与缓解措施 (Risks & Mitigation)', 1)

    risks = [
        ['风险类别', '具体风险', '影响', '缓解措施'],
        ['技术风险', '数据迁移失败或数据丢失', '高',
         '• 分阶段迁移\n• 完整备份\n• 并行运行验证'],
        ['技术风险', '性能未达预期', '中',
         '• POC阶段性能测试\n• 使用Snowflake优化建议\n• 专家咨询'],
        ['运营风险', '团队缺乏Snowflake经验', '中',
         '• Snowflake官方培训\n• 聘请顾问\n• 逐步迁移学习'],
        ['财务风险', '成本超出预算', '中',
         '• 设置资源监控器\n• 定期成本审查\n• 优化查询效率'],
        ['安全风险', '数据泄露或未授权访问', '高',
         '• 实施RBAC\n• 启用加密\n• 定期安全审计'],
        ['业务风险', '影响现有业务运行', '高',
         '• 灰度发布\n• 回滚计划\n• 24/7监控']
    ]
    add_table_with_header(doc, risks[0], risks[1:])

    doc.add_page_break()

    # ==================== SUCCESS METRICS ====================
    add_heading(doc, '九、成功指标 (Success Metrics)', 1)

    add_paragraph(doc, '9.1 技术指标', bold=True, font_size=12)
    tech_metrics = [
        '查询性能：复杂分析查询响应时间 < 5秒',
        '数据同步延迟：< 5分钟',
        '系统可用性：> 99.9%',
        '数据一致性：100% (通过自动化验证)',
        'API响应时间：P95 < 200ms',
        '并发支持：> 100个并发查询'
    ]
    for metric in tech_metrics:
        add_bullet_point(doc, metric)

    add_paragraph(doc, '9.2 业务指标', bold=True, font_size=12)
    business_metrics = [
        '报表生成时间减少 70%',
        '数据分析人员效率提升 50%',
        '数据存储成本降低 30%',
        '支持的同时评估人数增加 3倍',
        '新数据分析需求交付时间缩短 60%',
        'BI仪表板用户满意度 > 4.5/5'
    ]
    for metric in business_metrics:
        add_bullet_point(doc, metric)

    add_paragraph(doc, '9.3 用户体验指标', bold=True, font_size=12)
    ux_metrics = [
        '管理员报表查询满意度 > 90%',
        '数据分析师培训完成率 > 95%',
        '报表错误率 < 1%',
        '自助分析采用率 > 70%'
    ]
    for metric in ux_metrics:
        add_bullet_point(doc, metric)

    doc.add_page_break()

    # ==================== NEXT STEPS ====================
    add_heading(doc, '十、后续行动计划 (Next Steps)', 1)

    add_paragraph(doc, '10.1 即时行动 (未来1个月)', bold=True, font_size=12)
    immediate_actions = [
        '✅ 组建Snowflake集成项目团队',
        '✅ 申请Snowflake试用账户',
        '✅ 完成POC环境搭建',
        '✅ 评估当前数据规模和增长趋势',
        '✅ 设计数据仓库模型初版',
        '✅ 制定详细项目计划和预算'
    ]
    for action in immediate_actions:
        add_bullet_point(doc, action)

    add_paragraph(doc, '10.2 短期行动 (1-3个月)', bold=True, font_size=12)
    short_term_actions = [
        '🔄 完成Snowflake环境配置',
        '🔄 开发数据同步管道',
        '🔄 迁移历史数据到Snowflake',
        '🔄 开发分析API和BI报表',
        '🔄 进行性能和压力测试',
        '🔄 团队培训和知识转移'
    ]
    for action in short_term_actions:
        add_bullet_point(doc, action)

    add_paragraph(doc, '10.3 中期行动 (3-6个月)', bold=True, font_size=12)
    medium_term_actions = [
        '📈 生产环境上线',
        '📈 实施高级分析功能',
        '📈 集成机器学习模型',
        '📈 优化成本和性能',
        '📈 扩展数据源集成',
        '📈 建立数据治理框架'
    ]
    for action in medium_term_actions:
        add_bullet_point(doc, action)

    doc.add_page_break()

    # ==================== CONCLUSION ====================
    add_heading(doc, '十一、结论 (Conclusion)', 1)

    add_paragraph(doc,
        '邮轮员工英语评估平台已经建立了坚实的技术基础，通过FastAPI、PostgreSQL和AI服务的集成，'
        '成功实现了智能化的英语能力评估。当前系统架构清晰，代码质量良好，功能模块完整。')

    doc.add_paragraph()

    add_paragraph(doc,
        '集成Snowflake数据平台将为系统带来显著提升：', bold=True)

    conclusion_benefits = [
        '💡 分析能力：支持复杂的多维度数据分析和BI报表',
        '⚡ 性能提升：大幅提升查询性能，特别是聚合分析场景',
        '📊 数据洞察：通过高级分析发现评估趋势和优化机会',
        '🔐 企业级能力：增强数据治理、安全和合规能力',
        '💰 成本优化：通过存储计算分离和按需付费降低总体成本',
        '🚀 未来就绪：为AI/ML模型训练和大规模扩展做好准备'
    ]
    for benefit in conclusion_benefits:
        add_bullet_point(doc, benefit)

    doc.add_paragraph()

    add_paragraph(doc,
        '建议采用分阶段实施策略，从POC验证开始，逐步迁移数据和功能，确保平稳过渡。'
        '通过合理的成本控制和性能优化，Snowflake集成将为邮轮员工英语评估平台带来长期价值，'
        '支持业务持续增长和数据驱动决策。')

    doc.add_paragraph()
    doc.add_paragraph()

    final_para = doc.add_paragraph()
    final_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    final_para.add_run('--- 报告结束 ---\n').bold = True
    final_para.add_run(f'生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')

    # ==================== APPENDIX ====================
    doc.add_page_break()
    add_heading(doc, '附录A：技术参考资源 (Appendix A: Technical References)', 1)

    references = [
        'Snowflake官方文档: https://docs.snowflake.com/',
        'Snowflake Python Connector: https://docs.snowflake.com/en/developer-guide/python-connector/',
        'Snowflake SQLAlchemy: https://docs.snowflake.com/en/developer-guide/python-connector/sqlalchemy',
        'FastAPI文档: https://fastapi.tiangolo.com/',
        'SQLAlchemy异步文档: https://docs.sqlalchemy.org/en/20/orm/extensions/asyncio.html',
        'Snowpark Python: https://docs.snowflake.com/en/developer-guide/snowpark/python/index',
        'Snowflake成本优化指南: https://www.snowflake.com/en/data-cloud/workloads/cost-optimization/',
        'Snowflake安全最佳实践: https://docs.snowflake.com/en/user-guide/security-best-practices'
    ]
    for ref in references:
        add_bullet_point(doc, ref)

    doc.add_page_break()
    add_heading(doc, '附录B：项目文件清单 (Appendix B: Project Files)', 1)

    key_files = [
        'src/main/python/core/database.py - 数据库连接管理',
        'src/main/python/models/assessment.py - 数据模型定义',
        'src/main/python/api/routes/assessment.py - 评估API路由',
        'src/main/python/api/routes/analytics.py - 分析API路由',
        'src/main/python/services/ai_service.py - AI服务集成',
        'src/main/python/utils/scoring.py - 评分算法',
        'src/main/python/data/generate_question_bank.py - 题库生成',
        'requirements.txt - Python依赖清单',
        'README.md - 项目说明文档',
        'CLAUDE.md - 开发指南'
    ]
    for file in key_files:
        add_bullet_point(doc, file)

    # Save document
    output_path = "C:\\Users\\szh2051\\OneDrive - Carnival Corporation\\Desktop\\Python\\Claude Demo\\output\\CCL_English_Assessment_Snowflake_Integration_Report.docx"
    doc.save(output_path)
    print(f"Project summary document generated successfully: {output_path}")
    return output_path


if __name__ == "__main__":
    generate_project_summary()
