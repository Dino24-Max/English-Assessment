#!/usr/bin/env python3
"""
创建邮轮员工英语评估平台 - 运营部门和场景文档
Cruise Employee English Assessment Platform - Operations and Scenarios Documentation
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# 创建文档
doc = Document()

# 设置文档标题
title = doc.add_heading('邮轮员工英语评估平台', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.runs[0]
title_run.font.color.rgb = RGBColor(0, 51, 102)

subtitle = doc.add_heading('运营部门结构与评估场景说明', level=2)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 添加日期
date_para = doc.add_paragraph(f'文档生成日期: {datetime.now().strftime("%Y年%m月%d日")}')
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()  # 空行

# ==================== 概述部分 ====================
doc.add_heading('一、平台概述', level=1)

overview_text = """
本英语评估平台专为邮轮员工设计，涵盖三大运营部门的14个具体部门。评估内容包括6个核心模块，
共21道题目，总分100分。评估场景均基于真实邮轮运营环境，确保测试内容与实际工作密切相关。
"""
doc.add_paragraph(overview_text)

# 评估模块概览
doc.add_heading('评估模块结构', level=2)
modules_data = [
    ['模块名称', '题目数量', '分值', '测试内容'],
    ['听力理解 (Listening)', '3题', '12分', '理解客人请求、预订、投诉等对话'],
    ['时间与数字 (Time & Numbers)', '3题', '12分', '准确理解时间、数量、房间号等信息'],
    ['语法 (Grammar)', '4题', '16分', '邮轮服务场景的语法应用'],
    ['词汇 (Vocabulary)', '4题', '32分', '邮轮专业术语匹配'],
    ['阅读理解 (Reading)', '4题', '24分', '理解邮轮公告、政策、通知'],
    ['口语表达 (Speaking)', '3题', '20分', '真实场景客户服务对话']
]

table = doc.add_table(rows=len(modules_data), cols=len(modules_data[0]))
table.style = 'Light Grid Accent 1'

for i, row_data in enumerate(modules_data):
    row = table.rows[i]
    for j, cell_data in enumerate(row_data):
        cell = row.cells[j]
        cell.text = cell_data
        if i == 0:  # 表头加粗
            cell.paragraphs[0].runs[0].font.bold = True

doc.add_paragraph()

# ==================== 三大运营部门 ====================
doc.add_heading('二、三大运营部门详细说明', level=1)

# ========== 1. HOTEL OPERATIONS ==========
doc.add_heading('1. 酒店运营部 (HOTEL OPERATIONS)', level=2)
hotel_intro = doc.add_paragraph(
    '酒店运营部负责邮轮上的所有客户服务、餐饮、客房管理等事务，是邮轮运营的核心部门。'
)

doc.add_heading('部门细分 (8个部门):', level=3)

hotel_depts = {
    'Guest Services (客户服务部)': {
        '职责': '处理客人入住登记、咨询、投诉、礼宾服务',
        '评估场景': [
            '餐厅预订咨询 - "I\'d like to book a table for four people at seven PM"',
            '设施位置指引 - "How do I get to the spa?"',
            '客房问题处理 - "The air conditioning is too cold in my room"',
            '紧急情况应对 - 理解并回应客人需求'
        ]
    },
    'Housekeeping (客房部)': {
        '职责': '客舱清洁、床品更换、客房维护',
        '评估场景': [
            '房间号码识别 - "Room 8254"',
            '客房设施报修 - "My safe isn\'t working properly"',
            '清洁服务时间安排',
            '特殊要求处理'
        ]
    },
    'Food & Beverage (餐饮部)': {
        '职责': '餐厅服务、酒吧服务、宴会服务',
        '评估场景': [
            '餐厅预订处理 - "Table for eight people at six-thirty"',
            '用餐时间说明 - "Breakfast is served from 7 AM to 10:30 AM"',
            '特殊饮食需求 - "Dietary restrictions with 24-hour advance notice"',
            '酒单推荐 - Sommelier (侍酒师) 专业术语'
        ]
    },
    'Culinary (厨房部)': {
        '职责': '食品制作、菜单设计、食品安全',
        '评估场景': [
            '厨房术语 - "Galley" (船上厨房)',
            '餐饮类型 - "Buffet" (自助餐), "A la carte" (单点菜单)',
            '食品准备时间沟通',
            '过敏原信息传达'
        ]
    },
    'Bar & Lounge (酒吧休息厅)': {
        '职责': '酒水服务、调酒、休息区管理',
        '评估场景': [
            '营业时间告知',
            '酒水推荐',
            '年龄验证沟通 - "Guests must be 21+"',
            '服务礼仪用语'
        ]
    },
    'Entertainment (娱乐部)': {
        '职责': '活动策划、表演安排、游客娱乐',
        '评估场景': [
            '活动时间通知',
            '场地位置指引 - "Deck 12" or "Deck 14"',
            '节目预告说明',
            '参与须知解释'
        ]
    },
    'Spa & Fitness (水疗健身部)': {
        '职责': '水疗服务、健身指导、美容服务',
        '评估场景': [
            '设施位置指引 - "Follow signs to the spa"',
            '预约服务说明 - "Reservations required"',
            '服务项目介绍',
            '营业时间告知'
        ]
    },
    'Shore Excursions (岸上观光部)': {
        '职责': '岸上行程安排、导游服务、票务管理',
        '评估场景': [
            '上岸时间通知 - "Embark" (登船) 术语',
            '行程安排说明 - "Excursion" (岸上游览)',
            '码头代理联系 - "Contact the Port Agent"',
            '误船应对 - "Guests who miss departure must contact..."'
        ]
    }
}

for dept_name, details in hotel_depts.items():
    doc.add_heading(dept_name, level=4)
    doc.add_paragraph(f"✓ 职责范围: {details['职责']}")
    doc.add_paragraph('✓ 评估相关场景:')
    for scenario in details['评估场景']:
        doc.add_paragraph(scenario, style='List Bullet 2')
    doc.add_paragraph()

# ========== 2. MARINE OPERATIONS ==========
doc.add_page_break()
doc.add_heading('2. 海事运营部 (MARINE OPERATIONS)', level=2)
marine_intro = doc.add_paragraph(
    '海事运营部负责邮轮的航行、安全、技术维护，确保船只安全运营和法规遵守。'
)

doc.add_heading('部门细分 (3个部门):', level=3)

marine_depts = {
    'Deck Department (甲板部)': {
        '职责': '船舶航行、导航、甲板维护、安全巡逻',
        '评估场景': [
            '船舶术语理解 - "Bridge" (驾驶台), "Gangway" (舷梯)',
            '甲板位置说明 - "Deck 12", "Deck 14"',
            '安全演习指导 - "Muster drill" (安全演习)',
            '紧急集合 - "Muster" (紧急集合), "Assembly station" (集合点)'
        ]
    },
    'Engine Department (轮机部)': {
        '职责': '发动机维护、电力系统、技术支持',
        '评估场景': [
            '设备维修通知 - "Air conditioning issue", "Safe malfunction"',
            '技术术语理解 - "Maintenance" (维护)',
            '维修时间预估',
            '系统状态沟通'
        ]
    },
    'Safety & Security (安全保卫部)': {
        '职责': '安全培训、应急响应、监控、安保',
        '评估场景': [
            '安全演习通知 - "Mandatory muster drill at 4:00 PM"',
            '安全设备说明 - "Life jacket" (救生衣)',
            '紧急程序 - "Assembly station" (集合点)',
            '登船呼叫 - "All aboard" (最后登船通知)'
        ]
    }
}

for dept_name, details in marine_depts.items():
    doc.add_heading(dept_name, level=4)
    doc.add_paragraph(f"✓ 职责范围: {details['职责']}")
    doc.add_paragraph('✓ 评估相关场景:')
    for scenario in details['评估场景']:
        doc.add_paragraph(scenario, style='List Bullet 2')
    doc.add_paragraph()

# ========== 3. CASINO OPERATIONS ==========
doc.add_page_break()
doc.add_heading('3. 赌场运营部 (CASINO OPERATIONS)', level=2)
casino_intro = doc.add_paragraph(
    '赌场运营部负责船上赌场的运营管理、游戏服务、合规监管，需严格遵守国际海事法规。'
)

doc.add_heading('部门细分 (3个部门):', level=3)

casino_depts = {
    'Table Games (桌面游戏部)': {
        '职责': '扑克、21点、轮盘等桌面游戏服务',
        '评估场景': [
            '营业时间说明 - "8:00 AM - 2:00 AM"',
            '开放时段区分 - "Sea Days" (海上航行日) vs "Port Days" (停靠港口日)',
            '年龄限制沟通 - "Guests must be 21+ to gamble"',
            '证件查验 - "Valid ID required"'
        ]
    },
    'Slot Machines (老虎机部)': {
        '职责': '老虎机运营、机器维护、奖金支付',
        '评估场景': [
            '游戏规则解释',
            '机器操作指导',
            '支付方式说明',
            '技术问题报告'
        ]
    },
    'Casino Administration (赌场行政部)': {
        '职责': '赌场监管、合规检查、财务管理',
        '评估场景': [
            '法规限制说明 - "Subject to maritime laws"',
            '领海限制 - "Suspended while in territorial waters"',
            '营业时间政策 - 根据海事法律调整',
            '合规要求传达'
        ]
    }
}

for dept_name, details in casino_depts.items():
    doc.add_heading(dept_name, level=4)
    doc.add_paragraph(f"✓ 职责范围: {details['职责']}")
    doc.add_paragraph('✓ 评估相关场景:')
    for scenario in details['评估场景']:
        doc.add_paragraph(scenario, style='List Bullet 2')
    doc.add_paragraph()

# ==================== 评估场景详细分类 ====================
doc.add_page_break()
doc.add_heading('三、评估场景详细分类', level=1)

# 场景分类
doc.add_heading('1. 客户服务场景 (Customer Service Scenarios)', level=2)
service_scenarios = [
    '预订与咨询 - 餐厅预订、活动报名、水疗预约',
    '投诉处理 - 房间空调问题、设施故障、服务不满',
    '信息查询 - 营业时间、位置指引、活动安排',
    '特殊需求 - 饮食限制、医疗协助、无障碍服务'
]
for scenario in service_scenarios:
    doc.add_paragraph(scenario, style='List Bullet')

doc.add_heading('2. 时间与数字场景 (Time & Numbers Scenarios)', level=2)
time_scenarios = [
    '营业时间 - 早餐7:00 AM - 10:30 AM, 赌场营业至2:00 AM',
    '预订信息 - 8位客人6:30PM用餐',
    '房间编号 - 8254号房, 9173号客舱',
    '甲板楼层 - Deck 12, Deck 14'
]
for scenario in time_scenarios:
    doc.add_paragraph(scenario, style='List Bullet')

doc.add_heading('3. 安全与应急场景 (Safety & Emergency Scenarios)', level=2)
safety_scenarios = [
    '安全演习 - 强制性集合演习,下午4:00 PM',
    '紧急设备 - 救生衣、集合点、逃生路线',
    '误船处理 - 联系港口代理、自行前往下一港口',
    '安全公告 - 根据海事法律的限制说明'
]
for scenario in safety_scenarios:
    doc.add_paragraph(scenario, style='List Bullet')

doc.add_heading('4. 专业术语场景 (Technical Terminology Scenarios)', level=2)
term_scenarios = [
    '船舶术语 - Bridge(驾驶台), Gangway(舷梯), Tender(接驳船)',
    '餐饮术语 - Galley(厨房), Buffet(自助餐), Sommelier(侍酒师)',
    '服务术语 - Concierge(礼宾), Amenities(设施), Excursion(岸上游)',
    '安全术语 - Muster(集合), Assembly station(集合点), Embark(登船)'
]
for scenario in term_scenarios:
    doc.add_paragraph(scenario, style='List Bullet')

# ==================== 题目场景映射表 ====================
doc.add_page_break()
doc.add_heading('四、评估题目与场景映射', level=1)

question_mapping = [
    ['模块', '题号', '场景', '涉及部门'],
    ['听力', 'Q1', '餐厅预订 (晚7点4人用餐)', 'Food & Beverage'],
    ['听力', 'Q2', '客房空调报修 (8254号房)', 'Housekeeping, Engineering'],
    ['听力', 'Q3', '自助餐厅位置询问 (12层)', 'Guest Services, Entertainment'],
    ['时间数字', 'Q4', '早餐供应时间 (7:00-10:30)', 'Food & Beverage'],
    ['时间数字', 'Q5', '预订人数 (8人, 6:30PM)', 'Food & Beverage'],
    ['时间数字', 'Q6', '客舱号码 (9173号保险柜故障)', 'Housekeeping, Engineering'],
    ['语法', 'Q7', '行李服务礼貌用语', 'Guest Services'],
    ['语法', 'Q8', '客人抵达时态', 'Guest Services'],
    ['语法', 'Q9', '水疗位置指引', 'Spa & Fitness'],
    ['语法', 'Q10', '餐厅位置介词', 'Food & Beverage'],
    ['词汇', 'Q11', '船舶术语匹配', 'Deck Department'],
    ['词汇', 'Q12', '酒店服务术语', 'Guest Services'],
    ['词汇', 'Q13', '餐饮术语', 'Food & Beverage, Culinary'],
    ['词汇', 'Q14', '安全术语', 'Safety & Security'],
    ['阅读', 'Q15', '误船政策说明', 'Shore Excursions'],
    ['阅读', 'Q16', '赌场营业时间', 'Casino Operations'],
    ['阅读', 'Q17', '特色餐厅预订政策', 'Food & Beverage'],
    ['阅读', 'Q18', '安全演习时间', 'Safety & Security'],
    ['口语', 'Q19', '空调问题应对', 'Housekeeping, Engineering'],
    ['口语', 'Q20', '自助餐时间查询', 'Food & Beverage'],
    ['口语', 'Q21', '水疗位置指引', 'Spa & Fitness, Guest Services']
]

mapping_table = doc.add_table(rows=len(question_mapping), cols=len(question_mapping[0]))
mapping_table.style = 'Light List Accent 1'

for i, row_data in enumerate(question_mapping):
    row = mapping_table.rows[i]
    for j, cell_data in enumerate(row_data):
        cell = row.cells[j]
        cell.text = cell_data
        if i == 0:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(11)

# ==================== 评分标准 ====================
doc.add_page_break()
doc.add_heading('五、评分标准', level=1)

doc.add_heading('总分与通过标准', level=2)
scoring_overview = """
• 总分: 100分
• 及格分数: 65分 (65%)
• 安全相关题目通过率: 80%
• 口语模块最低分: 12分 (60%)
"""
doc.add_paragraph(scoring_overview)

doc.add_heading('各模块评分细则', level=2)
module_scores = [
    ['模块', '题目数', '每题分值', '总分', '备注'],
    ['听力理解', '3', '4分', '12分', '选择题, 正确选项得分'],
    ['时间与数字', '3', '4分', '12分', '填空题, 完全正确得分'],
    ['语法', '4', '4分', '16分', '选择题, 正确选项得分'],
    ['词汇', '4', '8分', '32分', '匹配题, 每对2分'],
    ['阅读理解', '4', '6分', '24分', '选择题, 正确选项得分'],
    ['口语表达', '3', '20分/3≈6.67', '20分', 'AI评分, 综合评估发音/流利度/内容']
]

score_table = doc.add_table(rows=len(module_scores), cols=len(module_scores[0]))
score_table.style = 'Light Grid Accent 1'

for i, row_data in enumerate(module_scores):
    row = score_table.rows[i]
    for j, cell_data in enumerate(row_data):
        cell = row.cells[j]
        cell.text = str(cell_data)
        if i == 0:
            cell.paragraphs[0].runs[0].font.bold = True

# ==================== 技术实现 ====================
doc.add_page_break()
doc.add_heading('六、技术实现', level=1)

doc.add_heading('前端技术', level=2)
frontend_tech = """
• HTML5 + CSS3: 响应式设计, 移动端友好
• JavaScript ES6+: 交互逻辑, 表单验证, 音频控制
• Web Speech API: 文本转语音 (TTS) 用于听力模块
• MediaRecorder API: 语音录制用于口语模块
• 拖拽 API: 词汇匹配模块的交互实现
"""
doc.add_paragraph(frontend_tech)

doc.add_heading('后端技术', level=2)
backend_tech = """
• FastAPI: Python异步Web框架
• SQLAlchemy: ORM数据库操作
• PostgreSQL: 关系型数据库
• Redis: 会话管理与缓存
• AI服务: OpenAI/Anthropic Claude用于口语评分
"""
doc.add_paragraph(backend_tech)

doc.add_heading('安全特性', level=2)
security_features = """
• CSRF保护: 防止跨站请求伪造
• 速率限制: 防止滥用和DOS攻击
• 输入验证: XSS和SQL注入防护
• 会话加密: Fernet对称加密
• 安全头部: HSTS, CSP, X-Frame-Options等
"""
doc.add_paragraph(security_features)

# ==================== 附录 ====================
doc.add_page_break()
doc.add_heading('七、附录: 关键英语术语表', level=1)

doc.add_heading('A. 船舶术语 (Vessel Terminology)', level=2)
vessel_terms = [
    'Bridge - 驾驶台',
    'Gangway - 舷梯/登船通道',
    'Tender - 接驳船',
    'Muster - 紧急集合',
    'Deck - 甲板/楼层',
    'Starboard - 右舷',
    'Port - 左舷',
    'Bow - 船头',
    'Stern - 船尾'
]
for term in vessel_terms:
    doc.add_paragraph(term, style='List Bullet')

doc.add_heading('B. 酒店服务术语 (Hospitality Terminology)', level=2)
hospitality_terms = [
    'Concierge - 礼宾服务',
    'Amenities - 设施/便利设施',
    'Housekeeping - 客房服务',
    'Turndown service - 开夜床服务',
    'Stateroom - 客舱',
    'Suite - 套房',
    'Balcony cabin - 阳台房'
]
for term in hospitality_terms:
    doc.add_paragraph(term, style='List Bullet')

doc.add_heading('C. 餐饮术语 (Food & Beverage Terminology)', level=2)
fb_terms = [
    'Galley - 船上厨房',
    'Buffet - 自助餐',
    'A la carte - 单点菜单',
    'Sommelier - 侍酒师',
    'Main dining room - 主餐厅',
    'Specialty restaurant - 特色餐厅',
    'Cover charge - 附加费'
]
for term in fb_terms:
    doc.add_paragraph(term, style='List Bullet')

doc.add_heading('D. 安全术语 (Safety Terminology)', level=2)
safety_terms = [
    'Muster drill - 安全演习',
    'Life jacket/vest - 救生衣',
    'Assembly station - 集合点',
    'All aboard - 最后登船通知',
    'Lifeboat - 救生艇',
    'Emergency exit - 紧急出口',
    'Evacuation - 疏散'
]
for term in safety_terms:
    doc.add_paragraph(term, style='List Bullet')

doc.add_heading('E. 旅游术语 (Travel Terminology)', level=2)
travel_terms = [
    'Embark - 登船',
    'Disembark - 下船',
    'Excursion - 岸上游览',
    'Port of call - 停靠港',
    'Itinerary - 行程',
    'Shore leave - 上岸时间',
    'Port Agent - 港口代理'
]
for term in travel_terms:
    doc.add_paragraph(term, style='List Bullet')

# ==================== 结尾 ====================
doc.add_page_break()
doc.add_heading('文档说明', level=1)

notes = """
本文档详细列出了邮轮员工英语评估平台涉及的三大运营部门 (Hotel, Marine, Casino)、
14个具体部门、以及所有21道评估题目对应的真实工作场景。

评估内容完全基于邮轮行业实际运营需求,确保员工通过评估后能够胜任相应岗位的英语沟通要求。

如需了解更多技术实现细节或评估流程,请参考项目GitHub仓库:
https://github.com/Dino24-Max/English-Assessment

生成时间: {}
平台版本: v1.0
""".format(datetime.now().strftime("%Y年%m月%d日 %H:%M"))

doc.add_paragraph(notes)

# 保存文档
output_path = r"C:\Users\szh2051\OneDrive - Carnival Corporation\Desktop\邮轮员工英语评估平台-运营部门与场景说明.docx"
doc.save(output_path)

print("Word document created successfully!")
print(f"File path: {output_path}")
print("Content: 3 operations, 14 departments, 21 question scenarios")
print("Total pages: about 15-20 pages")
