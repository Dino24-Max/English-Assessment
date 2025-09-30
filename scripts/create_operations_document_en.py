#!/usr/bin/env python3
"""
Cruise Employee English Assessment Platform - Operations and Scenarios Documentation (English Version)
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# Create document
doc = Document()

# Set document title
title = doc.add_heading('Cruise Employee English Assessment Platform', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.runs[0]
title_run.font.color.rgb = RGBColor(0, 51, 102)

subtitle = doc.add_heading('Operations Structure & Assessment Scenarios', level=2)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add date
date_para = doc.add_paragraph(f'Document Generated: {datetime.now().strftime("%B %d, %Y")}')
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()  # Blank line

# ==================== OVERVIEW ====================
doc.add_heading('I. Platform Overview', level=1)

overview_text = """
This English Assessment Platform is specifically designed for cruise ship employees, covering three major
operational divisions with 14 specific departments. The assessment includes 6 core modules with a total of
21 questions worth 100 points. All assessment scenarios are based on real cruise operations, ensuring test
content is closely aligned with actual work requirements.
"""
doc.add_paragraph(overview_text)

# Assessment Module Overview
doc.add_heading('Assessment Module Structure', level=2)
modules_data = [
    ['Module Name', 'Questions', 'Points', 'Test Content'],
    ['Listening', '3', '12', 'Understanding guest requests, reservations, complaints'],
    ['Time & Numbers', '3', '12', 'Accurate comprehension of times, quantities, room numbers'],
    ['Grammar', '4', '16', 'Grammar application in cruise service scenarios'],
    ['Vocabulary', '4', '32', 'Cruise industry terminology matching'],
    ['Reading', '4', '24', 'Understanding cruise announcements, policies, notices'],
    ['Speaking', '3', '20', 'Real scenario customer service dialogues']
]

table = doc.add_table(rows=len(modules_data), cols=len(modules_data[0]))
table.style = 'Light Grid Accent 1'

for i, row_data in enumerate(modules_data):
    row = table.rows[i]
    for j, cell_data in enumerate(row_data):
        cell = row.cells[j]
        cell.text = cell_data
        if i == 0:
            cell.paragraphs[0].runs[0].font.bold = True

doc.add_paragraph()

# ==================== THREE MAJOR OPERATIONS ====================
doc.add_heading('II. Three Major Operational Divisions', level=1)

# ========== 1. HOTEL OPERATIONS ==========
doc.add_heading('1. HOTEL OPERATIONS', level=2)
hotel_intro = doc.add_paragraph(
    'Hotel Operations is responsible for all customer service, dining, housekeeping, and guest-related '
    'activities onboard, serving as the core department of cruise operations.'
)

doc.add_heading('Department Breakdown (8 Departments):', level=3)

hotel_depts = {
    'Guest Services': {
        'Responsibilities': 'Handle check-in, inquiries, complaints, concierge services',
        'Assessment Scenarios': [
            'Restaurant reservations - "I\'d like to book a table for four people at seven PM"',
            'Facility directions - "How do I get to the spa?"',
            'Room issues - "The air conditioning is too cold in my room"',
            'Emergency response - Understanding and responding to guest needs'
        ]
    },
    'Housekeeping': {
        'Responsibilities': 'Cabin cleaning, linen changes, room maintenance',
        'Assessment Scenarios': [
            'Room number identification - "Room 8254"',
            'Room facility repairs - "My safe isn\'t working properly"',
            'Cleaning service scheduling',
            'Special request handling'
        ]
    },
    'Food & Beverage': {
        'Responsibilities': 'Restaurant service, bar service, banquet service',
        'Assessment Scenarios': [
            'Reservation processing - "Table for eight people at six-thirty"',
            'Dining time information - "Breakfast is served from 7 AM to 10:30 AM"',
            'Special dietary needs - "Dietary restrictions with 24-hour advance notice"',
            'Wine recommendations - Sommelier professional terminology'
        ]
    },
    'Culinary': {
        'Responsibilities': 'Food preparation, menu design, food safety',
        'Assessment Scenarios': [
            'Kitchen terminology - "Galley" (ship\'s kitchen)',
            'Dining types - "Buffet", "A la carte"',
            'Food preparation time communication',
            'Allergen information delivery'
        ]
    },
    'Bar & Lounge': {
        'Responsibilities': 'Beverage service, mixology, lounge management',
        'Assessment Scenarios': [
            'Operating hours notification',
            'Beverage recommendations',
            'Age verification communication - "Guests must be 21+"',
            'Service etiquette phrases'
        ]
    },
    'Entertainment': {
        'Responsibilities': 'Activity planning, show arrangements, guest entertainment',
        'Assessment Scenarios': [
            'Activity time notifications',
            'Venue location directions - "Deck 12" or "Deck 14"',
            'Program announcements',
            'Participation requirements explanation'
        ]
    },
    'Spa & Fitness': {
        'Responsibilities': 'Spa services, fitness instruction, beauty services',
        'Assessment Scenarios': [
            'Facility location directions - "Follow signs to the spa"',
            'Appointment service explanations - "Reservations required"',
            'Service offerings introduction',
            'Operating hours notification'
        ]
    },
    'Shore Excursions': {
        'Responsibilities': 'Shore tour arrangements, guide services, ticketing',
        'Assessment Scenarios': [
            'Disembarkation time notifications - "Embark" terminology',
            'Itinerary explanations - "Excursion"',
            'Port agent contact - "Contact the Port Agent"',
            'Missed ship procedures - "Guests who miss departure must contact..."'
        ]
    }
}

for dept_name, details in hotel_depts.items():
    doc.add_heading(dept_name, level=4)
    doc.add_paragraph(f"Responsibilities: {details['Responsibilities']}")
    doc.add_paragraph('Assessment Scenarios:')
    for scenario in details['Assessment Scenarios']:
        doc.add_paragraph(scenario, style='List Bullet 2')
    doc.add_paragraph()

# ========== 2. MARINE OPERATIONS ==========
doc.add_page_break()
doc.add_heading('2. MARINE OPERATIONS', level=2)
marine_intro = doc.add_paragraph(
    'Marine Operations is responsible for ship navigation, safety, technical maintenance, ensuring safe '
    'vessel operations and regulatory compliance.'
)

doc.add_heading('Department Breakdown (3 Departments):', level=3)

marine_depts = {
    'Deck Department': {
        'Responsibilities': 'Ship navigation, deck maintenance, safety patrols',
        'Assessment Scenarios': [
            'Vessel terminology - "Bridge" (wheelhouse), "Gangway" (boarding ramp)',
            'Deck location descriptions - "Deck 12", "Deck 14"',
            'Safety drill instructions - "Muster drill"',
            'Emergency assembly - "Muster", "Assembly station"'
        ]
    },
    'Engine Department': {
        'Responsibilities': 'Engine maintenance, electrical systems, technical support',
        'Assessment Scenarios': [
            'Equipment repair notifications - "Air conditioning issue", "Safe malfunction"',
            'Technical terminology - "Maintenance"',
            'Repair time estimation',
            'System status communication'
        ]
    },
    'Safety & Security': {
        'Responsibilities': 'Safety training, emergency response, surveillance, security',
        'Assessment Scenarios': [
            'Safety drill announcements - "Mandatory muster drill at 4:00 PM"',
            'Safety equipment explanations - "Life jacket"',
            'Emergency procedures - "Assembly station"',
            'Boarding calls - "All aboard" (final boarding notice)'
        ]
    }
}

for dept_name, details in marine_depts.items():
    doc.add_heading(dept_name, level=4)
    doc.add_paragraph(f"Responsibilities: {details['Responsibilities']}")
    doc.add_paragraph('Assessment Scenarios:')
    for scenario in details['Assessment Scenarios']:
        doc.add_paragraph(scenario, style='List Bullet 2')
    doc.add_paragraph()

# ========== 3. CASINO OPERATIONS ==========
doc.add_page_break()
doc.add_heading('3. CASINO OPERATIONS', level=2)
casino_intro = doc.add_paragraph(
    'Casino Operations is responsible for casino management, gaming services, compliance oversight, '
    'with strict adherence to international maritime regulations.'
)

doc.add_heading('Department Breakdown (3 Departments):', level=3)

casino_depts = {
    'Table Games': {
        'Responsibilities': 'Poker, blackjack, roulette, and other table game services',
        'Assessment Scenarios': [
            'Operating hours - "8:00 AM - 2:00 AM"',
            'Operating schedule differences - "Sea Days" vs "Port Days"',
            'Age restriction communication - "Guests must be 21+ to gamble"',
            'ID verification - "Valid ID required"'
        ]
    },
    'Slot Machines': {
        'Responsibilities': 'Slot machine operations, machine maintenance, prize payouts',
        'Assessment Scenarios': [
            'Game rules explanation',
            'Machine operation guidance',
            'Payment method explanations',
            'Technical issue reporting'
        ]
    },
    'Casino Administration': {
        'Responsibilities': 'Casino oversight, compliance checks, financial management',
        'Assessment Scenarios': [
            'Regulatory restrictions - "Subject to maritime laws"',
            'Territorial water restrictions - "Suspended while in territorial waters"',
            'Operating hours policy - Adjustments per maritime law',
            'Compliance requirement communication'
        ]
    }
}

for dept_name, details in casino_depts.items():
    doc.add_heading(dept_name, level=4)
    doc.add_paragraph(f"Responsibilities: {details['Responsibilities']}")
    doc.add_paragraph('Assessment Scenarios:')
    for scenario in details['Assessment Scenarios']:
        doc.add_paragraph(scenario, style='List Bullet 2')
    doc.add_paragraph()

# ==================== SCENARIO CATEGORIES ====================
doc.add_page_break()
doc.add_heading('III. Assessment Scenario Categories', level=1)

# Scenario categories
doc.add_heading('1. Customer Service Scenarios', level=2)
service_scenarios = [
    'Reservations & Inquiries - Restaurant bookings, activity registrations, spa appointments',
    'Complaint Handling - Room AC issues, facility malfunctions, service dissatisfaction',
    'Information Queries - Operating hours, location directions, activity schedules',
    'Special Needs - Dietary restrictions, medical assistance, accessibility services'
]
for scenario in service_scenarios:
    doc.add_paragraph(scenario, style='List Bullet')

doc.add_heading('2. Time & Numbers Scenarios', level=2)
time_scenarios = [
    'Operating Hours - Breakfast 7:00 AM - 10:30 AM, Casino open until 2:00 AM',
    'Reservation Information - 8 guests at 6:30 PM for dinner',
    'Room Numbers - Room 8254, Cabin 9173',
    'Deck Levels - Deck 12, Deck 14'
]
for scenario in time_scenarios:
    doc.add_paragraph(scenario, style='List Bullet')

doc.add_heading('3. Safety & Emergency Scenarios', level=2)
safety_scenarios = [
    'Safety Drills - Mandatory muster drill at 4:00 PM',
    'Emergency Equipment - Life jackets, assembly stations, escape routes',
    'Missed Ship Procedures - Contact port agent, reach next port independently',
    'Safety Announcements - Maritime law restrictions'
]
for scenario in safety_scenarios:
    doc.add_paragraph(scenario, style='List Bullet')

doc.add_heading('4. Technical Terminology Scenarios', level=2)
term_scenarios = [
    'Vessel Terms - Bridge (wheelhouse), Gangway (boarding ramp), Tender (shuttle boat)',
    'F&B Terms - Galley (kitchen), Buffet, Sommelier (wine expert)',
    'Service Terms - Concierge (guest services), Amenities (facilities), Excursion (shore tour)',
    'Safety Terms - Muster (assembly), Assembly station, Embark (board ship)'
]
for scenario in term_scenarios:
    doc.add_paragraph(scenario, style='List Bullet')

# ==================== QUESTION MAPPING ====================
doc.add_page_break()
doc.add_heading('IV. Question-Scenario Mapping', level=1)

question_mapping = [
    ['Module', 'Q#', 'Scenario', 'Related Department'],
    ['Listening', 'Q1', 'Restaurant reservation (4 people at 7 PM)', 'Food & Beverage'],
    ['Listening', 'Q2', 'Room AC repair (Room 8254)', 'Housekeeping, Engineering'],
    ['Listening', 'Q3', 'Buffet location inquiry (Deck 12)', 'Guest Services, Entertainment'],
    ['Time/Numbers', 'Q4', 'Breakfast hours (7:00-10:30)', 'Food & Beverage'],
    ['Time/Numbers', 'Q5', 'Reservation count (8 people, 6:30 PM)', 'Food & Beverage'],
    ['Time/Numbers', 'Q6', 'Cabin number (9173 safe malfunction)', 'Housekeeping, Engineering'],
    ['Grammar', 'Q7', 'Luggage service politeness', 'Guest Services'],
    ['Grammar', 'Q8', 'Guest arrival tense', 'Guest Services'],
    ['Grammar', 'Q9', 'Spa direction verb', 'Spa & Fitness'],
    ['Grammar', 'Q10', 'Restaurant location preposition', 'Food & Beverage'],
    ['Vocabulary', 'Q11', 'Vessel terminology matching', 'Deck Department'],
    ['Vocabulary', 'Q12', 'Hospitality service terms', 'Guest Services'],
    ['Vocabulary', 'Q13', 'Dining terminology', 'Food & Beverage, Culinary'],
    ['Vocabulary', 'Q14', 'Safety terminology', 'Safety & Security'],
    ['Reading', 'Q15', 'Missed ship policy', 'Shore Excursions'],
    ['Reading', 'Q16', 'Casino operating hours', 'Casino Operations'],
    ['Reading', 'Q17', 'Specialty restaurant reservations', 'Food & Beverage'],
    ['Reading', 'Q18', 'Muster drill schedule', 'Safety & Security'],
    ['Speaking', 'Q19', 'AC issue response', 'Housekeeping, Engineering'],
    ['Speaking', 'Q20', 'Buffet hours inquiry', 'Food & Beverage'],
    ['Speaking', 'Q21', 'Spa location directions', 'Spa & Fitness, Guest Services']
]

mapping_table = doc.add_table(rows=len(question_mapping), cols=len(question_mapping[0]))
mapping_table.style = 'Light List Accent 1'

for i, row_data in enumerate(question_mapping):
    row = mapping_table.rows[i]
    for j, cell_data in enumerate(row_data):
        cell = row.cells[j]
        cell.text = cell_data
        if i == 0:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(11)

# ==================== SCORING ====================
doc.add_page_break()
doc.add_heading('V. Scoring Standards', level=1)

doc.add_heading('Overall Score & Pass Criteria', level=2)
scoring_overview = """
Total Score: 100 points
Passing Score: 65 points (65%)
Safety-related Questions Pass Rate: 80%
Speaking Module Minimum Score: 12 points (60%)
"""
doc.add_paragraph(scoring_overview)

doc.add_heading('Module Scoring Details', level=2)
module_scores = [
    ['Module', 'Questions', 'Points Each', 'Total', 'Notes'],
    ['Listening', '3', '4', '12', 'Multiple choice, correct answer scores'],
    ['Time & Numbers', '3', '4', '12', 'Fill-in-blank, fully correct scores'],
    ['Grammar', '4', '4', '16', 'Multiple choice, correct answer scores'],
    ['Vocabulary', '4', '8', '32', 'Matching, 2 points per correct pair'],
    ['Reading', '4', '6', '24', 'Multiple choice, correct answer scores'],
    ['Speaking', '3', '20/3≈6.67', '20', 'AI-scored, comprehensive evaluation']
]

score_table = doc.add_table(rows=len(module_scores), cols=len(module_scores[0]))
score_table.style = 'Light Grid Accent 1'

for i, row_data in enumerate(module_scores):
    row = score_table.rows[i]
    for j, cell_data in enumerate(row_data):
        cell = row.cells[j]
        cell.text = str(cell_data)
        if i == 0:
            cell.paragraphs[0].runs[0].font.bold = True

# ==================== TECHNICAL IMPLEMENTATION ====================
doc.add_page_break()
doc.add_heading('VI. Technical Implementation', level=1)

doc.add_heading('Frontend Technologies', level=2)
frontend_tech = """
HTML5 + CSS3: Responsive design, mobile-friendly
JavaScript ES6+: Interactive logic, form validation, audio control
Web Speech API: Text-to-speech (TTS) for listening module
MediaRecorder API: Voice recording for speaking module
Drag & Drop API: Vocabulary matching module interaction
"""
doc.add_paragraph(frontend_tech)

doc.add_heading('Backend Technologies', level=2)
backend_tech = """
FastAPI: Python asynchronous web framework
SQLAlchemy: ORM database operations
PostgreSQL: Relational database
Redis: Session management and caching
AI Services: OpenAI/Anthropic Claude for speaking assessment
"""
doc.add_paragraph(backend_tech)

doc.add_heading('Security Features', level=2)
security_features = """
CSRF Protection: Prevent cross-site request forgery
Rate Limiting: Prevent abuse and DOS attacks
Input Validation: XSS and SQL injection protection
Session Encryption: Fernet symmetric encryption
Security Headers: HSTS, CSP, X-Frame-Options, etc.
"""
doc.add_paragraph(security_features)

# ==================== APPENDIX ====================
doc.add_page_break()
doc.add_heading('VII. Appendix: Key English Terminology', level=1)

doc.add_heading('A. Vessel Terminology', level=2)
vessel_terms = [
    'Bridge - Ship\'s control center/wheelhouse',
    'Gangway - Boarding ramp/walkway to shore',
    'Tender - Small shuttle boat for shore access',
    'Muster - Emergency assembly',
    'Deck - Ship floor/level',
    'Starboard - Right side of ship',
    'Port - Left side of ship',
    'Bow - Front of ship',
    'Stern - Rear of ship'
]
for term in vessel_terms:
    doc.add_paragraph(term, style='List Bullet')

doc.add_heading('B. Hospitality Terminology', level=2)
hospitality_terms = [
    'Concierge - Guest services specialist',
    'Amenities - Facilities/conveniences',
    'Housekeeping - Room cleaning services',
    'Turndown service - Evening bed preparation',
    'Stateroom - Guest cabin',
    'Suite - Luxury cabin with separate rooms',
    'Balcony cabin - Room with private balcony'
]
for term in hospitality_terms:
    doc.add_paragraph(term, style='List Bullet')

doc.add_heading('C. Food & Beverage Terminology', level=2)
fb_terms = [
    'Galley - Ship\'s kitchen',
    'Buffet - Self-service dining',
    'A la carte - Menu with individual item pricing',
    'Sommelier - Wine expert',
    'Main dining room - Primary restaurant',
    'Specialty restaurant - Premium dining venue',
    'Cover charge - Additional fee'
]
for term in fb_terms:
    doc.add_paragraph(term, style='List Bullet')

doc.add_heading('D. Safety Terminology', level=2)
safety_terms = [
    'Muster drill - Safety assembly practice',
    'Life jacket/vest - Personal flotation device',
    'Assembly station - Emergency gathering point',
    'All aboard - Final boarding call',
    'Lifeboat - Emergency evacuation vessel',
    'Emergency exit - Escape route',
    'Evacuation - Emergency disembarkation'
]
for term in safety_terms:
    doc.add_paragraph(term, style='List Bullet')

doc.add_heading('E. Travel Terminology', level=2)
travel_terms = [
    'Embark - Board the ship',
    'Disembark - Leave the ship',
    'Excursion - Shore tour/visit',
    'Port of call - Destination port',
    'Itinerary - Travel schedule',
    'Shore leave - Time allowed ashore',
    'Port Agent - Port representative'
]
for term in travel_terms:
    doc.add_paragraph(term, style='List Bullet')

# ==================== CLOSING ====================
doc.add_page_break()
doc.add_heading('Document Information', level=1)

notes = """
This document provides comprehensive details on the three major operational divisions (Hotel, Marine, Casino)
of the Cruise Employee English Assessment Platform, covering 14 specific departments and all 21 assessment
questions mapped to real-world work scenarios.

Assessment content is entirely based on actual cruise industry operational requirements, ensuring employees
who pass the assessment can competently handle English communication requirements in their respective positions.

For more technical implementation details or assessment procedures, please refer to the project GitHub repository:
https://github.com/Dino24-Max/English-Assessment

Generated: {}
Platform Version: v1.0
""".format(datetime.now().strftime("%B %d, %Y at %H:%M"))

doc.add_paragraph(notes)

# Save document
output_path = r"C:\Users\szh2051\OneDrive - Carnival Corporation\Desktop\Cruise-Employee-English-Assessment-Operations-Scenarios.docx"
doc.save(output_path)

print("Word document created successfully!")
print(f"File saved at: {output_path}")
print("Content: 3 major operations, 14 departments, 21 question scenarios")
print("Total pages: approximately 15-20 pages")
