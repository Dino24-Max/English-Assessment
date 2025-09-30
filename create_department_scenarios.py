"""
Create comprehensive scenario documentation for all 16 departments.
Each department will have 10 detailed scenarios.
Total: 160 scenarios across Hotel, Marine, and Casino Operations.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os


def create_chinese_scenarios_document():
    """Create Chinese version with 10 scenarios per department"""

    doc = Document()

    # Set document styling
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(11)

    # Title
    title = doc.add_heading('邮轮员工英语评估平台', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    title_run.font.bold = True

    # Subtitle
    subtitle = doc.add_heading('部门场景详细说明文档', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(18)
    subtitle_run.font.color.rgb = RGBColor(0, 102, 204)

    # Document info
    doc.add_paragraph('文档版本: 3.0')
    doc.add_paragraph('更新日期: 2025-09-30')
    doc.add_paragraph('总部门数: 16个部门')
    doc.add_paragraph('总场景数: 160个场景（每个部门10个）')
    doc.add_paragraph()

    # Overview
    doc.add_heading('一、概述', level=1)
    p = doc.add_paragraph('本文档为16个部门提供详细的工作场景说明，每个部门包含10个真实工作场景。这些场景将用于英语评估平台的题目设计，涵盖听力、口语、阅读、词汇和语法等多个模块。')
    doc.add_paragraph()

    # Hotel Operations Scenarios
    doc.add_heading('二、酒店运营部门场景 (Hotel Operations)', level=1)

    hotel_scenarios = {
        '前台接待 (Front Desk)': [
            {'title': '场景1：客人登船办理入住', 'description': '协助客人办理登船手续，验证身份文件，发放房卡，解释船上设施位置'},
            {'title': '场景2：处理房间升级请求', 'description': '客人要求升级到海景套房，查询可用房间，说明价格差异，办理升级手续'},
            {'title': '场景3：解答岸上观光问题', 'description': '客人咨询下一个港口的岸上游览选项，推荐热门景点，说明预订流程和集合时间'},
            {'title': '场景4：处理客人投诉', 'description': '客人对房间噪音问题投诉，记录详细信息，提供解决方案，安排换房或补偿'},
            {'title': '场景5：协助丢失物品查询', 'description': '客人报告房卡或随身物品丢失，核实身份，补发房卡或联系失物招领处'},
            {'title': '场景6：解释船上消费系统', 'description': '向客人说明船卡消费功能，如何查看账单，设置消费限额，结账流程'},
            {'title': '场景7：预订船上特色餐厅', 'description': '客人想预订高级餐厅，查询可用时间，说明着装要求，确认预订详情'},
            {'title': '场景8：处理医疗紧急情况', 'description': '客人需要医疗帮助，联系船上医务室，安排就诊时间，提供位置指引'},
            {'title': '场景9：协助离船手续', 'description': '解释离船流程，行李托运安排，海关申报要求，最后结账确认'},
            {'title': '场景10：处理特殊需求', 'description': '客人有饮食过敏或行动不便等特殊需求，记录详情，协调相关部门提供支持'}
        ],
        '客房服务 (Housekeeping)': [
            {'title': '场景1：日常客房清洁', 'description': '按标准流程清洁客房，更换床单毛巾，补充洗漱用品，整理房间'},
            {'title': '场景2：处理额外物品请求', 'description': '客人要求额外枕头、毛毯或衣架，及时送达并确认客人满意'},
            {'title': '场景3：处理客房维修报告', 'description': '发现房间设施损坏（如漏水、灯泡坏了），填写维修单，通知维修部门'},
            {'title': '场景4：提供开夜床服务', 'description': '晚间为客人提供开夜床服务，整理床铺，补充瓶装水，放置巧克力'},
            {'title': '场景5：清洁卫生间', 'description': '深度清洁卫生间，消毒马桶和洗手池，补充卫生纸和洗浴用品'},
            {'title': '场景6：处理客人隐私要求', 'description': '客人要求不打扰或特定时间清洁，记录偏好，按客人要求安排清洁时间'},
            {'title': '场景7：准备VIP客房', 'description': '为VIP客人准备特别欢迎礼品，摆放鲜花，确保房间完美无瑕'},
            {'title': '场景8：处理遗留物品', 'description': '清洁时发现客人遗留物品，按规定送交失物招领，填写物品清单'},
            {'title': '场景9：应对客人反馈', 'description': '客人对清洁质量有意见，道歉并立即重新清洁，确保达到标准'},
            {'title': '场景10：管理清洁用品库存', 'description': '检查清洁用品库存，填写补给申请单，确保物资充足'}
        ],
        '餐饮服务 (Food & Beverage)': [
            {'title': '场景1：早餐自助服务', 'description': '引导客人入座，介绍自助餐台布局，提供饮料服务，及时清理餐盘'},
            {'title': '场景2：推荐每日特色菜', 'description': '向客人介绍今日特色菜品，说明烹饪方式和食材，根据客人口味推荐'},
            {'title': '场景3：处理食物过敏问题', 'description': '客人有食物过敏（如海鲜、坚果），核实菜单成分，提供安全替代选项'},
            {'title': '场景4：提供儿童餐服务', 'description': '为带小孩的家庭提供儿童菜单，准备儿童座椅，提供蜡笔和涂色纸'},
            {'title': '场景5：处理菜品投诉', 'description': '客人对食物温度或口味不满，道歉并重新制作，必要时提供补偿'},
            {'title': '场景6：解释特殊饮食菜单', 'description': '为素食、清真或犹太饮食客人介绍专门菜单，确认饮食限制'},
            {'title': '场景7：提供葡萄酒配对建议', 'description': '根据客人点的主菜推荐合适的葡萄酒，介绍酒的产地和特点'},
            {'title': '场景8：处理大型团体用餐', 'description': '协调大型团体的座位安排，统一上菜时间，确保服务流畅'},
            {'title': '场景9：庆祝特殊场合', 'description': '为客人准备生日或周年庆祝，安排蛋糕和唱歌，拍照留念'},
            {'title': '场景10：解释餐厅营业时间', 'description': '向客人说明各餐厅的营业时间，预订要求，着装规定'}
        ],
        '酒吧服务 (Bar Service)': [
            {'title': '场景1：调制经典鸡尾酒', 'description': '根据客人要求调制玛格丽特、莫吉托等经典鸡尾酒，解释配料和制作过程'},
            {'title': '场景2：推荐船上特色饮品', 'description': '介绍船上独家特调饮品，说明口味特点，提供试饮小样'},
            {'title': '场景3：处理醉酒客人', 'description': '礼貌地拒绝为醉酒客人继续提供酒精饮料，提供水和小食，必要时通知安保'},
            {'title': '场景4：介绍欢乐时光优惠', 'description': '向客人说明欢乐时光的时间段和优惠饮品，推荐热门选择'},
            {'title': '场景5：提供无酒精饮品', 'description': '为不喝酒的客人或儿童推荐无酒精鸡尾酒（mocktails），制作果汁饮品'},
            {'title': '场景6：解释酒水套餐', 'description': '介绍船上的酒水套餐选项，说明价格和包含的饮品类型，帮助客人选择'},
            {'title': '场景7：处理饮品投诉', 'description': '客人觉得饮品太甜或太淡，道歉并重新调制，询问客人偏好'},
            {'title': '场景8：组织品酒活动', 'description': '主持威士忌或葡萄酒品鉴活动，介绍酒的历史和风味特点'},
            {'title': '场景9：管理酒吧库存', 'description': '检查酒精饮料和调酒材料库存，填写补给单，确保热门饮品不断货'},
            {'title': '场景10：创建定制鸡尾酒', 'description': '根据客人口味偏好创造个性化鸡尾酒，记录配方，命名并加入特色菜单'}
        ],
        '宾客服务 (Guest Services)': [
            {'title': '场景1：协助预订岸上游览', 'description': '向客人介绍各港口的岸上游览选项，处理预订，提供集合时间和地点'},
            {'title': '场景2：解答船上活动问题', 'description': '向客人介绍今日船上活动安排，如演出、健身课程、讲座等'},
            {'title': '场景3：处理邮轮积分计划', 'description': '解释忠诚度计划的积分累积和兑换规则，帮助客人注册或查询积分'},
            {'title': '场景4：协调特殊庆祝活动', 'description': '帮客人安排船上的生日派对、蜜月庆祝或周年纪念活动'},
            {'title': '场景5：提供港口信息', 'description': '向客人介绍下一个停靠港口的天气、货币、交通和注意事项'},
            {'title': '场景6：处理通讯服务', 'description': '解释船上WiFi套餐选项，协助购买上网时长，解决连接问题'},
            {'title': '场景7：安排Spa预约', 'description': '帮客人预约船上水疗中心的按摩或美容服务，说明项目和价格'},
            {'title': '场景8：解决客人纠纷', 'description': '调解客人之间的小纠纷（如座位争执），提供公平解决方案'},
            {'title': '场景9：提供翻译协助', 'description': '为不同语言的客人提供基本翻译帮助，协调多语言服务'},
            {'title': '场景10：处理紧急家庭联络', 'description': '协助客人在紧急情况下联系家人，安排紧急通讯，提供支持'}
        ],
        '客舱服务 (Cabin Service)': [
            {'title': '场景1：客房送餐服务', 'description': '接收客房送餐订单，确认菜品和送达时间，准时送餐到客房'},
            {'title': '场景2：早餐送舱服务', 'description': '处理客人悬挂在门把手上的早餐订单卡，按时送达热早餐'},
            {'title': '场景3：处理客舱维修请求', 'description': '接到客人关于空调、电视或其他设施的维修报告，联系维修部门，跟进进度'},
            {'title': '场景4：提供客舱深夜服务', 'description': '深夜为客人提供额外毛毯、枕头或小食品，确保不打扰其他客人'},
            {'title': '场景5：协调客舱特殊布置', 'description': '为客人生日或纪念日在客舱内布置气球、鲜花和贺卡'},
            {'title': '场景6：处理客舱噪音投诉', 'description': '接到客人关于邻舱噪音的投诉，礼貌提醒噪音来源，提供换房选项'},
            {'title': '场景7：解释客舱设施使用', 'description': '向客人演示如何使用保险箱、电视、温度控制等客舱设施'},
            {'title': '场景8：提供客舱清洁偏好', 'description': '记录客人的清洁时间偏好，是否需要开夜床，枕头软硬度选择'},
            {'title': '场景9：处理客舱紧急情况', 'description': '应对客舱内的医疗紧急情况或安全问题，联系相关部门，确保客人安全'},
            {'title': '场景10：协调客舱用品补充', 'description': '检查并补充客舱的咖啡、茶、瓶装水等物品，确保充足'}
        ],
        '辅助服务 (Auxiliary Service)': [
            {'title': '场景1：协调部门间物资调配', 'description': '根据各部门需求调配物资，确保运营顺畅，记录调配明细'},
            {'title': '场景2：处理紧急物资需求', 'description': '应对突发的物资短缺，快速协调备用库存，联系供应商补货'},
            {'title': '场景3：支持大型活动准备', 'description': '为船上大型活动（如船长晚宴）提供额外人手和物资支持'},
            {'title': '场景4：管理临时员工调配', 'description': '根据业务高峰期需求，协调临时员工支援繁忙部门'},
            {'title': '场景5：处理特殊服务请求', 'description': '应对VIP客人的特殊服务需求，如私人导游、专属服务员等'},
            {'title': '场景6：协调维修和保养', 'description': '安排船上设施的定期维护和紧急维修，协调时间避免影响客人'},
            {'title': '场景7：支持新员工培训', 'description': '为新员工提供岗位辅助，解答疑问，确保快速适应工作'},
            {'title': '场景8：管理物资运输路线', 'description': '优化船上物资运输路线，减少运输时间，提高效率'},
            {'title': '场景9：处理部门协调会议', 'description': '组织跨部门协调会议，解决运营中的协作问题'},
            {'title': '场景10：提供质量控制支持', 'description': '协助各部门进行服务质量检查，收集反馈，提出改进建议'}
        ],
        '洗衣房 (Laundry)': [
            {'title': '场景1：处理客人洗衣订单', 'description': '接收客人的洗衣袋，核对洗衣单，说明价格和取衣时间'},
            {'title': '场景2：提供快速洗衣服务', 'description': '客人需要当天取衣的加急服务，说明额外费用，优先处理'},
            {'title': '场景3：处理特殊面料衣物', 'description': '识别丝绸、羊毛等特殊面料，使用专门洗涤程序，确保不损坏'},
            {'title': '场景4：解释洗衣服务价格', 'description': '向客人说明按件收费和按袋收费的价格差异，推荐套餐服务'},
            {'title': '场景5：处理衣物污渍', 'description': '客人衣物有顽固污渍（如红酒、油渍），使用专业去污剂处理'},
            {'title': '场景6：提供熨烫和折叠服务', 'description': '为客人的正式服装提供专业熨烫，按要求折叠或悬挂'},
            {'title': '场景7：处理洗衣事故', 'description': '如衣物在洗涤过程中损坏或褪色，道歉并按规定赔偿'},
            {'title': '场景8：管理洗衣房设备', 'description': '检查洗衣机、烘干机运行状况，报告故障，安排维修'},
            {'title': '场景9：处理员工制服洗涤', 'description': '按时完成各部门员工制服的洗涤和配送，确保员工仪容整洁'},
            {'title': '场景10：提供干洗服务', 'description': '为需要干洗的礼服、西装提供专业干洗，确保质量'}
        ],
        '摄影服务 (Photo)': [
            {'title': '场景1：拍摄登船纪念照', 'description': '在登船处为客人拍摄欢迎照片，调整姿势和背景，确保照片质量'},
            {'title': '场景2：组织船长晚宴摄影', 'description': '在船长晚宴现场为客人和船长拍照，调整灯光，快速处理照片'},
            {'title': '场景3：销售照片套餐', 'description': '向客人介绍各种照片产品和套餐，展示样品，处理购买订单'},
            {'title': '场景4：提供家庭照服务', 'description': '为家庭客人安排专业摄影时段，选择合适背景，拍摄家庭合影'},
            {'title': '场景5：处理照片打印订单', 'description': '接收客人的数码照片打印请求，确认尺寸和数量，按时交付'},
            {'title': '场景6：组织主题摄影活动', 'description': '安排日落摄影、正式晚宴摄影等主题活动，宣传并协调参与者'},
            {'title': '场景7：提供照片编辑服务', 'description': '为客人提供基础照片编辑（如修色、裁剪），满足个性化需求'},
            {'title': '场景8：管理照片展示区', 'description': '在公共区域展示当天拍摄的照片，方便客人浏览和购买'},
            {'title': '场景9：处理照片投诉', 'description': '客人对照片质量不满，提供重拍或退款服务'},
            {'title': '场景10：提供数字照片下载', 'description': '协助客人通过APP或网站下载购买的数字照片，解决技术问题'}
        ],
        '物资供应 (Provisions)': [
            {'title': '场景1：管理食品库存', 'description': '检查冷藏和干货食品库存，记录过期日期，填写补给申请'},
            {'title': '场景2：协调港口物资补给', 'description': '在停靠港口时接收新鲜食材和物资，核对订单，检查质量'},
            {'title': '场景3：处理紧急食材需求', 'description': '餐厅临时缺少食材，快速从备用库存调配或联系供应商加急配送'},
            {'title': '场景4：管理饮料库存', 'description': '检查酒精和非酒精饮料库存，确保热门饮品充足，处理补给订单'},
            {'title': '场景5：控制食品安全', 'description': '确保食品储存温度符合标准，定期检查，记录温度日志'},
            {'title': '场景6：处理供应商订单', 'description': '根据各部门需求汇总采购清单，联系供应商下单，跟踪交付'},
            {'title': '场景7：管理清洁用品库存', 'description': '检查清洁剂、洗涤用品库存，确保各部门供应充足'},
            {'title': '场景8：处理物资浪费问题', 'description': '分析物资使用数据，识别浪费环节，提出节约建议'},
            {'title': '场景9：协调特殊饮食食材', 'description': '为特殊饮食需求（素食、无麸质等）采购专门食材'},
            {'title': '场景10：管理物资储藏空间', 'description': '优化仓库布局，确保物资分类清晰，方便快速取用'}
        ]
    }

    # Add Hotel Operations scenarios
    for dept_name, scenarios in hotel_scenarios.items():
        doc.add_heading(dept_name, level=2)
        for i, scenario in enumerate(scenarios, 1):
            p = doc.add_paragraph()
            run = p.add_run(scenario['title'])
            run.bold = True
            doc.add_paragraph(scenario['description'], style='List Bullet')
        doc.add_paragraph()

    doc.add_page_break()

    # Marine Operations Scenarios
    doc.add_heading('三、海务运营部门场景 (Marine Operations)', level=1)

    marine_scenarios = {
        '甲板部 (Deck Department)': [
            {'title': '场景1：执行船舶靠港操作', 'description': '协助船长指挥靠港，操作缆绳，与港口人员沟通，确保安全靠岸'},
            {'title': '场景2：进行安全演习', 'description': '组织救生艇演习，向客人讲解安全程序，演示救生衣穿戴'},
            {'title': '场景3：维护甲板设备', 'description': '检查和维护救生设备、锚链、舷梯等甲板设施，记录维护日志'},
            {'title': '场景4：监控航行安全', 'description': '值班时监控雷达和导航设备，报告船舶位置，注意海上交通'},
            {'title': '场景5：处理恶劣天气', 'description': '在风暴天气下加固甲板设备，关闭舱门，确保客人安全'},
            {'title': '场景6：协调引航员登船', 'description': '在港口协助引航员登船，提供梯子，确保安全交接'},
            {'title': '场景7：执行海上搜救', 'description': '参与海上搜救演习或实际搜救，操作救生艇，配合救援指令'},
            {'title': '场景8：维护船舶外观', 'description': '清洁和涂漆船体外部，保持船舶外观整洁，防止锈蚀'},
            {'title': '场景9：管理货物装卸', 'description': '在港口监督物资和行李装卸，确保安全操作，防止损坏'},
            {'title': '场景10：执行夜间值班', 'description': '夜间值班监控船舶周围环境，巡查甲板，记录值班日志'}
        ],
        '机舱部 (Engine Department)': [
            {'title': '场景1：监控主机运行', 'description': '在机舱控制室监控主推进系统，记录参数，及时发现异常'},
            {'title': '场景2：执行预防性维护', 'description': '按计划对发电机、泵浦等设备进行定期保养，更换滤芯和油液'},
            {'title': '场景3：处理设备故障', 'description': '设备报警时快速诊断问题，执行紧急修复，必要时切换备用系统'},
            {'title': '场景4：管理燃油供应', 'description': '监控燃油消耗，协调港口加油，检查燃油质量，记录库存'},
            {'title': '场景5：维护空调系统', 'description': '检查和维护船上中央空调系统，确保客舱和公共区域温度舒适'},
            {'title': '场景6：处理污水系统', 'description': '操作和维护污水处理系统，确保符合环保法规，记录处理数据'},
            {'title': '场景7：协调停机维修', 'description': '在港口停机时执行大型设备维修，协调外部技术人员，监督工作'},
            {'title': '场景8：管理备件库存', 'description': '检查机械备件库存，填写采购申请，确保关键备件充足'},
            {'title': '场景9：执行安全检查', 'description': '定期检查机舱消防设备、应急照明、逃生通道，确保安全标准'},
            {'title': '场景10：培训新机工', 'description': '指导新入职机工熟悉设备操作，讲解安全规程，监督实操练习'}
        ],
        '安全部 (Security Department)': [
            {'title': '场景1：执行登船安全检查', 'description': '在登船口检查客人身份和行李，使用安检设备，防止危险物品上船'},
            {'title': '场景2：巡逻公共区域', 'description': '定时巡查甲板、走廊、餐厅等公共区域，观察可疑行为，确保秩序'},
            {'title': '场景3：处理客人纠纷', 'description': '介入客人争吵或打斗，冷静劝阻，必要时隔离当事人，记录事件'},
            {'title': '场景4：监控安防系统', 'description': '在安全控制室监控摄像头画面，记录异常事件，协调现场响应'},
            {'title': '场景5：协助医疗紧急情况', 'description': '在医疗紧急情况下协助医务人员，维持现场秩序，疏散围观者'},
            {'title': '场景6：执行消防演习', 'description': '组织船员和客人参加消防演习，讲解逃生路线，检查灭火设备'},
            {'title': '场景7：调查失窃报告', 'description': '接到客人财物失窃报告，记录详情，调查监控录像，协调警方'},
            {'title': '场景8：管理人员通行控制', 'description': '控制员工通道和受限区域，检查员工证件，防止未授权进入'},
            {'title': '场景9：处理醉酒客人', 'description': '应对醉酒滋事客人，礼貌但坚定地维持秩序，护送回房或拘留'},
            {'title': '场景10：执行港口安全协调', 'description': '在港口与当地安保和海关合作，协调登离船安全，防止走私'}
        ]
    }

    # Add Marine Operations scenarios
    for dept_name, scenarios in marine_scenarios.items():
        doc.add_heading(dept_name, level=2)
        for i, scenario in enumerate(scenarios, 1):
            p = doc.add_paragraph()
            run = p.add_run(scenario['title'])
            run.bold = True
            doc.add_paragraph(scenario['description'], style='List Bullet')
        doc.add_paragraph()

    doc.add_page_break()

    # Casino Operations Scenarios
    doc.add_heading('四、娱乐场运营部门场景 (Casino Operations)', level=1)

    casino_scenarios = {
        '赌桌游戏 (Table Games)': [
            {'title': '场景1：操作二十一点赌桌', 'description': '发牌、计分、支付赢钱，解释规则，维持游戏公平性'},
            {'title': '场景2：教新客人玩轮盘', 'description': '向第一次玩的客人解释轮盘游戏规则，下注选项，赔率计算'},
            {'title': '场景3：处理筹码兑换', 'description': '为客人兑换现金为筹码，核对金额，确保准确无误'},
            {'title': '场景4：监控作弊行为', 'description': '观察玩家行为，识别可疑的作弊企图，及时报告主管'},
            {'title': '场景5：管理百家乐赌桌', 'description': '操作百家乐游戏，快速计算点数，支付赢家，保持游戏节奏'},
            {'title': '场景6：处理客人争议', 'description': '客人对游戏结果有异议，冷静解释规则，调取监控录像核实'},
            {'title': '场景7：执行赌桌开关流程', 'description': '开台前检查设备和筹码，关台后清点筹码，填写交接表'},
            {'title': '场景8：提供游戏技巧建议', 'description': '在不违反规则前提下，为客人提供基本策略建议'},
            {'title': '场景9：管理高额投注', 'description': '为VIP高额投注客人提供专属服务，保持专业和礼貌'},
            {'title': '场景10：执行换班交接', 'description': '与接班荷官交接赌桌情况，清点筹码，报告特殊情况'}
        ],
        '老虎机 (Slot Machines)': [
            {'title': '场景1：协助客人使用机器', 'description': '教客人如何插卡、投币、选择游戏，解释支付线和下注选项'},
            {'title': '场景2：处理机器故障', 'description': '机器卡币或死机时重启系统，联系技术人员，安抚客人情绪'},
            {'title': '场景3：支付大奖', 'description': '客人中大奖时核实金额，填写税表（如需要），安排现金或支票支付'},
            {'title': '场景4：介绍累积奖池', 'description': '向客人解释累积奖池机器的玩法，展示当前奖池金额'},
            {'title': '场景5：处理客人投诉', 'description': '客人认为机器未正确支付，核查机器记录，提供合理解释或补偿'},
            {'title': '场景6：补充机器纸币', 'description': '定期检查机器内纸币和硬币存量，及时补充，确保正常运行'},
            {'title': '场景7：推广新款机器', 'description': '向客人介绍新上线的老虎机主题，演示特色功能，吸引试玩'},
            {'title': '场景8：清洁和维护机器', 'description': '定期清洁机器屏幕和外壳，检查按钮和读卡器功能'},
            {'title': '场景9：监控机器区域', 'description': '巡视老虎机区域，观察可疑行为，防止欺诈和破坏'},
            {'title': '场景10：协助忠诚度计划', 'description': '帮客人注册会员卡，解释积分累积和兑换规则'}
        ],
        '娱乐场服务 (Casino Services)': [
            {'title': '场景1：接待新客人', 'description': '欢迎客人进入娱乐场，介绍布局和游戏选项，提供会员卡注册'},
            {'title': '场景2：管理会员计划', 'description': '处理会员注册、升级和福利兑换，解释各级别的优惠'},
            {'title': '场景3：组织娱乐场锦标赛', 'description': '安排老虎机或扑克锦标赛，宣传活动，协调参赛者'},
            {'title': '场景4：提供免费饮品', 'description': '为游戏中的客人提供免费酒水饮料，记录客人偏好'},
            {'title': '场景5：处理客人投诉', 'description': '接受客人对服务或游戏的投诉，记录详情，提供解决方案'},
            {'title': '场景6：推广娱乐场活动', 'description': '向客人介绍本周特别促销和抽奖活动，鼓励参与'},
            {'title': '场景7：协调VIP服务', 'description': '为高额投注客人提供专属接待，安排私人游戏室，提供礼遇'},
            {'title': '场景8：监控游戏区氛围', 'description': '确保娱乐场环境友好，音量适中，及时处理冲突'},
            {'title': '场景9：处理支付和兑现', 'description': '协助客人兑换筹码为现金，处理信用额度申请'},
            {'title': '场景10：收集客人反馈', 'description': '主动询问客人体验，收集改进建议，提交管理层'}
        ]
    }

    # Add Casino Operations scenarios
    for dept_name, scenarios in casino_scenarios.items():
        doc.add_heading(dept_name, level=2)
        for i, scenario in enumerate(scenarios, 1):
            p = doc.add_paragraph()
            run = p.add_run(scenario['title'])
            run.bold = True
            doc.add_paragraph(scenario['description'], style='List Bullet')
        doc.add_paragraph()

    # Footer
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_p.add_run('---  文档结束  ---')
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    footer_run.italic = True

    # Save document
    output_path = r"C:\Users\szh2051\OneDrive - Carnival Corporation\Desktop\邮轮部门场景详细说明.docx"
    doc.save(output_path)
    print("Chinese scenarios document created successfully!")
    return output_path


def create_english_scenarios_document():
    """Create English version with 10 scenarios per department"""

    doc = Document()

    # Set document styling
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title
    title = doc.add_heading('Cruise Employee English Assessment Platform', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    title_run.font.bold = True

    # Subtitle
    subtitle = doc.add_heading('Department Scenarios Detailed Documentation', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(18)
    subtitle_run.font.color.rgb = RGBColor(0, 102, 204)

    # Document info
    doc.add_paragraph('Document Version: 3.0')
    doc.add_paragraph('Last Updated: September 30, 2025')
    doc.add_paragraph('Total Departments: 16')
    doc.add_paragraph('Total Scenarios: 160 (10 per department)')
    doc.add_paragraph()

    # Overview
    doc.add_heading('1. Overview', level=1)
    p = doc.add_paragraph('This document provides detailed work scenarios for all 16 departments, with 10 real-world scenarios per department. These scenarios will be used for designing questions in the English assessment platform, covering listening, speaking, reading, vocabulary, and grammar modules.')
    doc.add_paragraph()

    # Hotel Operations Scenarios
    doc.add_heading('2. Hotel Operations Department Scenarios', level=1)

    hotel_scenarios = {
        'Front Desk': [
            {'title': 'Scenario 1: Guest Embarkation Check-in', 'description': 'Assist guests with embarkation procedures, verify identification documents, issue room keys, explain ship facility locations'},
            {'title': 'Scenario 2: Processing Room Upgrade Requests', 'description': 'Guest requests upgrade to ocean view suite, check availability, explain price difference, process upgrade'},
            {'title': 'Scenario 3: Answering Shore Excursion Questions', 'description': 'Guest inquires about shore excursion options at next port, recommend popular attractions, explain booking process and meeting times'},
            {'title': 'Scenario 4: Handling Guest Complaints', 'description': 'Guest complains about room noise issue, record detailed information, provide solutions, arrange room change or compensation'},
            {'title': 'Scenario 5: Assisting with Lost Item Inquiries', 'description': 'Guest reports lost room key or personal belongings, verify identity, reissue room key or contact lost and found'},
            {'title': 'Scenario 6: Explaining Onboard Billing System', 'description': 'Explain ship card billing function to guest, how to check statements, set spending limits, checkout process'},
            {'title': 'Scenario 7: Booking Specialty Restaurants', 'description': 'Guest wants to book fine dining restaurant, check available times, explain dress code, confirm reservation details'},
            {'title': 'Scenario 8: Handling Medical Emergencies', 'description': 'Guest needs medical assistance, contact ship medical center, arrange appointment, provide location directions'},
            {'title': 'Scenario 9: Assisting with Disembarkation', 'description': 'Explain disembarkation process, luggage handling arrangements, customs declaration requirements, final bill confirmation'},
            {'title': 'Scenario 10: Processing Special Needs', 'description': 'Guest has special needs like food allergies or mobility issues, record details, coordinate with relevant departments for support'}
        ],
        'Housekeeping': [
            {'title': 'Scenario 1: Daily Cabin Cleaning', 'description': 'Clean cabin according to standard procedures, change linens and towels, replenish toiletries, organize room'},
            {'title': 'Scenario 2: Processing Extra Item Requests', 'description': 'Guest requests extra pillows, blankets, or hangers, deliver promptly and confirm guest satisfaction'},
            {'title': 'Scenario 3: Reporting Cabin Maintenance Issues', 'description': 'Discover room facility damage (like leaks, broken lights), fill out repair form, notify maintenance department'},
            {'title': 'Scenario 4: Providing Turndown Service', 'description': 'Provide evening turndown service for guests, arrange bed, replenish bottled water, place chocolate'},
            {'title': 'Scenario 5: Cleaning Bathrooms', 'description': 'Deep clean bathrooms, disinfect toilets and sinks, replenish toilet paper and bath amenities'},
            {'title': 'Scenario 6: Handling Guest Privacy Requests', 'description': 'Guest requests do not disturb or specific cleaning times, record preferences, arrange cleaning according to guest wishes'},
            {'title': 'Scenario 7: Preparing VIP Cabins', 'description': 'Prepare special welcome gifts for VIP guests, arrange flowers, ensure room is perfect'},
            {'title': 'Scenario 8: Handling Left Behind Items', 'description': 'Find items left by guests during cleaning, submit to lost and found per policy, fill out item list'},
            {'title': 'Scenario 9: Responding to Guest Feedback', 'description': 'Guest has comments about cleaning quality, apologize and re-clean immediately, ensure standards are met'},
            {'title': 'Scenario 10: Managing Cleaning Supply Inventory', 'description': 'Check cleaning supply inventory, fill out resupply request, ensure adequate materials'}
        ],
        'Food & Beverage': [
            {'title': 'Scenario 1: Breakfast Buffet Service', 'description': 'Guide guests to seats, introduce buffet layout, provide beverage service, clear plates promptly'},
            {'title': 'Scenario 2: Recommending Daily Specials', 'description': 'Introduce today\'s special dishes to guests, explain cooking methods and ingredients, recommend based on taste'},
            {'title': 'Scenario 3: Handling Food Allergy Issues', 'description': 'Guest has food allergies (like seafood, nuts), verify menu ingredients, provide safe alternatives'},
            {'title': 'Scenario 4: Providing Children\'s Meal Service', 'description': 'Provide children\'s menu for families with kids, prepare high chairs, offer crayons and coloring sheets'},
            {'title': 'Scenario 5: Handling Food Complaints', 'description': 'Guest unhappy with food temperature or taste, apologize and remake dish, offer compensation if necessary'},
            {'title': 'Scenario 6: Explaining Special Diet Menus', 'description': 'Introduce specialized menus for vegetarian, halal, or kosher guests, confirm dietary restrictions'},
            {'title': 'Scenario 7: Providing Wine Pairing Suggestions', 'description': 'Recommend appropriate wine based on guest\'s main course, introduce wine origin and characteristics'},
            {'title': 'Scenario 8: Managing Large Group Dining', 'description': 'Coordinate seating for large groups, synchronize service timing, ensure smooth operations'},
            {'title': 'Scenario 9: Celebrating Special Occasions', 'description': 'Prepare birthday or anniversary celebrations for guests, arrange cake and singing, take photos'},
            {'title': 'Scenario 10: Explaining Restaurant Operating Hours', 'description': 'Explain operating hours of various restaurants to guests, reservation requirements, dress codes'}
        ],
        'Bar Service': [
            {'title': 'Scenario 1: Mixing Classic Cocktails', 'description': 'Mix classic cocktails like Margaritas and Mojitos per guest request, explain ingredients and preparation'},
            {'title': 'Scenario 2: Recommending Signature Ship Drinks', 'description': 'Introduce exclusive specialty drinks on ship, explain flavor profiles, offer tasting samples'},
            {'title': 'Scenario 3: Handling Intoxicated Guests', 'description': 'Politely refuse to serve more alcohol to intoxicated guests, provide water and snacks, notify security if necessary'},
            {'title': 'Scenario 4: Introducing Happy Hour Specials', 'description': 'Explain happy hour time slots and featured drinks to guests, recommend popular choices'},
            {'title': 'Scenario 5: Providing Non-Alcoholic Beverages', 'description': 'Recommend mocktails for non-drinking guests or children, make juice drinks'},
            {'title': 'Scenario 6: Explaining Beverage Packages', 'description': 'Introduce beverage package options on ship, explain pricing and included drink types, help guest choose'},
            {'title': 'Scenario 7: Handling Drink Complaints', 'description': 'Guest finds drink too sweet or weak, apologize and remake, ask about preferences'},
            {'title': 'Scenario 8: Organizing Tasting Events', 'description': 'Host whiskey or wine tasting events, introduce history and flavor characteristics'},
            {'title': 'Scenario 9: Managing Bar Inventory', 'description': 'Check alcohol and mixing material inventory, fill out resupply orders, ensure popular drinks are stocked'},
            {'title': 'Scenario 10: Creating Custom Cocktails', 'description': 'Create personalized cocktails based on guest taste preferences, record recipe, name and add to specialty menu'}
        ],
        'Guest Services': [
            {'title': 'Scenario 1: Assisting with Shore Excursion Bookings', 'description': 'Introduce shore excursion options at various ports to guests, process bookings, provide meeting times and locations'},
            {'title': 'Scenario 2: Answering Onboard Activity Questions', 'description': 'Introduce today\'s onboard activities to guests, such as shows, fitness classes, lectures'},
            {'title': 'Scenario 3: Managing Cruise Loyalty Program', 'description': 'Explain loyalty program point accumulation and redemption rules, help guests register or check points'},
            {'title': 'Scenario 4: Coordinating Special Celebrations', 'description': 'Help guests arrange birthday parties, honeymoon celebrations, or anniversary events on ship'},
            {'title': 'Scenario 5: Providing Port Information', 'description': 'Introduce weather, currency, transportation, and precautions for next port to guests'},
            {'title': 'Scenario 6: Handling Communication Services', 'description': 'Explain WiFi package options on ship, assist with purchasing internet time, resolve connection issues'},
            {'title': 'Scenario 7: Arranging Spa Appointments', 'description': 'Help guests book massage or beauty services at ship spa center, explain services and pricing'},
            {'title': 'Scenario 8: Resolving Guest Disputes', 'description': 'Mediate minor disputes between guests (like seating conflicts), provide fair solutions'},
            {'title': 'Scenario 9: Providing Translation Assistance', 'description': 'Provide basic translation help for guests speaking different languages, coordinate multilingual services'},
            {'title': 'Scenario 10: Handling Emergency Family Contacts', 'description': 'Assist guests with contacting family in emergencies, arrange emergency communications, provide support'}
        ],
        'Cabin Service': [
            {'title': 'Scenario 1: In-Cabin Dining Service', 'description': 'Receive in-cabin dining orders, confirm dishes and delivery time, deliver meals to cabin on time'},
            {'title': 'Scenario 2: Breakfast In-Cabin Delivery', 'description': 'Process breakfast order cards hung on door handles by guests, deliver hot breakfast on time'},
            {'title': 'Scenario 3: Processing Cabin Repair Requests', 'description': 'Receive guest reports about AC, TV, or other facility repairs, contact maintenance, follow up on progress'},
            {'title': 'Scenario 4: Providing Late Night Cabin Service', 'description': 'Provide extra blankets, pillows, or snacks to guests late at night, ensure not disturbing other guests'},
            {'title': 'Scenario 5: Coordinating Special Cabin Arrangements', 'description': 'Arrange balloons, flowers, and cards in cabin for guest birthdays or anniversaries'},
            {'title': 'Scenario 6: Handling Cabin Noise Complaints', 'description': 'Receive guest complaints about neighbor cabin noise, politely remind noise source, offer room change option'},
            {'title': 'Scenario 7: Explaining Cabin Amenity Usage', 'description': 'Demonstrate to guests how to use safe, TV, temperature control, and other cabin amenities'},
            {'title': 'Scenario 8: Recording Cabin Cleaning Preferences', 'description': 'Record guest cleaning time preferences, whether turndown service needed, pillow firmness choices'},
            {'title': 'Scenario 9: Handling Cabin Emergencies', 'description': 'Respond to medical emergencies or safety issues in cabins, contact relevant departments, ensure guest safety'},
            {'title': 'Scenario 10: Coordinating Cabin Amenity Replenishment', 'description': 'Check and replenish cabin coffee, tea, bottled water, and other items, ensure adequate supply'}
        ],
        'Auxiliary Service': [
            {'title': 'Scenario 1: Coordinating Inter-Department Supply Distribution', 'description': 'Distribute supplies according to department needs, ensure smooth operations, record distribution details'},
            {'title': 'Scenario 2: Handling Emergency Supply Needs', 'description': 'Respond to sudden supply shortages, quickly coordinate backup inventory, contact suppliers for replenishment'},
            {'title': 'Scenario 3: Supporting Large Event Preparation', 'description': 'Provide extra manpower and supply support for large ship events (like captain\'s dinner)'},
            {'title': 'Scenario 4: Managing Temporary Staff Allocation', 'description': 'Coordinate temporary staff support for busy departments according to peak demand'},
            {'title': 'Scenario 5: Handling Special Service Requests', 'description': 'Respond to VIP guest special service needs, like private guides, dedicated service staff'},
            {'title': 'Scenario 6: Coordinating Maintenance and Repairs', 'description': 'Arrange regular maintenance and emergency repairs of ship facilities, coordinate timing to avoid guest impact'},
            {'title': 'Scenario 7: Supporting New Employee Training', 'description': 'Provide position assistance to new employees, answer questions, ensure quick adaptation to work'},
            {'title': 'Scenario 8: Managing Supply Transport Routes', 'description': 'Optimize ship supply transport routes, reduce transport time, improve efficiency'},
            {'title': 'Scenario 9: Managing Inter-Department Coordination Meetings', 'description': 'Organize cross-department coordination meetings, resolve operational collaboration issues'},
            {'title': 'Scenario 10: Providing Quality Control Support', 'description': 'Assist departments with service quality inspections, collect feedback, propose improvements'}
        ],
        'Laundry': [
            {'title': 'Scenario 1: Processing Guest Laundry Orders', 'description': 'Receive guest laundry bags, verify laundry list, explain pricing and pickup time'},
            {'title': 'Scenario 2: Providing Express Laundry Service', 'description': 'Guest needs same-day pickup express service, explain extra charges, prioritize processing'},
            {'title': 'Scenario 3: Handling Special Fabric Garments', 'description': 'Identify special fabrics like silk and wool, use specialized washing procedures, ensure no damage'},
            {'title': 'Scenario 4: Explaining Laundry Service Pricing', 'description': 'Explain price difference between per-item and per-bag charging to guests, recommend package services'},
            {'title': 'Scenario 5: Treating Garment Stains', 'description': 'Guest garments have stubborn stains (like wine, grease), use professional stain removers'},
            {'title': 'Scenario 6: Providing Ironing and Folding Service', 'description': 'Provide professional ironing for guest formal wear, fold or hang as requested'},
            {'title': 'Scenario 7: Handling Laundry Accidents', 'description': 'If garments damaged or faded during washing, apologize and compensate per policy'},
            {'title': 'Scenario 8: Managing Laundry Room Equipment', 'description': 'Check washing machine and dryer operation, report malfunctions, arrange repairs'},
            {'title': 'Scenario 9: Processing Staff Uniform Laundry', 'description': 'Complete washing and delivery of staff uniforms for all departments on time, ensure staff appearance is neat'},
            {'title': 'Scenario 10: Providing Dry Cleaning Service', 'description': 'Provide professional dry cleaning for evening gowns and suits requiring dry cleaning, ensure quality'}
        ],
        'Photo': [
            {'title': 'Scenario 1: Taking Embarkation Souvenir Photos', 'description': 'Take welcome photos for guests at embarkation area, adjust poses and background, ensure photo quality'},
            {'title': 'Scenario 2: Organizing Captain\'s Dinner Photography', 'description': 'Take photos of guests with captain at captain\'s dinner venue, adjust lighting, quickly process photos'},
            {'title': 'Scenario 3: Selling Photo Packages', 'description': 'Introduce various photo products and packages to guests, display samples, process purchase orders'},
            {'title': 'Scenario 4: Providing Family Photo Service', 'description': 'Arrange professional photography sessions for family guests, select appropriate backgrounds, take family portraits'},
            {'title': 'Scenario 5: Processing Photo Printing Orders', 'description': 'Receive guest digital photo printing requests, confirm sizes and quantities, deliver on time'},
            {'title': 'Scenario 6: Organizing Themed Photography Events', 'description': 'Arrange themed activities like sunset photography and formal dinner photography, promote and coordinate participants'},
            {'title': 'Scenario 7: Providing Photo Editing Service', 'description': 'Provide basic photo editing for guests (like color correction, cropping), meet personalized needs'},
            {'title': 'Scenario 8: Managing Photo Display Area', 'description': 'Display photos taken that day in public areas, facilitate guest browsing and purchasing'},
            {'title': 'Scenario 9: Handling Photo Complaints', 'description': 'Guest unhappy with photo quality, provide retake or refund service'},
            {'title': 'Scenario 10: Providing Digital Photo Downloads', 'description': 'Assist guests with downloading purchased digital photos via app or website, resolve technical issues'}
        ],
        'Provisions': [
            {'title': 'Scenario 1: Managing Food Inventory', 'description': 'Check refrigerated and dry goods food inventory, record expiration dates, fill out resupply requests'},
            {'title': 'Scenario 2: Coordinating Port Supply Replenishment', 'description': 'Receive fresh ingredients and supplies when docking at ports, verify orders, check quality'},
            {'title': 'Scenario 3: Handling Emergency Ingredient Needs', 'description': 'Restaurant temporarily short on ingredients, quickly allocate from backup inventory or contact suppliers for rush delivery'},
            {'title': 'Scenario 4: Managing Beverage Inventory', 'description': 'Check alcoholic and non-alcoholic beverage inventory, ensure popular drinks are adequately stocked, process resupply orders'},
            {'title': 'Scenario 5: Controlling Food Safety', 'description': 'Ensure food storage temperatures meet standards, regular inspections, record temperature logs'},
            {'title': 'Scenario 6: Processing Supplier Orders', 'description': 'Compile purchasing lists based on department needs, contact suppliers to place orders, track delivery'},
            {'title': 'Scenario 7: Managing Cleaning Supply Inventory', 'description': 'Check detergent and cleaning product inventory, ensure adequate supply for all departments'},
            {'title': 'Scenario 8: Addressing Supply Waste Issues', 'description': 'Analyze supply usage data, identify waste areas, propose conservation recommendations'},
            {'title': 'Scenario 9: Coordinating Special Diet Ingredients', 'description': 'Purchase specialized ingredients for special dietary needs (vegetarian, gluten-free, etc.)'},
            {'title': 'Scenario 10: Managing Supply Storage Space', 'description': 'Optimize warehouse layout, ensure supply classification is clear, facilitate quick retrieval'}
        ]
    }

    # Add Hotel Operations scenarios
    for dept_name, scenarios in hotel_scenarios.items():
        doc.add_heading(dept_name, level=2)
        for i, scenario in enumerate(scenarios, 1):
            p = doc.add_paragraph()
            run = p.add_run(scenario['title'])
            run.bold = True
            doc.add_paragraph(scenario['description'], style='List Bullet')
        doc.add_paragraph()

    doc.add_page_break()

    # Marine Operations Scenarios
    doc.add_heading('3. Marine Operations Department Scenarios', level=1)

    marine_scenarios = {
        'Deck Department': [
            {'title': 'Scenario 1: Executing Ship Docking Operations', 'description': 'Assist captain with docking maneuvers, operate mooring lines, communicate with port personnel, ensure safe berthing'},
            {'title': 'Scenario 2: Conducting Safety Drills', 'description': 'Organize lifeboat drills, explain safety procedures to guests, demonstrate life jacket usage'},
            {'title': 'Scenario 3: Maintaining Deck Equipment', 'description': 'Inspect and maintain lifesaving equipment, anchor chains, gangways and other deck facilities, record maintenance logs'},
            {'title': 'Scenario 4: Monitoring Navigation Safety', 'description': 'Monitor radar and navigation equipment during watch, report vessel position, observe maritime traffic'},
            {'title': 'Scenario 5: Handling Severe Weather', 'description': 'Secure deck equipment during storm weather, close hatches, ensure guest safety'},
            {'title': 'Scenario 6: Coordinating Pilot Boarding', 'description': 'Assist pilot boarding at port, provide ladder, ensure safe handover'},
            {'title': 'Scenario 7: Executing Sea Rescue', 'description': 'Participate in sea rescue drills or actual rescues, operate lifeboats, follow rescue commands'},
            {'title': 'Scenario 8: Maintaining Ship Appearance', 'description': 'Clean and paint ship exterior, maintain neat appearance, prevent rust'},
            {'title': 'Scenario 9: Managing Cargo Loading/Unloading', 'description': 'Supervise supply and luggage loading/unloading at ports, ensure safe operations, prevent damage'},
            {'title': 'Scenario 10: Executing Night Watch', 'description': 'Monitor ship surroundings during night watch, patrol deck, record watch log'}
        ],
        'Engine Department': [
            {'title': 'Scenario 1: Monitoring Main Engine Operation', 'description': 'Monitor main propulsion system in engine control room, record parameters, detect anomalies promptly'},
            {'title': 'Scenario 2: Executing Preventive Maintenance', 'description': 'Perform scheduled maintenance on generators, pumps and other equipment, replace filters and oils'},
            {'title': 'Scenario 3: Handling Equipment Failures', 'description': 'Quickly diagnose issues when equipment alarms, perform emergency repairs, switch to backup systems if necessary'},
            {'title': 'Scenario 4: Managing Fuel Supply', 'description': 'Monitor fuel consumption, coordinate port refueling, check fuel quality, record inventory'},
            {'title': 'Scenario 5: Maintaining HVAC Systems', 'description': 'Inspect and maintain ship central air conditioning, ensure comfortable cabin and public area temperatures'},
            {'title': 'Scenario 6: Managing Sewage Systems', 'description': 'Operate and maintain sewage treatment system, ensure environmental compliance, record treatment data'},
            {'title': 'Scenario 7: Coordinating Shutdown Maintenance', 'description': 'Execute major equipment repairs during port shutdown, coordinate external technicians, supervise work'},
            {'title': 'Scenario 8: Managing Spare Parts Inventory', 'description': 'Check mechanical spare parts inventory, fill out purchase requests, ensure critical parts are adequate'},
            {'title': 'Scenario 9: Executing Safety Inspections', 'description': 'Regularly inspect engine room fire equipment, emergency lighting, escape routes, ensure safety standards'},
            {'title': 'Scenario 10: Training New Engineers', 'description': 'Guide new engineers on equipment operation, explain safety protocols, supervise practical exercises'}
        ],
        'Security Department': [
            {'title': 'Scenario 1: Executing Embarkation Security Checks', 'description': 'Check guest identification and luggage at embarkation, use screening equipment, prevent dangerous items on board'},
            {'title': 'Scenario 2: Patrolling Public Areas', 'description': 'Regularly patrol deck, corridors, restaurants and other public areas, observe suspicious behavior, maintain order'},
            {'title': 'Scenario 3: Handling Guest Disputes', 'description': 'Intervene in guest arguments or fights, calmly mediate, separate parties if necessary, record incident'},
            {'title': 'Scenario 4: Monitoring Security Systems', 'description': 'Monitor camera feeds in security control room, record unusual incidents, coordinate on-site response'},
            {'title': 'Scenario 5: Assisting Medical Emergencies', 'description': 'Assist medical personnel during medical emergencies, maintain scene order, disperse onlookers'},
            {'title': 'Scenario 6: Executing Fire Drills', 'description': 'Organize crew and guest participation in fire drills, explain escape routes, check firefighting equipment'},
            {'title': 'Scenario 7: Investigating Theft Reports', 'description': 'Receive guest reports of stolen belongings, record details, investigate surveillance footage, coordinate with police'},
            {'title': 'Scenario 8: Managing Personnel Access Control', 'description': 'Control staff passages and restricted areas, check employee credentials, prevent unauthorized entry'},
            {'title': 'Scenario 9: Handling Intoxicated Guests', 'description': 'Deal with disruptive intoxicated guests, politely but firmly maintain order, escort to room or detain'},
            {'title': 'Scenario 10: Executing Port Security Coordination', 'description': 'Cooperate with local security and customs at ports, coordinate embarkation/disembarkation security, prevent smuggling'}
        ]
    }

    # Add Marine Operations scenarios
    for dept_name, scenarios in marine_scenarios.items():
        doc.add_heading(dept_name, level=2)
        for i, scenario in enumerate(scenarios, 1):
            p = doc.add_paragraph()
            run = p.add_run(scenario['title'])
            run.bold = True
            doc.add_paragraph(scenario['description'], style='List Bullet')
        doc.add_paragraph()

    doc.add_page_break()

    # Casino Operations Scenarios
    doc.add_heading('4. Casino Operations Department Scenarios', level=1)

    casino_scenarios = {
        'Table Games': [
            {'title': 'Scenario 1: Operating Blackjack Table', 'description': 'Deal cards, count points, pay winnings, explain rules, maintain game fairness'},
            {'title': 'Scenario 2: Teaching New Guests Roulette', 'description': 'Explain roulette game rules to first-time players, betting options, odds calculation'},
            {'title': 'Scenario 3: Processing Chip Exchanges', 'description': 'Exchange cash for chips for guests, verify amounts, ensure accuracy'},
            {'title': 'Scenario 4: Monitoring Cheating Behavior', 'description': 'Observe player behavior, identify suspicious cheating attempts, report to supervisor promptly'},
            {'title': 'Scenario 5: Managing Baccarat Table', 'description': 'Operate baccarat game, quickly calculate points, pay winners, maintain game pace'},
            {'title': 'Scenario 6: Handling Guest Disputes', 'description': 'Guest disputes game results, calmly explain rules, review surveillance footage to verify'},
            {'title': 'Scenario 7: Executing Table Opening/Closing Procedures', 'description': 'Check equipment and chips before opening, count chips after closing, fill out handover forms'},
            {'title': 'Scenario 8: Providing Game Strategy Advice', 'description': 'Provide basic strategy advice for guests without violating rules'},
            {'title': 'Scenario 9: Managing High-Stakes Betting', 'description': 'Provide exclusive service for VIP high-stakes guests, maintain professionalism and courtesy'},
            {'title': 'Scenario 10: Executing Shift Handover', 'description': 'Hand over table status to relief dealer, count chips, report special situations'}
        ],
        'Slot Machines': [
            {'title': 'Scenario 1: Assisting Guests with Machine Use', 'description': 'Teach guests how to insert cards, coins, select games, explain paylines and betting options'},
            {'title': 'Scenario 2: Handling Machine Malfunctions', 'description': 'Restart system when machine jams coins or freezes, contact technicians, calm guest emotions'},
            {'title': 'Scenario 3: Paying Jackpots', 'description': 'Verify amount when guest wins jackpot, fill out tax forms (if needed), arrange cash or check payment'},
            {'title': 'Scenario 4: Introducing Progressive Jackpots', 'description': 'Explain progressive jackpot machine gameplay to guests, display current jackpot amount'},
            {'title': 'Scenario 5: Handling Guest Complaints', 'description': 'Guest believes machine did not pay correctly, check machine records, provide reasonable explanation or compensation'},
            {'title': 'Scenario 6: Replenishing Machine Currency', 'description': 'Regularly check machine bill and coin levels, replenish promptly, ensure normal operation'},
            {'title': 'Scenario 7: Promoting New Machines', 'description': 'Introduce new slot machine themes to guests, demonstrate special features, attract trial play'},
            {'title': 'Scenario 8: Cleaning and Maintaining Machines', 'description': 'Regularly clean machine screens and casings, check button and card reader functionality'},
            {'title': 'Scenario 9: Monitoring Machine Area', 'description': 'Patrol slot machine area, observe suspicious behavior, prevent fraud and vandalism'},
            {'title': 'Scenario 10: Assisting Loyalty Program', 'description': 'Help guests register for membership cards, explain point accumulation and redemption rules'}
        ],
        'Casino Services': [
            {'title': 'Scenario 1: Welcoming New Guests', 'description': 'Welcome guests to casino, introduce layout and gaming options, provide membership card registration'},
            {'title': 'Scenario 2: Managing Membership Program', 'description': 'Process membership registration, upgrades and benefit redemption, explain tier benefits'},
            {'title': 'Scenario 3: Organizing Casino Tournaments', 'description': 'Arrange slot or poker tournaments, promote events, coordinate participants'},
            {'title': 'Scenario 4: Providing Complimentary Beverages', 'description': 'Provide free drinks to gaming guests, record guest preferences'},
            {'title': 'Scenario 5: Handling Guest Complaints', 'description': 'Accept guest complaints about service or games, record details, provide solutions'},
            {'title': 'Scenario 6: Promoting Casino Events', 'description': 'Introduce this week\'s special promotions and prize draws to guests, encourage participation'},
            {'title': 'Scenario 7: Coordinating VIP Services', 'description': 'Provide exclusive reception for high-stakes guests, arrange private gaming rooms, offer amenities'},
            {'title': 'Scenario 8: Monitoring Gaming Area Atmosphere', 'description': 'Ensure casino environment is friendly, volume is appropriate, handle conflicts promptly'},
            {'title': 'Scenario 9: Processing Payments and Cash-outs', 'description': 'Assist guests with exchanging chips for cash, process credit line applications'},
            {'title': 'Scenario 10: Collecting Guest Feedback', 'description': 'Proactively ask guests about their experience, collect improvement suggestions, submit to management'}
        ]
    }

    # Add Casino Operations scenarios
    for dept_name, scenarios in casino_scenarios.items():
        doc.add_heading(dept_name, level=2)
        for i, scenario in enumerate(scenarios, 1):
            p = doc.add_paragraph()
            run = p.add_run(scenario['title'])
            run.bold = True
            doc.add_paragraph(scenario['description'], style='List Bullet')
        doc.add_paragraph()

    # Footer
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_p.add_run('---  End of Document  ---')
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    footer_run.italic = True

    # Save document
    output_path = r"C:\Users\szh2051\OneDrive - Carnival Corporation\Desktop\Cruise-Department-Scenarios-Detailed.docx"
    doc.save(output_path)
    print("English scenarios document created successfully!")
    return output_path


if __name__ == "__main__":
    print("Creating comprehensive scenario documents...")
    print("16 departments × 10 scenarios = 160 total scenarios")
    print("-" * 60)

    chinese_path = create_chinese_scenarios_document()
    print("-" * 60)
    english_path = create_english_scenarios_document()
    print("-" * 60)

    print("\nBoth scenario documents created successfully!")
    print(f"Chinese: {chinese_path}")
    print(f"English: {english_path}")
