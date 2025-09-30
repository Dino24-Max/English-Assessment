"""
Update Word documents with new Hotel Operations department structure.
Removes: Entertainment, Spa & Fitness
Adds: Auxiliary Service, Laundry, Photo, Provisions
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def update_chinese_document():
    """Update Chinese version of operations document"""

    # Create new document
    doc = Document()

    # Set document styling
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(11)

    # Title
    title = doc.add_heading('邮轮员工英语评估平台', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    title_run.font.bold = True

    # Subtitle
    subtitle = doc.add_heading('运营部门与评估场景说明文档', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(18)
    subtitle_run.font.color.rgb = RGBColor(0, 102, 204)

    # Document info
    doc.add_paragraph('文档版本: 2.0')
    doc.add_paragraph('更新日期: 2025-09-30')
    doc.add_paragraph('适用对象: 邮轮酒店运营、海务运营、娱乐场运营员工')
    doc.add_paragraph()

    # Section 1: Overview
    doc.add_heading('一、平台概述', level=1)
    p = doc.add_paragraph()
    p.add_run('CCL英语评估平台').bold = True
    p.add_run('是专为邮轮员工设计的综合性英语能力评估系统。该平台通过六大核心模块，全面测试员工在实际工作场景中的英语应用能力。')

    doc.add_heading('评估模块:', level=2)
    modules = [
        ('听力理解 (Listening)', '16分', '3道题', 'AI生成的真实工作场景对话'),
        ('时间与数字 (Time & Numbers)', '16分', '3道题', '填空题，测试时间表达和数字理解'),
        ('语法应用 (Grammar)', '16分', '4道题', '多选题，评估语法准确性'),
        ('词汇匹配 (Vocabulary)', '16分', '4道题', '分类匹配题，测试行业词汇'),
        ('阅读理解 (Reading)', '16分', '4道题', '标题选择题，评估阅读能力'),
        ('口语表达 (Speaking)', '20分', '3道题', 'AI分析的语音回答')
    ]

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '模块'
    hdr_cells[1].text = '分值'
    hdr_cells[2].text = '题量'
    hdr_cells[3].text = '评估方式'

    for module in modules:
        row_cells = table.add_row().cells
        row_cells[0].text = module[0]
        row_cells[1].text = module[1]
        row_cells[2].text = module[2]
        row_cells[3].text = module[3]

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('总分: 100分').bold = True
    p.add_run(' | ')
    p.add_run('题目总数: 21题').bold = True
    p.add_run(' | ')
    p.add_run('预计时长: 15-20分钟').bold = True

    doc.add_page_break()

    # Section 2: Three Operations
    doc.add_heading('二、三大运营部门架构', level=1)

    operations = {
        '酒店运营部门 (Hotel Operations)': {
            'count': 10,
            'departments': [
                '前台接待 (Front Desk)',
                '客房服务 (Housekeeping)',
                '餐饮服务 (Food & Beverage)',
                '酒吧服务 (Bar Service)',
                '宾客服务 (Guest Services)',
                '客舱服务 (Cabin Service)',
                '辅助服务 (Auxiliary Service)',
                '洗衣房 (Laundry)',
                '摄影服务 (Photo)',
                '物资供应 (Provisions)'
            ]
        },
        '海务运营部门 (Marine Operations)': {
            'count': 3,
            'departments': [
                '甲板部 (Deck Department)',
                '机舱部 (Engine Department)',
                '安全部 (Security Department)'
            ]
        },
        '娱乐场运营部门 (Casino Operations)': {
            'count': 3,
            'departments': [
                '赌桌游戏 (Table Games)',
                '老虎机 (Slot Machines)',
                '娱乐场服务 (Casino Services)'
            ]
        }
    }

    for op_name, op_data in operations.items():
        doc.add_heading(op_name, level=2)
        p = doc.add_paragraph()
        p.add_run(f"部门数量: {op_data['count']}个").bold = True

        for i, dept in enumerate(op_data['departments'], 1):
            doc.add_paragraph(f"{i}. {dept}", style='List Number')

    doc.add_page_break()

    # Section 3: Hotel Operations Details
    doc.add_heading('三、酒店运营部门详解 (Hotel Operations)', level=1)

    hotel_depts = {
        '前台接待 (Front Desk)': {
            'responsibilities': [
                '办理客人登船/离船手续',
                '处理客舱预订和升级请求',
                '解答客人关于船上设施的问题',
                '处理投诉和特殊需求'
            ],
            'scenarios': [
                '听力模块: 客人咨询岸上观光信息对话',
                '时间与数字: 客人询问餐厅营业时间',
                '口语模块: 处理客人房卡丢失问题'
            ]
        },
        '客房服务 (Housekeeping)': {
            'responsibilities': [
                '维护客舱清洁和整洁',
                '补充客舱用品和设施',
                '响应客人的客舱服务请求',
                '确保符合卫生标准'
            ],
            'scenarios': [
                '听力模块: 客人要求额外毛巾和枕头',
                '词汇模块: 匹配清洁用品和房间物品',
                '语法模块: 描述客舱清洁流程'
            ]
        },
        '餐饮服务 (Food & Beverage)': {
            'responsibilities': [
                '提供餐厅和自助餐服务',
                '介绍菜单和特色菜品',
                '处理饮食限制和过敏需求',
                '确保用餐体验质量'
            ],
            'scenarios': [
                '听力模块: 客人询问特殊饮食菜单',
                '阅读模块: 理解菜单描述和食材',
                '口语模块: 推荐今日特色菜'
            ]
        },
        '酒吧服务 (Bar Service)': {
            'responsibilities': [
                '调制和提供饮品',
                '介绍饮料菜单和促销活动',
                '管理酒吧库存',
                '确保负责任的酒精服务'
            ],
            'scenarios': [
                '词汇模块: 匹配鸡尾酒和配料',
                '时间与数字: 告知欢乐时光时间',
                '语法模块: 描述饮品制作过程'
            ]
        },
        '宾客服务 (Guest Services)': {
            'responsibilities': [
                '协调岸上游览和活动',
                '处理客人查询和请求',
                '提供船上活动信息',
                '解决客人问题'
            ],
            'scenarios': [
                '听力模块: 客人预订岸上游览',
                '阅读模块: 理解每日活动时间表',
                '口语模块: 解释港口停靠信息'
            ]
        },
        '客舱服务 (Cabin Service)': {
            'responsibilities': [
                '提供客舱送餐服务',
                '响应客舱维护请求',
                '协调特殊服务安排',
                '确保客舱舒适度'
            ],
            'scenarios': [
                '听力模块: 客人订购客舱早餐',
                '时间与数字: 确认送餐时间',
                '语法模块: 描述客舱服务选项'
            ]
        },
        '辅助服务 (Auxiliary Service)': {
            'responsibilities': [
                '提供船上辅助支持服务',
                '协助其他部门运营',
                '处理特殊物流需求',
                '维护船上物资流转'
            ],
            'scenarios': [
                '词汇模块: 匹配服务类型和需求',
                '阅读模块: 理解服务请求单',
                '口语模块: 协调部门间服务'
            ]
        },
        '洗衣房 (Laundry)': {
            'responsibilities': [
                '处理客人和船员衣物洗涤',
                '提供干洗和熨烫服务',
                '管理洗衣时间表',
                '确保衣物质量和及时交付'
            ],
            'scenarios': [
                '听力模块: 客人询问洗衣服务价格',
                '时间与数字: 告知衣物取回时间',
                '语法模块: 解释洗涤标签说明'
            ]
        },
        '摄影服务 (Photo)': {
            'responsibilities': [
                '提供专业摄影服务',
                '组织照片拍摄活动',
                '销售和展示照片产品',
                '处理数码照片订单'
            ],
            'scenarios': [
                '词汇模块: 匹配摄影术语和产品',
                '口语模块: 介绍照片套餐选项',
                '阅读模块: 理解照片订单详情'
            ]
        },
        '物资供应 (Provisions)': {
            'responsibilities': [
                '管理船上物资库存',
                '协调物资补给和采购',
                '确保食品和用品储存',
                '维护库存管理系统'
            ],
            'scenarios': [
                '时间与数字: 记录库存数量和日期',
                '阅读模块: 理解采购订单',
                '语法模块: 描述库存管理流程'
            ]
        }
    }

    for dept_name, dept_info in hotel_depts.items():
        doc.add_heading(dept_name, level=2)

        doc.add_heading('主要职责:', level=3)
        for resp in dept_info['responsibilities']:
            doc.add_paragraph(resp, style='List Bullet')

        doc.add_heading('评估场景示例:', level=3)
        for scenario in dept_info['scenarios']:
            doc.add_paragraph(scenario, style='List Bullet')

        doc.add_paragraph()

    doc.add_page_break()

    # Section 4: Marine Operations Details
    doc.add_heading('四、海务运营部门详解 (Marine Operations)', level=1)

    marine_depts = {
        '甲板部 (Deck Department)': {
            'responsibilities': [
                '船舶导航和航行操作',
                '维护甲板设备和安全',
                '管理救生设备和演习',
                '监督靠港和离港操作'
            ],
            'scenarios': [
                '听力模块: 船长发布航行通知',
                '时间与数字: 报告船舶位置坐标',
                '口语模块: 进行安全演习指导'
            ]
        },
        '机舱部 (Engine Department)': {
            'responsibilities': [
                '维护船舶推进系统',
                '管理发电和配电系统',
                '监控机械设备运行',
                '执行预防性维护计划'
            ],
            'scenarios': [
                '词汇模块: 匹配机械部件和功能',
                '阅读模块: 理解技术手册',
                '语法模块: 描述维修程序'
            ]
        },
        '安全部 (Security Department)': {
            'responsibilities': [
                '维护船上安全和秩序',
                '监控安全系统',
                '处理安全事件',
                '执行安全协议和演习'
            ],
            'scenarios': [
                '听力模块: 收到安全警报通知',
                '口语模块: 向客人解释安全程序',
                '时间与数字: 记录事件时间和位置'
            ]
        }
    }

    for dept_name, dept_info in marine_depts.items():
        doc.add_heading(dept_name, level=2)

        doc.add_heading('主要职责:', level=3)
        for resp in dept_info['responsibilities']:
            doc.add_paragraph(resp, style='List Bullet')

        doc.add_heading('评估场景示例:', level=3)
        for scenario in dept_info['scenarios']:
            doc.add_paragraph(scenario, style='List Bullet')

        doc.add_paragraph()

    doc.add_page_break()

    # Section 5: Casino Operations Details
    doc.add_heading('五、娱乐场运营部门详解 (Casino Operations)', level=1)

    casino_depts = {
        '赌桌游戏 (Table Games)': {
            'responsibilities': [
                '操作和管理赌桌游戏',
                '解释游戏规则给客人',
                '监督游戏公平性',
                '处理筹码兑换'
            ],
            'scenarios': [
                '听力模块: 向客人解释二十一点规则',
                '时间与数字: 计算和支付赔率',
                '口语模块: 介绍不同赌桌游戏'
            ]
        },
        '老虎机 (Slot Machines)': {
            'responsibilities': [
                '维护老虎机设备',
                '协助客人使用机器',
                '处理支付和技术问题',
                '监控机器性能'
            ],
            'scenarios': [
                '词汇模块: 匹配老虎机术语',
                '阅读模块: 理解机器支付表',
                '语法模块: 解释奖金功能'
            ]
        },
        '娱乐场服务 (Casino Services)': {
            'responsibilities': [
                '提供客户服务和支持',
                '管理会员计划',
                '协调娱乐场活动',
                '处理客人查询'
            ],
            'scenarios': [
                '听力模块: 客人询问会员福利',
                '口语模块: 介绍娱乐场促销活动',
                '时间与数字: 告知锦标赛时间表'
            ]
        }
    }

    for dept_name, dept_info in casino_depts.items():
        doc.add_heading(dept_name, level=2)

        doc.add_heading('主要职责:', level=3)
        for resp in dept_info['responsibilities']:
            doc.add_paragraph(resp, style='List Bullet')

        doc.add_heading('评估场景示例:', level=3)
        for scenario in dept_info['scenarios']:
            doc.add_paragraph(scenario, style='List Bullet')

        doc.add_paragraph()

    doc.add_page_break()

    # Section 6: Question Distribution
    doc.add_heading('六、题目分布与运营部门映射', level=1)

    p = doc.add_paragraph()
    p.add_run('评估平台的21道题目均衡分布于三大运营部门，确保全面评估各部门员工的英语能力。').bold = True
    doc.add_paragraph()

    doc.add_heading('题目分布原则:', level=2)
    principles = [
        '每个模块的题目涵盖酒店、海务、娱乐场三大运营',
        '情景设计基于真实工作场景',
        '难度梯度从基础到高级',
        '综合评估听说读写能力'
    ]
    for principle in principles:
        doc.add_paragraph(principle, style='List Bullet')

    doc.add_paragraph()

    doc.add_heading('分模块题目映射:', level=2)

    module_mapping = [
        ('听力理解', '1-3', [
            'Q1: 酒店运营 - 前台接待客人咨询',
            'Q2: 海务运营 - 安全部演习通知',
            'Q3: 娱乐场运营 - 赌桌游戏规则解释'
        ]),
        ('时间与数字', '4-6', [
            'Q4: 酒店运营 - 餐饮服务营业时间',
            'Q5: 海务运营 - 甲板部位置报告',
            'Q6: 娱乐场运营 - 老虎机支付计算'
        ]),
        ('语法应用', '7-10', [
            'Q7: 酒店运营 - 客房服务流程描述',
            'Q8: 海务运营 - 机舱部维修程序',
            'Q9: 娱乐场运营 - 娱乐场服务解释',
            'Q10: 酒店运营 - 洗衣房服务说明'
        ]),
        ('词汇匹配', '11-14', [
            'Q11: 酒店运营 - 餐饮术语匹配',
            'Q12: 海务运营 - 机械部件匹配',
            'Q13: 娱乐场运营 - 赌场术语匹配',
            'Q14: 酒店运营 - 摄影服务术语'
        ]),
        ('阅读理解', '15-18', [
            'Q15: 酒店运营 - 客舱服务菜单',
            'Q16: 海务运营 - 安全程序手册',
            'Q17: 娱乐场运营 - 游戏规则说明',
            'Q18: 酒店运营 - 物资供应订单'
        ]),
        ('口语表达', '19-21', [
            'Q19: 酒店运营 - 宾客服务查询',
            'Q20: 海务运营 - 甲板部安全指导',
            'Q21: 娱乐场运营 - 会员福利介绍'
        ])
    ]

    for module_name, q_range, questions in module_mapping:
        doc.add_heading(f'{module_name} (题目 {q_range})', level=3)
        for q in questions:
            doc.add_paragraph(q, style='List Bullet')
        doc.add_paragraph()

    doc.add_page_break()

    # Section 7: Implementation Notes
    doc.add_heading('七、实施说明', level=1)

    doc.add_heading('评分标准:', level=2)
    scoring = [
        '选择题 (听力、时间数字、语法、阅读): 答对得全分，答错不得分',
        '词汇匹配题: 完全正确得全分，部分正确按比例得分',
        '口语题: AI分析发音、流利度、语法、内容相关性综合评分'
    ]
    for item in scoring:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()

    doc.add_heading('通过标准:', level=2)
    passing = [
        '60-69分: 基础通过 - 基本沟通能力',
        '70-79分: 良好 - 胜任日常工作交流',
        '80-89分: 优秀 - 熟练应对复杂场景',
        '90-100分: 卓越 - 专业级英语能力'
    ]
    for item in passing:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()

    doc.add_heading('技术特点:', level=2)
    tech_features = [
        'AI驱动的语音识别和评分',
        '实时进度追踪',
        '自适应用户界面',
        '详细的模块化成绩报告',
        '安全的会话管理',
        '移动端响应式设计'
    ]
    for feature in tech_features:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph()

    # Footer
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_p.add_run('---  文档结束  ---')
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    footer_run.italic = True

    # Save document
    output_path = r"C:\Users\szh2051\OneDrive - Carnival Corporation\Desktop\邮轮员工英语评估平台-运营部门与场景说明.docx"
    doc.save(output_path)
    print("Chinese document updated successfully!")


def update_english_document():
    """Update English version of operations document"""

    # Create new document
    doc = Document()

    # Set document styling
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title
    title = doc.add_heading('Cruise Employee English Assessment Platform', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    title_run.font.bold = True

    # Subtitle
    subtitle = doc.add_heading('Operations, Departments & Assessment Scenarios', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(18)
    subtitle_run.font.color.rgb = RGBColor(0, 102, 204)

    # Document info
    doc.add_paragraph('Document Version: 2.0')
    doc.add_paragraph('Last Updated: September 30, 2025')
    doc.add_paragraph('Target Audience: Hotel Operations, Marine Operations, Casino Operations Staff')
    doc.add_paragraph()

    # Section 1: Overview
    doc.add_heading('1. Platform Overview', level=1)
    p = doc.add_paragraph()
    p.add_run('The CCL English Assessment Platform ').bold = True
    p.add_run('is a comprehensive English proficiency evaluation system designed specifically for cruise ship employees. The platform tests practical English skills through six core modules covering real-world work scenarios.')

    doc.add_heading('Assessment Modules:', level=2)
    modules = [
        ('Listening Comprehension', '16 points', '3 questions', 'AI-generated workplace dialogues'),
        ('Time & Numbers', '16 points', '3 questions', 'Fill-in-the-blank context questions'),
        ('Grammar Application', '16 points', '4 questions', 'Multiple choice gap-fill exercises'),
        ('Vocabulary Matching', '16 points', '4 questions', 'Category matching exercises'),
        ('Reading Comprehension', '16 points', '4 questions', 'Title selection questions'),
        ('Speaking Expression', '20 points', '3 questions', 'AI-analyzed voice responses')
    ]

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Module'
    hdr_cells[1].text = 'Points'
    hdr_cells[2].text = 'Questions'
    hdr_cells[3].text = 'Assessment Type'

    for module in modules:
        row_cells = table.add_row().cells
        row_cells[0].text = module[0]
        row_cells[1].text = module[1]
        row_cells[2].text = module[2]
        row_cells[3].text = module[3]

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Total Score: 100 points').bold = True
    p.add_run(' | ')
    p.add_run('Total Questions: 21').bold = True
    p.add_run(' | ')
    p.add_run('Estimated Duration: 15-20 minutes').bold = True

    doc.add_page_break()

    # Section 2: Three Operations
    doc.add_heading('2. Three Operations Architecture', level=1)

    operations = {
        'Hotel Operations': {
            'count': 10,
            'departments': [
                'Front Desk',
                'Housekeeping',
                'Food & Beverage',
                'Bar Service',
                'Guest Services',
                'Cabin Service',
                'Auxiliary Service',
                'Laundry',
                'Photo',
                'Provisions'
            ]
        },
        'Marine Operations': {
            'count': 3,
            'departments': [
                'Deck Department',
                'Engine Department',
                'Security Department'
            ]
        },
        'Casino Operations': {
            'count': 3,
            'departments': [
                'Table Games',
                'Slot Machines',
                'Casino Services'
            ]
        }
    }

    for op_name, op_data in operations.items():
        doc.add_heading(op_name, level=2)
        p = doc.add_paragraph()
        p.add_run(f"Number of Departments: {op_data['count']}").bold = True

        for i, dept in enumerate(op_data['departments'], 1):
            doc.add_paragraph(f"{i}. {dept}", style='List Number')

    doc.add_page_break()

    # Section 3: Hotel Operations Details
    doc.add_heading('3. Hotel Operations - Detailed Breakdown', level=1)

    hotel_depts = {
        'Front Desk': {
            'responsibilities': [
                'Process guest embarkation and disembarkation',
                'Handle cabin reservations and upgrade requests',
                'Answer guest inquiries about ship facilities',
                'Resolve complaints and special requests'
            ],
            'scenarios': [
                'Listening: Guest inquiry about shore excursions',
                'Time & Numbers: Guest asking about restaurant hours',
                'Speaking: Handling lost room key issue'
            ]
        },
        'Housekeeping': {
            'responsibilities': [
                'Maintain cabin cleanliness and tidiness',
                'Replenish cabin amenities and supplies',
                'Respond to guest cabin service requests',
                'Ensure compliance with hygiene standards'
            ],
            'scenarios': [
                'Listening: Guest requesting extra towels and pillows',
                'Vocabulary: Matching cleaning supplies and room items',
                'Grammar: Describing cabin cleaning procedures'
            ]
        },
        'Food & Beverage': {
            'responsibilities': [
                'Provide restaurant and buffet service',
                'Present menus and specialty dishes',
                'Handle dietary restrictions and allergies',
                'Ensure quality dining experience'
            ],
            'scenarios': [
                'Listening: Guest asking about special diet menu',
                'Reading: Understanding menu descriptions',
                'Speaking: Recommending daily specials'
            ]
        },
        'Bar Service': {
            'responsibilities': [
                'Mix and serve beverages',
                'Present drink menus and promotions',
                'Manage bar inventory',
                'Ensure responsible alcohol service'
            ],
            'scenarios': [
                'Vocabulary: Matching cocktails and ingredients',
                'Time & Numbers: Informing happy hour times',
                'Grammar: Describing drink preparation'
            ]
        },
        'Guest Services': {
            'responsibilities': [
                'Coordinate shore excursions and activities',
                'Handle guest inquiries and requests',
                'Provide ship activity information',
                'Resolve guest issues'
            ],
            'scenarios': [
                'Listening: Guest booking shore excursion',
                'Reading: Understanding daily activity schedule',
                'Speaking: Explaining port information'
            ]
        },
        'Cabin Service': {
            'responsibilities': [
                'Provide in-cabin dining service',
                'Respond to cabin maintenance requests',
                'Coordinate special service arrangements',
                'Ensure cabin comfort'
            ],
            'scenarios': [
                'Listening: Guest ordering cabin breakfast',
                'Time & Numbers: Confirming delivery time',
                'Grammar: Describing cabin service options'
            ]
        },
        'Auxiliary Service': {
            'responsibilities': [
                'Provide auxiliary support services shipboard',
                'Assist other department operations',
                'Handle special logistics requirements',
                'Maintain shipboard supply flow'
            ],
            'scenarios': [
                'Vocabulary: Matching service types and needs',
                'Reading: Understanding service request forms',
                'Speaking: Coordinating inter-department services'
            ]
        },
        'Laundry': {
            'responsibilities': [
                'Process guest and crew laundry',
                'Provide dry cleaning and pressing services',
                'Manage laundry schedules',
                'Ensure garment quality and timely delivery'
            ],
            'scenarios': [
                'Listening: Guest inquiring about laundry pricing',
                'Time & Numbers: Informing garment pickup time',
                'Grammar: Explaining care label instructions'
            ]
        },
        'Photo': {
            'responsibilities': [
                'Provide professional photography services',
                'Organize photo shooting events',
                'Sell and display photo products',
                'Process digital photo orders'
            ],
            'scenarios': [
                'Vocabulary: Matching photography terms and products',
                'Speaking: Introducing photo package options',
                'Reading: Understanding photo order details'
            ]
        },
        'Provisions': {
            'responsibilities': [
                'Manage shipboard supply inventory',
                'Coordinate supply replenishment and procurement',
                'Ensure food and supply storage',
                'Maintain inventory management systems'
            ],
            'scenarios': [
                'Time & Numbers: Recording inventory quantities and dates',
                'Reading: Understanding procurement orders',
                'Grammar: Describing inventory management processes'
            ]
        }
    }

    for dept_name, dept_info in hotel_depts.items():
        doc.add_heading(dept_name, level=2)

        doc.add_heading('Key Responsibilities:', level=3)
        for resp in dept_info['responsibilities']:
            doc.add_paragraph(resp, style='List Bullet')

        doc.add_heading('Assessment Scenario Examples:', level=3)
        for scenario in dept_info['scenarios']:
            doc.add_paragraph(scenario, style='List Bullet')

        doc.add_paragraph()

    doc.add_page_break()

    # Section 4: Marine Operations Details
    doc.add_heading('4. Marine Operations - Detailed Breakdown', level=1)

    marine_depts = {
        'Deck Department': {
            'responsibilities': [
                'Ship navigation and sailing operations',
                'Maintain deck equipment and safety',
                'Manage lifesaving equipment and drills',
                'Supervise docking and undocking operations'
            ],
            'scenarios': [
                'Listening: Captain issuing navigation announcement',
                'Time & Numbers: Reporting ship position coordinates',
                'Speaking: Conducting safety drill instructions'
            ]
        },
        'Engine Department': {
            'responsibilities': [
                'Maintain ship propulsion systems',
                'Manage power generation and distribution',
                'Monitor mechanical equipment operation',
                'Execute preventive maintenance plans'
            ],
            'scenarios': [
                'Vocabulary: Matching mechanical parts and functions',
                'Reading: Understanding technical manuals',
                'Grammar: Describing repair procedures'
            ]
        },
        'Security Department': {
            'responsibilities': [
                'Maintain shipboard safety and order',
                'Monitor security systems',
                'Handle security incidents',
                'Execute security protocols and drills'
            ],
            'scenarios': [
                'Listening: Receiving security alert notification',
                'Speaking: Explaining safety procedures to guests',
                'Time & Numbers: Recording incident time and location'
            ]
        }
    }

    for dept_name, dept_info in marine_depts.items():
        doc.add_heading(dept_name, level=2)

        doc.add_heading('Key Responsibilities:', level=3)
        for resp in dept_info['responsibilities']:
            doc.add_paragraph(resp, style='List Bullet')

        doc.add_heading('Assessment Scenario Examples:', level=3)
        for scenario in dept_info['scenarios']:
            doc.add_paragraph(scenario, style='List Bullet')

        doc.add_paragraph()

    doc.add_page_break()

    # Section 5: Casino Operations Details
    doc.add_heading('5. Casino Operations - Detailed Breakdown', level=1)

    casino_depts = {
        'Table Games': {
            'responsibilities': [
                'Operate and manage table games',
                'Explain game rules to guests',
                'Monitor game fairness',
                'Handle chip exchanges'
            ],
            'scenarios': [
                'Listening: Explaining blackjack rules to guest',
                'Time & Numbers: Calculating and paying odds',
                'Speaking: Introducing different table games'
            ]
        },
        'Slot Machines': {
            'responsibilities': [
                'Maintain slot machine equipment',
                'Assist guests with machine usage',
                'Handle payouts and technical issues',
                'Monitor machine performance'
            ],
            'scenarios': [
                'Vocabulary: Matching slot machine terminology',
                'Reading: Understanding machine pay tables',
                'Grammar: Explaining bonus features'
            ]
        },
        'Casino Services': {
            'responsibilities': [
                'Provide customer service and support',
                'Manage loyalty programs',
                'Coordinate casino events',
                'Handle guest inquiries'
            ],
            'scenarios': [
                'Listening: Guest asking about membership benefits',
                'Speaking: Introducing casino promotions',
                'Time & Numbers: Informing tournament schedule'
            ]
        }
    }

    for dept_name, dept_info in casino_depts.items():
        doc.add_heading(dept_name, level=2)

        doc.add_heading('Key Responsibilities:', level=3)
        for resp in dept_info['responsibilities']:
            doc.add_paragraph(resp, style='List Bullet')

        doc.add_heading('Assessment Scenario Examples:', level=3)
        for scenario in dept_info['scenarios']:
            doc.add_paragraph(scenario, style='List Bullet')

        doc.add_paragraph()

    doc.add_page_break()

    # Section 6: Question Distribution
    doc.add_heading('6. Question Distribution & Operations Mapping', level=1)

    p = doc.add_paragraph()
    p.add_run('The 21 assessment questions are evenly distributed across Hotel, Marine, and Casino operations to ensure comprehensive evaluation of English skills for all departments.').bold = True
    doc.add_paragraph()

    doc.add_heading('Distribution Principles:', level=2)
    principles = [
        'Each module covers all three operations',
        'Scenarios based on real workplace situations',
        'Difficulty progression from basic to advanced',
        'Comprehensive evaluation of listening, speaking, reading, writing'
    ]
    for principle in principles:
        doc.add_paragraph(principle, style='List Bullet')

    doc.add_paragraph()

    doc.add_heading('Module-by-Module Question Mapping:', level=2)

    module_mapping = [
        ('Listening Comprehension', '1-3', [
            'Q1: Hotel Operations - Front Desk guest inquiry',
            'Q2: Marine Operations - Security drill announcement',
            'Q3: Casino Operations - Table game rules explanation'
        ]),
        ('Time & Numbers', '4-6', [
            'Q4: Hotel Operations - Food & Beverage operating hours',
            'Q5: Marine Operations - Deck position reporting',
            'Q6: Casino Operations - Slot machine payout calculation'
        ]),
        ('Grammar Application', '7-10', [
            'Q7: Hotel Operations - Housekeeping procedure description',
            'Q8: Marine Operations - Engine repair procedure',
            'Q9: Casino Operations - Casino services explanation',
            'Q10: Hotel Operations - Laundry service description'
        ]),
        ('Vocabulary Matching', '11-14', [
            'Q11: Hotel Operations - F&B terminology matching',
            'Q12: Marine Operations - Mechanical parts matching',
            'Q13: Casino Operations - Gaming terminology matching',
            'Q14: Hotel Operations - Photography service terms'
        ]),
        ('Reading Comprehension', '15-18', [
            'Q15: Hotel Operations - Cabin service menu',
            'Q16: Marine Operations - Safety procedure manual',
            'Q17: Casino Operations - Game rules explanation',
            'Q18: Hotel Operations - Provisions supply order'
        ]),
        ('Speaking Expression', '19-21', [
            'Q19: Hotel Operations - Guest services inquiry',
            'Q20: Marine Operations - Deck safety instructions',
            'Q21: Casino Operations - Membership benefits introduction'
        ])
    ]

    for module_name, q_range, questions in module_mapping:
        doc.add_heading(f'{module_name} (Questions {q_range})', level=3)
        for q in questions:
            doc.add_paragraph(q, style='List Bullet')
        doc.add_paragraph()

    doc.add_page_break()

    # Section 7: Implementation Notes
    doc.add_heading('7. Implementation Guidelines', level=1)

    doc.add_heading('Scoring Standards:', level=2)
    scoring = [
        'Multiple Choice (Listening, Time & Numbers, Grammar, Reading): Full points for correct answer, zero for incorrect',
        'Vocabulary Matching: Full points for completely correct, partial credit proportional to accuracy',
        'Speaking: AI analysis of pronunciation, fluency, grammar, and content relevance for comprehensive scoring'
    ]
    for item in scoring:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()

    doc.add_heading('Passing Standards:', level=2)
    passing = [
        '60-69 points: Basic Pass - Fundamental communication ability',
        '70-79 points: Good - Competent in daily work communication',
        '80-89 points: Excellent - Proficient in handling complex scenarios',
        '90-100 points: Outstanding - Professional-level English proficiency'
    ]
    for item in passing:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()

    doc.add_heading('Technical Features:', level=2)
    tech_features = [
        'AI-powered speech recognition and scoring',
        'Real-time progress tracking',
        'Adaptive user interface',
        'Detailed modular score reporting',
        'Secure session management',
        'Mobile-responsive design'
    ]
    for feature in tech_features:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph()

    # Footer
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_p.add_run('---  End of Document  ---')
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    footer_run.italic = True

    # Save document with temporary name to avoid permission error
    output_path = r"C:\Users\szh2051\OneDrive - Carnival Corporation\Desktop\Cruise-Employee-English-Assessment-Operations-Scenarios-Updated.docx"
    doc.save(output_path)
    print("English document updated successfully!")
    print(f"Saved as: {output_path}")
    print("Please close the old file and rename this file to replace it.")


if __name__ == "__main__":
    print("Updating Word documents with new Hotel Operations structure...")
    print("\nChanges:")
    print("  ADDED: Auxiliary Service, Laundry, Photo, Provisions")
    print("  REMOVED: Entertainment, Spa & Fitness")
    print("\nNew Hotel Operations total: 10 departments")
    print("Total departments across all operations: 16")
    print("-" * 60)

    update_chinese_document()
    print("-" * 60)
    update_english_document()
    print("-" * 60)
    print("\nBoth documents have been successfully updated!")
