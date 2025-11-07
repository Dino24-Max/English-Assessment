#!/usr/bin/env python3
"""
Generate Complete Phases Documentation
Cruise Employee English Assessment Platform - All Phases Details
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def add_heading(doc, text, level=1):
    """Add a styled heading"""
    heading = doc.add_heading(text, level=level)
    return heading

def add_status_box(doc, status_text, color='green'):
    """Add a colored status box"""
    p = doc.add_paragraph()
    run = p.add_run(status_text)
    run.bold = True
    run.font.size = Pt(12)
    if color == 'green':
        run.font.color.rgb = RGBColor(0, 128, 0)
    elif color == 'orange':
        run.font.color.rgb = RGBColor(255, 140, 0)
    elif color == 'red':
        run.font.color.rgb = RGBColor(255, 0, 0)
    return p

def add_table(doc, headers, rows):
    """Add a formatted table"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'

    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)

    return table

def create_phases_document():
    """Create complete phases documentation"""

    doc = Document()

    # ============================================
    # COVER PAGE
    # ============================================
    title = doc.add_heading('Complete Development Phases', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading('Cruise Employee English Assessment Platform', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    date_para = doc.add_paragraph(f'Generated: {datetime.now().strftime("%B %d, %Y at %I:%M %p")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    timeline_para = doc.add_paragraph('8-Week Development Plan | 6 Phases | Detailed Implementation Guide')
    timeline_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    timeline_para.runs[0].font.bold = True
    timeline_para.runs[0].font.size = Pt(12)

    doc.add_page_break()

    # ============================================
    # EXECUTIVE SUMMARY
    # ============================================
    add_heading(doc, 'Executive Summary', level=1)

    doc.add_paragraph(
        'This document provides a comprehensive breakdown of all 6 development phases for the '
        'Cruise Employee English Assessment Platform. Each phase includes detailed task lists, '
        'technical requirements, implementation steps, and success criteria.'
    )

    doc.add_heading('Current Status', level=2)
    add_status_box(doc, 'PHASE 0: COMPLETE (100%)', 'green')
    doc.add_paragraph('All critical fixes completed with comprehensive test coverage (60%) and anti-cheating system fully implemented.')

    doc.add_paragraph()
    add_status_box(doc, 'PHASES 1-6: PENDING', 'orange')
    doc.add_paragraph('Ready to begin Phase 1: Admin Dashboard development.')

    doc.add_heading('Timeline Overview', level=2)

    timeline_data = [
        ['Phase 0', 'Critical Fixes', 'Week 1', '5 days', 'COMPLETE'],
        ['Phase 1', 'Admin Dashboard', 'Week 2-3', '10 days', 'PENDING'],
        ['Phase 2', 'Analytics & Reporting', 'Week 4', '5 days', 'PENDING'],
        ['Phase 3', 'Email Notifications', 'Week 5', '5 days', 'PENDING'],
        ['Phase 4', 'Export & Certificates', 'Week 6', '5 days', 'PENDING'],
        ['Phase 5', 'Testing & QA', 'Week 7', '5 days', 'PENDING'],
        ['Phase 6', 'Documentation', 'Week 8', '5 days', 'PENDING']
    ]
    add_table(doc, ['Phase', 'Name', 'Timeline', 'Duration', 'Status'], timeline_data)

    doc.add_page_break()

    # ============================================
    # PHASE 0: CRITICAL FIXES (COMPLETE)
    # ============================================
    add_heading(doc, 'Phase 0: Critical Fixes & Foundation', level=1)
    add_status_box(doc, 'STATUS: COMPLETE', 'green')

    doc.add_paragraph()
    doc.add_paragraph('Duration: 5 days (Week 1)')
    doc.add_paragraph('Priority: CRITICAL')
    doc.add_paragraph('Completion Date: November 6, 2025')

    doc.add_heading('Overview', level=2)
    doc.add_paragraph(
        'Phase 0 addressed all production-blocking issues with 100% internal solutions (no paid services). '
        'Critical database persistence, missing templates, audio recording, and anti-cheating validation '
        'were successfully implemented and tested.'
    )

    doc.add_heading('Completed Tasks', level=2)

    # Task 1
    doc.add_heading('Task 1: Database Answer Persistence', level=3)
    doc.add_paragraph('File Modified: src/main/python/api/routes/ui.py')
    doc.add_paragraph()
    doc.add_paragraph('Implementations:')
    implementations = [
        'Automatic guest user creation for demo/testing',
        'Assessment session auto-creation',
        'Answer persistence via AssessmentEngine.submit_response()',
        'Atomic database transactions with error handling',
        'Results page now fetches from database (not session)',
        'Real-time score calculation and storage'
    ]
    for impl in implementations:
        doc.add_paragraph(impl, style='List Bullet')

    # Task 2
    doc.add_heading('Task 2: Missing HTML Templates', level=3)
    doc.add_paragraph('Files Created: 3 new templates')
    doc.add_paragraph()

    doc.add_paragraph('2.1 Registration Template (registration.html)')
    features_reg = [
        'Professional design with gradient background',
        'Division selection: Hotel, Marine, Casino',
        'Dynamic department dropdown (14 total departments)',
        'Client-side and server-side validation',
        'Responsive design (mobile, tablet, desktop)',
        'Accessibility compliant (WCAG 2.1 AA)'
    ]
    for feature in features_reg:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('2.2 Instructions Template (instructions.html)')
    features_inst = [
        'Comprehensive assessment guidelines',
        '6 module descriptions with visual icons',
        'Pass/fail criteria (70 points, 80% safety, 12 speaking)',
        'Anti-cheating notice and rules',
        'Technical requirements list',
        'Responsive card-based layout'
    ]
    for feature in features_inst:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('2.3 Speaking Question Template (speaking_question.html)')
    features_speak = [
        'Browser-native MediaRecorder API (no external services)',
        'Real-time recording timer with countdown',
        'Waveform visualization during recording',
        'Playback functionality for review',
        'Re-record capability',
        'Microphone permission handling',
        '20-second maximum recording limit',
        'Auto-stop on time limit',
        'Form submission with audio blob'
    ]
    for feature in features_speak:
        doc.add_paragraph(feature, style='List Bullet')

    # Task 3
    doc.add_heading('Task 3: Anti-Cheating Validation', level=3)
    doc.add_paragraph('File Modified: src/main/python/utils/anti_cheating.py (complete rewrite)')
    doc.add_paragraph()

    doc.add_paragraph('3.1 Session Tracking')
    tracking = [
        'IP address capture from X-Forwarded-For, X-Real-IP, or direct client',
        'User agent (browser/device) tracking',
        'Session start metadata recording',
        'Proxy chain support for accurate IP detection'
    ]
    for item in tracking:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('3.2 Integrity Validation')
    validation = [
        'IP consistency checking (detects network changes)',
        'User agent consistency checking (detects device switches)',
        'Real-time session validation on each request',
        'Suspicious behavior detection and flagging'
    ]
    for item in validation:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('3.3 Behavior Monitoring')
    monitoring = [
        'Tab Switching: Frontend visibility API, counter with warnings, threshold: 3 switches',
        'Copy/Paste: Clipboard event monitoring, attempt counter, threshold: 5 attempts',
        'Automatic event logging to analytics_data JSON field',
        'Suspicious event timeline with timestamps'
    ]
    for item in monitoring:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('3.4 Suspicious Score Algorithm (0-100 scale)')
    doc.add_paragraph('Scoring breakdown:')
    scoring_table = [
        ['IP address changed', '+40 points', 'Critical indicator'],
        ['User agent changed', '+30 points', 'High concern'],
        ['Tab switches', '+5 each', 'Max +20 points'],
        ['Copy/paste attempts', '+3 each', 'Max +15 points'],
        ['Suspicious events', '+10 each', 'Max +20 points']
    ]
    add_table(doc, ['Risk Factor', 'Points', 'Severity'], scoring_table)

    doc.add_paragraph()
    doc.add_paragraph('Risk Levels:')
    risk_levels = [
        '0 points: Clean (no issues)',
        '1-19 points: Low concern (minor behavior)',
        '20-39 points: Medium concern (watch closely)',
        '40-69 points: High concern (requires review)',
        '70-100 points: Critical (likely cheating)'
    ]
    for level in risk_levels:
        doc.add_paragraph(level, style='List Bullet')

    # Task 4
    doc.add_heading('Task 4: Anti-Cheating Admin Endpoints', level=3)
    doc.add_paragraph('File Modified: src/main/python/api/routes/admin.py')
    doc.add_paragraph()
    doc.add_paragraph('New API Endpoints:')
    endpoints = [
        'GET /api/v1/admin/anti-cheating/assessments - View all assessments with anti-cheating data',
        'GET /api/v1/admin/anti-cheating/assessment/{id} - Detailed risk breakdown for specific assessment',
        'Returns: IP addresses, user agents, tab switches, copy/paste counts, suspicious scores, risk levels'
    ]
    for ep in endpoints:
        doc.add_paragraph(ep, style='List Bullet')

    # Task 5
    doc.add_heading('Task 5: Comprehensive Test Suite', level=3)
    doc.add_paragraph('Files Created: 5 test files + documentation')
    doc.add_paragraph()

    doc.add_paragraph('5.1 Unit Tests (test_anti_cheating.py)')
    unit_tests = [
        '25+ test cases covering anti-cheating service',
        '~95% coverage of anti_cheating.py',
        'Test classes: IP tracking, session validation, tab switches, copy/paste, scoring, flagging',
        'Mock database and request objects for isolated testing'
    ]
    for test in unit_tests:
        doc.add_paragraph(test, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('5.2 Integration Tests (test_ui_routes.py)')
    int_tests = [
        '20+ test cases for UI routes and database integration',
        '~80% coverage of ui.py routes',
        'Test classes: Answer submission, results retrieval, end-to-end flows, error handling',
        'Uses in-memory SQLite for fast, isolated tests'
    ]
    for test in int_tests:
        doc.add_paragraph(test, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('5.3 Test Infrastructure')
    infra = [
        'run_tests.py - Automated test runner with coverage reporting',
        'TESTING.md - Comprehensive testing documentation',
        'TEST_SUITE_SUMMARY.md - Quick reference guide',
        'pytest configuration with async support',
        'Coverage target: 50%+ (Actual: ~60%)'
    ]
    for item in infra:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Phase 0 Success Metrics', level=2)

    metrics_table = [
        ['Database Persistence', 'Fixed', 'COMPLETE', 'SUCCESS'],
        ['Missing Templates', '2-3 templates', '3 templates created', 'SUCCESS'],
        ['Audio Recording', 'Basic interface', 'Full-featured UI', 'EXCEEDED'],
        ['Anti-Cheating', 'Basic validation', 'Comprehensive system', 'EXCEEDED'],
        ['Test Coverage', '50%+', '~60%', 'EXCEEDED'],
        ['Paid Services', '$0', '$0', 'SUCCESS']
    ]
    add_table(doc, ['Goal', 'Target', 'Actual', 'Status'], metrics_table)

    doc.add_heading('Phase 0 Deliverables', level=2)
    deliverables = [
        '2 modified Python files (ui.py, anti_cheating.py)',
        '3 new HTML templates (registration, instructions, speaking)',
        '2 new test files (unit + integration)',
        '3 documentation files (PHASE_0_COMPLETION.md, TESTING.md, TEST_SUITE_SUMMARY.md)',
        '1 test runner script (run_tests.py)',
        'All committed to GitHub with detailed commit messages'
    ]
    for deliverable in deliverables:
        doc.add_paragraph(deliverable, style='List Bullet')

    doc.add_page_break()

    # ============================================
    # PHASE 1: ADMIN DASHBOARD
    # ============================================
    add_heading(doc, 'Phase 1: Admin Dashboard', level=1)
    add_status_box(doc, 'STATUS: PENDING', 'orange')

    doc.add_paragraph()
    doc.add_paragraph('Duration: 10 days (Week 2-3)')
    doc.add_paragraph('Priority: HIGH')
    doc.add_paragraph('Dependencies: Phase 0 complete')

    doc.add_heading('Overview', level=2)
    doc.add_paragraph(
        'Phase 1 builds a comprehensive admin dashboard for managing users, assessments, questions, '
        'and monitoring anti-cheating activities. This provides administrators with full control '
        'over the platform and visibility into all operations.'
    )

    doc.add_heading('Feature 1: User Management System', level=2)
    doc.add_paragraph('Estimated Time: 3 days')
    doc.add_paragraph()
    doc.add_paragraph('Functionality:')
    user_mgmt = [
        'View all users with pagination and search',
        'Filter by division, department, registration date',
        'Edit user details (name, email, division, department)',
        'Deactivate/reactivate user accounts',
        'View user assessment history',
        'Export user data (CSV/Excel)',
        'User activity tracking (last login, assessments taken)',
        'Bulk user import from CSV'
    ]
    for func in user_mgmt:
        doc.add_paragraph(func, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Technical Implementation:')
    tech_user = [
        'API Endpoints: GET/POST/PUT/DELETE /api/v1/admin/users',
        'Database: User model queries with SQLAlchemy',
        'Frontend: HTML templates with DataTables.js for sorting/filtering',
        'Authentication: Admin role checking middleware',
        'Validation: Pydantic schemas for input validation'
    ]
    for tech in tech_user:
        doc.add_paragraph(tech, style='List Bullet')

    doc.add_heading('Feature 2: Assessment Monitoring', level=2)
    doc.add_paragraph('Estimated Time: 3 days')
    doc.add_paragraph()
    doc.add_paragraph('Functionality:')
    assess_mon = [
        'View all assessments (in-progress, completed, flagged)',
        'Real-time assessment status updates',
        'Filter by division, status, date range, score range',
        'Search by user name or email',
        'View detailed assessment responses',
        'Monitor time spent per question',
        'Identify struggling candidates (low scores, long times)',
        'Manual assessment review and override'
    ]
    for func in assess_mon:
        doc.add_paragraph(func, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Technical Implementation:')
    tech_assess = [
        'API Endpoints: GET /api/v1/admin/assessments with query parameters',
        'WebSocket support for real-time updates',
        'Database: Complex joins (Assessment + User + Responses)',
        'Caching: Redis for frequently accessed data',
        'Frontend: Dynamic table with live updates'
    ]
    for tech in tech_assess:
        doc.add_paragraph(tech, style='List Bullet')

    doc.add_heading('Feature 3: Question Bank Management', level=2)
    doc.add_paragraph('Estimated Time: 2 days')
    doc.add_paragraph()
    doc.add_paragraph('Functionality:')
    question_mgmt = [
        'View all questions by module (Listening, Grammar, etc.)',
        'Create new questions with rich text editor',
        'Edit existing questions',
        'Delete questions (with usage check)',
        'Bulk import questions from CSV/JSON',
        'Bulk export questions',
        'Question difficulty tagging (easy, medium, hard)',
        'Question usage statistics (how often used, pass rate)',
        'Question preview functionality'
    ]
    for func in question_mgmt:
        doc.add_paragraph(func, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Technical Implementation:')
    tech_question = [
        'API Endpoints: Full CRUD for /api/v1/admin/questions',
        'Database: Question model with metadata fields',
        'Frontend: Form builder with TinyMCE or Quill for rich text',
        'File upload: Audio files for listening questions',
        'Validation: Answer format validation by question type'
    ]
    for tech in tech_question:
        doc.add_paragraph(tech, style='List Bullet')

    doc.add_heading('Feature 4: Anti-Cheating Dashboard', level=2)
    doc.add_paragraph('Estimated Time: 2 days')
    doc.add_paragraph()
    doc.add_paragraph('Functionality:')
    anti_cheat_dash = [
        'View all flagged assessments (score >= 40)',
        'Suspicious score distribution chart',
        'Risk factor breakdown (IP changes, tab switches, etc.)',
        'Manual review workflow (approve, reject, investigate)',
        'IP address tracking and geolocation',
        'User agent analysis (detect device/browser patterns)',
        'Timeline view of suspicious events',
        'Auto-flag configuration (threshold adjustment)',
        'Export anti-cheating reports'
    ]
    for func in anti_cheat_dash:
        doc.add_paragraph(func, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Technical Implementation:')
    tech_anti = [
        'API Endpoints: Already created in Phase 0 (extend as needed)',
        'Frontend: Chart.js for score distribution visualization',
        'Database: Query analytics_data JSON field',
        'IP Geolocation: Free API (ip-api.com) or GeoLite2',
        'Review workflow: Add review_status field to Assessment model'
    ]
    for tech in tech_anti:
        doc.add_paragraph(tech, style='List Bullet')

    doc.add_heading('Phase 1 Technical Stack', level=2)
    tech_stack_p1 = [
        'Backend: FastAPI, SQLAlchemy, Pydantic',
        'Database: PostgreSQL (or SQLite for dev)',
        'Frontend: HTML/CSS/JavaScript, Bootstrap 5',
        'Tables: DataTables.js (sorting, filtering, pagination)',
        'Charts: Chart.js (for visualizations)',
        'Real-time: WebSocket for live updates (optional)',
        'Authentication: JWT tokens or session-based',
        'File uploads: FastAPI UploadFile handling'
    ]
    for tech in tech_stack_p1:
        doc.add_paragraph(tech, style='List Bullet')

    doc.add_heading('Phase 1 Deliverables', level=2)
    deliverables_p1 = [
        'Admin dashboard landing page (/admin)',
        'User management interface with CRUD operations',
        'Assessment monitoring dashboard with filters',
        'Question bank editor with rich text support',
        'Anti-cheating monitoring dashboard',
        'API endpoints for all admin operations',
        'Admin authentication and authorization',
        'Unit and integration tests for admin features',
        'Admin user documentation'
    ]
    for deliverable in deliverables_p1:
        doc.add_paragraph(deliverable, style='List Bullet')

    doc.add_heading('Phase 1 Success Criteria', level=2)
    success_p1 = [
        'Admins can create, edit, and delete users',
        'All assessments viewable with real-time status',
        'Question bank fully manageable with bulk operations',
        'Flagged assessments clearly identified and reviewable',
        'All admin pages load in < 2 seconds',
        'Mobile-responsive admin interface',
        'Test coverage: 70%+ for admin features',
        'Zero critical bugs in admin workflows'
    ]
    for criteria in success_p1:
        doc.add_paragraph(criteria, style='List Bullet')

    doc.add_page_break()

    # ============================================
    # PHASE 2: ANALYTICS & REPORTING
    # ============================================
    add_heading(doc, 'Phase 2: Analytics & Reporting', level=1)
    add_status_box(doc, 'STATUS: PENDING', 'orange')

    doc.add_paragraph()
    doc.add_paragraph('Duration: 5 days (Week 4)')
    doc.add_paragraph('Priority: HIGH')
    doc.add_paragraph('Dependencies: Phase 1 complete')

    doc.add_heading('Overview', level=2)
    doc.add_paragraph(
        'Phase 2 implements a comprehensive analytics engine with interactive dashboards, '
        'performance metrics, and predictive insights. This provides data-driven visibility '
        'into assessment performance and trends.'
    )

    doc.add_heading('Feature 1: Performance Analytics Engine', level=2)
    doc.add_paragraph('Estimated Time: 2 days')
    doc.add_paragraph()
    doc.add_paragraph('Metrics Tracked:')
    metrics = [
        'Pass/Fail Rates: Overall, by division, by department, by module',
        'Average Scores: Total score, per module, trending over time',
        'Time Metrics: Average completion time, time per module, time per question',
        'Module Difficulty: Pass rates per module, average scores per module',
        'Question Analytics: Most missed questions, easiest/hardest questions',
        'User Performance: Score distribution, improvement over retakes',
        'Anti-Cheating Stats: Flagged assessment rate, average suspicious scores',
        'Completion Rates: Started vs completed, abandonment points'
    ]
    for metric in metrics:
        doc.add_paragraph(metric, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Technical Implementation:')
    tech_analytics = [
        'Database: Aggregation queries with SQLAlchemy',
        'Caching: Redis for computed metrics (refresh every 15 minutes)',
        'API Endpoints: GET /api/v1/analytics/* with query parameters',
        'Background Jobs: Celery for periodic metric calculation',
        'Data Warehouse: Optional denormalized tables for fast queries'
    ]
    for tech in tech_analytics:
        doc.add_paragraph(tech, style='List Bullet')

    doc.add_heading('Feature 2: Interactive Dashboards', level=2)
    doc.add_paragraph('Estimated Time: 2 days')
    doc.add_paragraph()
    doc.add_paragraph('Dashboard Views:')
    dashboards = [
        'Overview Dashboard: KPIs, pass rate, average score, total assessments',
        'Division Performance: Comparison charts for Hotel, Marine, Casino',
        'Module Analysis: Score breakdown by 6 modules with bar charts',
        'Trend Analysis: Line charts showing performance over time',
        'User Leaderboard: Top performers, most improved',
        'Question Heatmap: Visual difficulty map of all questions',
        'Real-Time Stats: Live counter of active assessments',
        'Custom Reports: User-defined filters and date ranges'
    ]
    for dash in dashboards:
        doc.add_paragraph(dash, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Visualization Tools:')
    viz_tools = [
        'Chart.js: Bar charts, line charts, pie charts',
        'Plotly (optional): Interactive 3D charts, heatmaps',
        'DataTables: Sortable, filterable data grids',
        'Export Options: PDF, PNG, CSV, Excel',
        'Responsive Design: Mobile-friendly charts'
    ]
    for tool in viz_tools:
        doc.add_paragraph(tool, style='List Bullet')

    doc.add_heading('Feature 3: Predictive Analytics (Optional)', level=2)
    doc.add_paragraph('Estimated Time: 1 day')
    doc.add_paragraph()
    doc.add_paragraph('AI-Powered Insights:')
    predictive = [
        'Score Prediction: Predict final score after first 3 modules',
        'Failure Risk: Identify candidates likely to fail (early intervention)',
        'Improvement Recommendations: Suggest modules for focused study',
        'Trend Forecasting: Predict future pass rates based on historical data',
        'Anomaly Detection: Flag unusual scoring patterns',
        'Question Difficulty Adjustment: Recommend question pool changes'
    ]
    for pred in predictive:
        doc.add_paragraph(pred, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Technical Implementation:')
    tech_pred = [
        'Machine Learning: Scikit-learn for regression models',
        'Training Data: Historical assessment results',
        'Model Storage: Pickle or joblib for model persistence',
        'API Endpoint: POST /api/v1/analytics/predict',
        'Refresh: Retrain models weekly with new data'
    ]
    for tech in tech_pred:
        doc.add_paragraph(tech, style='List Bullet')

    doc.add_heading('Phase 2 Technical Stack', level=2)
    tech_stack_p2 = [
        'Analytics Engine: SQLAlchemy aggregations, Pandas (optional)',
        'Caching: Redis for computed metrics',
        'Background Jobs: Celery + Redis for async metric calculation',
        'Visualization: Chart.js, Plotly (optional)',
        'ML (Optional): Scikit-learn, NumPy, Pandas',
        'Export: ReportLab (PDF), openpyxl (Excel)'
    ]
    for tech in tech_stack_p2:
        doc.add_paragraph(tech, style='List Bullet')

    doc.add_heading('Phase 2 Deliverables', level=2)
    deliverables_p2 = [
        'Analytics API endpoints (/api/v1/analytics/*)',
        'Overview dashboard with KPIs',
        'Division and module performance dashboards',
        'Interactive charts with drill-down capability',
        'Custom report builder with filters',
        'Export functionality (CSV, Excel, PDF)',
        'Background job for metric calculation',
        'Predictive analytics endpoint (optional)',
        'Analytics documentation for admins'
    ]
    for deliverable in deliverables_p2:
        doc.add_paragraph(deliverable, style='List Bullet')

    doc.add_heading('Phase 2 Success Criteria', level=2)
    success_p2 = [
        'All key metrics calculated accurately',
        'Dashboards load in < 3 seconds',
        'Charts render correctly on all devices',
        'Export generates valid files',
        'Metrics refresh automatically every 15 minutes',
        'Test coverage: 65%+ including analytics',
        'Predictive models achieve 80%+ accuracy (if implemented)',
        'Admin user feedback: "Very useful" or higher'
    ]
    for criteria in success_p2:
        doc.add_paragraph(criteria, style='List Bullet')

    doc.add_page_break()

    # ============================================
    # PHASE 3: EMAIL NOTIFICATIONS
    # ============================================
    add_heading(doc, 'Phase 3: Email Notifications', level=1)
    add_status_box(doc, 'STATUS: PENDING', 'orange')

    doc.add_paragraph()
    doc.add_paragraph('Duration: 5 days (Week 5)')
    doc.add_paragraph('Priority: MEDIUM-HIGH')
    doc.add_paragraph('Dependencies: Phase 0-2 complete')

    doc.add_heading('Overview', level=2)
    doc.add_paragraph(
        'Phase 3 implements a comprehensive email notification system with automated triggers, '
        'professional HTML templates, and multi-language support. This keeps users and admins '
        'informed about assessment status and important events.'
    )

    doc.add_heading('Feature 1: User Notifications', level=2)
    doc.add_paragraph('Estimated Time: 2 days')
    doc.add_paragraph()
    doc.add_paragraph('Email Types:')
    user_emails = [
        'Registration Confirmation: Welcome email with assessment link',
        'Assessment Started: Confirmation email when assessment begins',
        'Assessment Completed: Immediate notification upon submission',
        'Results Available: Email with score summary and pass/fail status',
        'Certificate Delivery: PDF certificate attached (Phase 4 integration)',
        'Assessment Reminder: Reminder if not started within 24 hours',
        'Incomplete Assessment: Reminder if started but not completed',
        'Retake Invitation: If failed, invitation to retake with study tips'
    ]
    for email in user_emails:
        doc.add_paragraph(email, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Email Content:')
    email_content = [
        'Personalized greeting (user name)',
        'Assessment details (division, date, time)',
        'Score breakdown by module',
        'Pass/fail status with clear messaging',
        'Next steps and action items',
        'Support contact information',
        'Branded footer with company logo'
    ]
    for content in email_content:
        doc.add_paragraph(content, style='List Bullet')

    doc.add_heading('Feature 2: Admin Notifications', level=2)
    doc.add_paragraph('Estimated Time: 1 day')
    doc.add_paragraph()
    doc.add_paragraph('Email Types:')
    admin_emails = [
        'Flagged Assessment Alert: Immediate notification when suspicious score >= 40',
        'Daily Summary Report: Total assessments, pass rate, flagged count',
        'Weekly Analytics Report: Performance trends, insights',
        'System Health Alert: Error notifications, downtime alerts',
        'New User Registration: Notification of new registrations',
        'Failed Assessment Threshold: Alert if failure rate exceeds 50%',
        'Certificate Request: Manual certificate issuance requests'
    ]
    for email in admin_emails:
        doc.add_paragraph(email, style='List Bullet')

    doc.add_heading('Feature 3: Email Templates', level=2)
    doc.add_paragraph('Estimated Time: 1 day')
    doc.add_paragraph()
    doc.add_paragraph('Template Features:')
    template_features = [
        'Professional HTML design with inline CSS',
        'Responsive layout (mobile-friendly)',
        'Company branding (logo, colors)',
        'Personalization tokens ({user_name}, {score}, {date})',
        'Multi-language support (English, Spanish, more)',
        'Dynamic content based on assessment outcome',
        'CTA buttons (View Results, Retake Assessment)',
        'Plain text fallback for compatibility'
    ]
    for feature in template_features:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Template System:')
    template_system = [
        'Jinja2 templates for dynamic content',
        'Template inheritance for consistent styling',
        'Template versioning for A/B testing',
        'Preview functionality in admin dashboard',
        'Admin-editable templates (optional)'
    ]
    for sys in template_system:
        doc.add_paragraph(sys, style='List Bullet')

    doc.add_heading('Feature 4: Email Service Integration', level=2)
    doc.add_paragraph('Estimated Time: 1 day')
    doc.add_paragraph()
    doc.add_paragraph('Email Service Options:')

    doc.add_paragraph()
    doc.add_paragraph('Option 1: SMTP (Free)')
    smtp_features = [
        'Cost: Free (use company SMTP server or Gmail)',
        'Pros: No external service, full control',
        'Cons: Limited to ~500 emails/day (Gmail), deliverability issues',
        'Best for: Development, small-scale deployments',
        'Implementation: Python smtplib or aiosmtplib'
    ]
    for feature in smtp_features:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Option 2: SendGrid (Paid - Recommended for Production)')
    sendgrid_features = [
        'Cost: Free tier (100 emails/day), $15/month (40k emails)',
        'Pros: High deliverability, analytics, templates, API',
        'Cons: Paid service (beyond free tier)',
        'Best for: Production deployments',
        'Implementation: sendgrid-python library'
    ]
    for feature in sendgrid_features:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Option 3: Amazon SES (Paid - Cost-Effective)')
    ses_features = [
        'Cost: $0.10 per 1,000 emails',
        'Pros: Very cheap at scale, AWS integration',
        'Cons: Requires AWS account setup',
        'Best for: High-volume production',
        'Implementation: boto3 library'
    ]
    for feature in ses_features:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_heading('Feature 5: Email Queue and Retry Logic', level=2)
    doc.add_paragraph('Estimated Time: 0.5 days')
    doc.add_paragraph()
    doc.add_paragraph('Queue Implementation:')
    queue_features = [
        'Celery + Redis for background email sending',
        'Retry logic for failed sends (3 attempts with backoff)',
        'Dead letter queue for permanently failed emails',
        'Email status tracking (sent, failed, opened)',
        'Rate limiting to avoid spam detection',
        'Priority queue for urgent emails'
    ]
    for feature in queue_features:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_heading('Phase 3 Technical Stack', level=2)
    tech_stack_p3 = [
        'Email Sending: aiosmtplib (SMTP) or sendgrid-python',
        'Templates: Jinja2 for HTML generation',
        'Queue: Celery + Redis for background jobs',
        'Tracking: Database table for email logs',
        'Testing: Mailtrap or MailHog for dev testing',
        'Validation: email-validator library'
    ]
    for tech in tech_stack_p3:
        doc.add_paragraph(tech, style='List Bullet')

    doc.add_heading('Phase 3 Deliverables', level=2)
    deliverables_p3 = [
        'Email service integration (SMTP or SendGrid)',
        '8 user email templates (HTML + plain text)',
        '7 admin email templates',
        'Email queue with Celery + Redis',
        'Email status tracking dashboard',
        'Template preview in admin panel',
        'Email sending API endpoint',
        'Retry logic for failed sends',
        'Email notification configuration settings',
        'Unit tests for email functionality'
    ]
    for deliverable in deliverables_p3:
        doc.add_paragraph(deliverable, style='List Bullet')

    doc.add_heading('Phase 3 Success Criteria', level=2)
    success_p3 = [
        'Emails sent within 30 seconds of trigger event',
        'Email deliverability rate: 95%+',
        'Templates render correctly in Gmail, Outlook, Apple Mail',
        'Failed emails retry automatically',
        'Admin can preview and test templates',
        'Email logs viewable in admin dashboard',
        'Test coverage: 70%+ including email features',
        'No user complaints about missing emails'
    ]
    for criteria in success_p3:
        doc.add_paragraph(criteria, style='List Bullet')

    doc.add_page_break()

    # ============================================
    # PHASE 4: EXPORT & CERTIFICATES
    # ============================================
    add_heading(doc, 'Phase 4: Export & Certificate Generation', level=1)
    add_status_box(doc, 'STATUS: PENDING', 'orange')

    doc.add_paragraph()
    doc.add_paragraph('Duration: 5 days (Week 6)')
    doc.add_paragraph('Priority: MEDIUM')
    doc.add_paragraph('Dependencies: Phase 0-3 complete')

    doc.add_heading('Overview', level=2)
    doc.add_paragraph(
        'Phase 4 implements PDF certificate generation, detailed assessment reports, and bulk '
        'export functionality. This provides users with professional certificates and admins '
        'with powerful data export capabilities.'
    )

    doc.add_heading('Feature 1: PDF Certificate Generation', level=2)
    doc.add_paragraph('Estimated Time: 2 days')
    doc.add_paragraph()
    doc.add_paragraph('Certificate Features:')
    cert_features = [
        'Professional branded design (company logo, colors)',
        'User information (name, division, department)',
        'Assessment details (date, total score, pass/fail)',
        'Module score breakdown (all 6 modules)',
        'QR code for verification (links to verification page)',
        'Digital signature or stamp',
        'Certificate ID (unique identifier)',
        'Expiration date (optional, e.g., valid for 1 year)',
        'Watermark for authenticity'
    ]
    for feature in cert_features:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Certificate Templates:')
    cert_templates = [
        'Pass Certificate: Green theme, congratulatory message',
        'Fail Certificate: Optional "Participation Certificate"',
        'Excellence Certificate: For scores >= 90',
        'Multi-language support (English, Spanish)',
        'Customizable per division (Hotel, Marine, Casino)'
    ]
    for template in cert_templates:
        doc.add_paragraph(template, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Technical Implementation:')
    tech_cert = [
        'PDF Generation: ReportLab or WeasyPrint (HTML to PDF)',
        'QR Codes: python-qrcode library',
        'Storage: Save to data/certificates/{user_id}_{assessment_id}.pdf',
        'API Endpoint: GET /api/v1/certificate/{assessment_id}',
        'Email Integration: Attach to Phase 3 completion email',
        'Verification Page: Public page to verify certificate by ID or QR'
    ]
    for tech in tech_cert:
        doc.add_paragraph(tech, style='List Bullet')

    doc.add_heading('Feature 2: Assessment Detail Reports', level=2)
    doc.add_paragraph('Estimated Time: 1.5 days')
    doc.add_paragraph()
    doc.add_paragraph('Report Content:')
    report_content = [
        'User Information: Name, email, division, department',
        'Assessment Overview: Date, duration, status',
        'Total Score: Overall score and pass/fail determination',
        'Module Breakdown: Score for each of 6 modules',
        'Question-by-Question: All questions, user answers, correct answers, points',
        'Time Analysis: Time spent per module and per question',
        'Anti-Cheating Report: IP address, suspicious score, risk factors',
        'Strengths and Weaknesses: Analysis of performance',
        'Recommendations: Areas for improvement'
    ]
    for content in report_content:
        doc.add_paragraph(content, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Report Formats:')
    report_formats = [
        'PDF: Professional formatted report for printing/sharing',
        'HTML: Web-viewable version',
        'JSON: Machine-readable format for integrations',
        'CSV: Data export for analysis (question-level detail)'
    ]
    for fmt in report_formats:
        doc.add_paragraph(fmt, style='List Bullet')

    doc.add_heading('Feature 3: Bulk Export Functionality', level=2)
    doc.add_paragraph('Estimated Time: 1 day')
    doc.add_paragraph()
    doc.add_paragraph('Export Options:')
    export_options = [
        'User Data Export: All users with registration details',
        'Assessment Results Export: All assessments with scores',
        'Question Performance Export: Aggregated question statistics',
        'Anti-Cheating Report: All flagged assessments',
        'Custom Filtered Export: User-defined filters (date, division, score)',
        'Response Detail Export: All individual responses (question-level)',
        'Audio Files Export: Bulk download of speaking responses'
    ]
    for option in export_options:
        doc.add_paragraph(option, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Export Formats:')
    export_formats = [
        'CSV: Excel-compatible, easy data analysis',
        'Excel (XLSX): Multiple sheets, formatted output',
        'JSON: API integration, data interchange',
        'ZIP: Bundled files (e.g., certificates, audio)'
    ]
    for fmt in export_formats:
        doc.add_paragraph(fmt, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Technical Implementation:')
    tech_export = [
        'CSV: Python csv module or pandas',
        'Excel: openpyxl library',
        'PDF: ReportLab for bulk generation',
        'ZIP: Python zipfile module',
        'Background Jobs: Celery for large exports (async)',
        'Download Links: Temporary signed URLs (expire after 1 hour)',
        'API Endpoints: POST /api/v1/export/{type} with filters'
    ]
    for tech in tech_export:
        doc.add_paragraph(tech, style='List Bullet')

    doc.add_heading('Feature 4: Certificate Verification System', level=2)
    doc.add_paragraph('Estimated Time: 0.5 days')
    doc.add_paragraph()
    doc.add_paragraph('Verification Features:')
    verify_features = [
        'Public Verification Page: /verify/{certificate_id}',
        'QR Code Scan: Mobile-friendly QR verification',
        'Certificate Lookup: Search by ID or user details',
        'Display Verified Details: Name, score, date, validity',
        'Tamper Detection: Flag invalid or expired certificates',
        'API Endpoint: GET /api/v1/verify/{certificate_id} (public)',
        'Verification Log: Track verification attempts'
    ]
    for feature in verify_features:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_heading('Phase 4 Technical Stack', level=2)
    tech_stack_p4 = [
        'PDF Generation: ReportLab or WeasyPrint',
        'QR Codes: python-qrcode',
        'Excel: openpyxl',
        'CSV: pandas (optional for complex exports)',
        'ZIP: Python zipfile',
        'Background Jobs: Celery + Redis for large exports',
        'Storage: Local filesystem or S3-compatible storage'
    ]
    for tech in tech_stack_p4:
        doc.add_paragraph(tech, style='List Bullet')

    doc.add_heading('Phase 4 Deliverables', level=2)
    deliverables_p4 = [
        'PDF certificate generation with QR codes',
        'Certificate templates (3+ designs)',
        'Detailed assessment report (PDF + HTML)',
        'Bulk export functionality (CSV, Excel, JSON, ZIP)',
        'Certificate verification page',
        'QR code verification system',
        'Admin export dashboard with filters',
        'Email integration (attach certificates)',
        'API endpoints for export and verification',
        'Unit tests for certificate and export features'
    ]
    for deliverable in deliverables_p4:
        doc.add_paragraph(deliverable, style='List Bullet')

    doc.add_heading('Phase 4 Success Criteria', level=2)
    success_p4 = [
        'Certificates generate in < 5 seconds',
        'PDFs render correctly in all PDF readers',
        'QR codes scan successfully on mobile devices',
        'Exports complete in < 30 seconds for 1000 records',
        'Certificate verification works 100% of the time',
        'All export formats validate correctly',
        'Test coverage: 75%+ including certificate features',
        'User feedback: "Professional and useful" certificates'
    ]
    for criteria in success_p4:
        doc.add_paragraph(criteria, style='List Bullet')

    doc.add_page_break()

    # ============================================
    # PHASE 5: TESTING & QA
    # ============================================
    add_heading(doc, 'Phase 5: Testing & Quality Assurance', level=1)
    add_status_box(doc, 'STATUS: PENDING', 'orange')

    doc.add_paragraph()
    doc.add_paragraph('Duration: 5 days (Week 7)')
    doc.add_paragraph('Priority: CRITICAL')
    doc.add_paragraph('Dependencies: Phase 0-4 complete')

    doc.add_heading('Overview', level=2)
    doc.add_paragraph(
        'Phase 5 focuses on comprehensive testing, quality assurance, and performance optimization. '
        'The goal is to achieve 85%+ code coverage and ensure the platform is production-ready '
        'with zero critical bugs.'
    )

    doc.add_heading('Feature 1: Test Coverage Expansion', level=2)
    doc.add_paragraph('Estimated Time: 2 days')
    doc.add_paragraph('Target Coverage: 85%+')
    doc.add_paragraph()
    doc.add_paragraph('Test Types:')
    test_types = [
        'Unit Tests: All business logic functions and classes',
        'Integration Tests: API endpoints, database interactions',
        'End-to-End Tests: Complete user workflows (Selenium/Playwright)',
        'Frontend Tests: JavaScript functionality (Jest, optional)',
        'API Tests: All REST endpoints with various inputs',
        'Database Tests: Complex queries, transactions, rollbacks'
    ]
    for test in test_types:
        doc.add_paragraph(test, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Priority Test Areas:')
    priority_tests = [
        'Admin Dashboard: All CRUD operations',
        'Analytics Engine: Metric calculations accuracy',
        'Email System: Template rendering, sending, queueing',
        'Certificate Generation: PDF creation, QR codes',
        'Export Functionality: All formats (CSV, Excel, PDF)',
        'Anti-Cheating: Score calculation, flagging logic',
        'Authentication: Login, permissions, token validation',
        'Assessment Flow: Full user journey from registration to certificate'
    ]
    for test in priority_tests:
        doc.add_paragraph(test, style='List Bullet')

    doc.add_heading('Feature 2: Load Testing', level=2)
    doc.add_paragraph('Estimated Time: 1 day')
    doc.add_paragraph()
    doc.add_paragraph('Load Testing Scenarios:')
    load_scenarios = [
        'Concurrent Users: 100 simultaneous assessments',
        'Database Load: 10,000 assessment records query performance',
        'API Stress Test: 1000 requests/second to critical endpoints',
        'File Generation: Bulk certificate generation (500 PDFs)',
        'Email Queue: 1000 emails in queue processing time',
        'Admin Dashboard: Multiple admins accessing simultaneously'
    ]
    for scenario in load_scenarios:
        doc.add_paragraph(scenario, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Load Testing Tools:')
    load_tools = [
        'Locust: Python-based load testing framework',
        'Apache Bench (ab): Simple HTTP benchmarking',
        'k6 (optional): Modern load testing tool',
        'PostgreSQL EXPLAIN ANALYZE: Query performance analysis',
        'New Relic or DataDog (optional): Application performance monitoring'
    ]
    for tool in load_tools:
        doc.add_paragraph(tool, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Performance Targets:')
    perf_targets = [
        'API Response Time: < 200ms for 95% of requests',
        'Page Load Time: < 2 seconds for all pages',
        'Database Query Time: < 100ms for standard queries',
        'Certificate Generation: < 5 seconds per PDF',
        'Email Sending: < 30 seconds from trigger to delivery',
        'Concurrent Users: Support 100 simultaneous users without degradation'
    ]
    for target in perf_targets:
        doc.add_paragraph(target, style='List Bullet')

    doc.add_heading('Feature 3: Security Testing', level=2)
    doc.add_paragraph('Estimated Time: 1 day')
    doc.add_paragraph()
    doc.add_paragraph('Security Test Types:')
    sec_tests = [
        'SQL Injection: Test all database inputs',
        'XSS (Cross-Site Scripting): Test all user inputs and outputs',
        'CSRF (Cross-Site Request Forgery): Verify CSRF tokens',
        'Authentication Bypass: Test permission checks',
        'Authorization: Verify role-based access control',
        'Session Management: Test session hijacking, fixation',
        'File Upload: Test malicious file uploads',
        'API Security: Test API authentication, rate limiting'
    ]
    for test in sec_tests:
        doc.add_paragraph(test, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Security Testing Tools:')
    sec_tools = [
        'OWASP ZAP: Automated vulnerability scanner',
        'Bandit: Python security linter',
        'Safety: Check dependencies for known vulnerabilities',
        'Manual Penetration Testing: Targeted security assessment',
        'SSL Labs: Test HTTPS configuration (production only)'
    ]
    for tool in sec_tools:
        doc.add_paragraph(tool, style='List Bullet')

    doc.add_heading('Feature 4: Cross-Browser & Device Testing', level=2)
    doc.add_paragraph('Estimated Time: 0.5 days')
    doc.add_paragraph()
    doc.add_paragraph('Browser Testing:')
    browsers = [
        'Chrome (Windows, Mac, Linux)',
        'Firefox (Windows, Mac, Linux)',
        'Safari (Mac, iOS)',
        'Edge (Windows)',
        'Mobile Browsers (Chrome Mobile, Safari Mobile)'
    ]
    for browser in browsers:
        doc.add_paragraph(browser, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Device Testing:')
    devices = [
        'Desktop: 1920x1080, 1366x768',
        'Tablet: iPad (1024x768), Android tablets',
        'Mobile: iPhone (375x667), Android (360x640)',
        'Responsive breakpoints: 320px, 768px, 1024px, 1920px'
    ]
    for device in devices:
        doc.add_paragraph(device, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Special Focus:')
    special_focus = [
        'Audio Recording: MediaRecorder API compatibility',
        'File Downloads: PDF, CSV, Excel downloads',
        'Forms: All input validation and submission',
        'Navigation: All links and redirects',
        'Charts: Chart.js rendering on different screens'
    ]
    for focus in special_focus:
        doc.add_paragraph(focus, style='List Bullet')

    doc.add_heading('Feature 5: Bug Fixing & Optimization', level=2)
    doc.add_paragraph('Estimated Time: 0.5 days')
    doc.add_paragraph()
    doc.add_paragraph('Bug Priority Levels:')
    bug_priorities = [
        'P0 (Critical): Blocks core functionality, immediate fix',
        'P1 (High): Impacts major features, fix within 24 hours',
        'P2 (Medium): Impacts minor features, fix within week',
        'P3 (Low): Cosmetic issues, fix before production'
    ]
    for priority in bug_priorities:
        doc.add_paragraph(priority, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Optimization Focus:')
    optimizations = [
        'Database Queries: Add indexes, optimize joins',
        'API Endpoints: Reduce response payloads, add caching',
        'Frontend: Minify JS/CSS, optimize images',
        'Background Jobs: Optimize Celery task performance',
        'Memory Usage: Profile and fix memory leaks',
        'Database Connections: Pool sizing and connection management'
    ]
    for opt in optimizations:
        doc.add_paragraph(opt, style='List Bullet')

    doc.add_heading('Phase 5 Technical Stack', level=2)
    tech_stack_p5 = [
        'Testing: pytest, pytest-asyncio, pytest-cov',
        'E2E Testing: Selenium or Playwright',
        'Load Testing: Locust',
        'Security: OWASP ZAP, Bandit, Safety',
        'Coverage: coverage.py, codecov',
        'Performance: cProfile, memory_profiler'
    ]
    for tech in tech_stack_p5:
        doc.add_paragraph(tech, style='List Bullet')

    doc.add_heading('Phase 5 Deliverables', level=2)
    deliverables_p5 = [
        'Test coverage report: 85%+',
        'Load testing report with performance metrics',
        'Security audit report with findings',
        'Cross-browser compatibility report',
        'Bug tracking document with all fixes',
        'Performance optimization report',
        'End-to-end test suite (Selenium/Playwright)',
        'CI/CD pipeline with automated tests (optional)',
        'Quality assurance checklist',
        'Pre-production deployment verification'
    ]
    for deliverable in deliverables_p5:
        doc.add_paragraph(deliverable, style='List Bullet')

    doc.add_heading('Phase 5 Success Criteria', level=2)
    success_p5 = [
        'Test coverage: 85%+ achieved',
        'Zero P0 (critical) bugs remaining',
        'All P1 (high) bugs fixed',
        'Load tests pass at 100 concurrent users',
        'Security scan shows zero high-severity vulnerabilities',
        'All major browsers render correctly',
        'Mobile responsive on all tested devices',
        'Performance targets met (response time, page load)',
        'QA sign-off for production deployment'
    ]
    for criteria in success_p5:
        doc.add_paragraph(criteria, style='List Bullet')

    doc.add_page_break()

    # ============================================
    # PHASE 6: DOCUMENTATION & DEPLOYMENT
    # ============================================
    add_heading(doc, 'Phase 6: Documentation & Deployment', level=1)
    add_status_box(doc, 'STATUS: PENDING', 'orange')

    doc.add_paragraph()
    doc.add_paragraph('Duration: 5 days (Week 8)')
    doc.add_paragraph('Priority: MEDIUM')
    doc.add_paragraph('Dependencies: Phase 0-5 complete')

    doc.add_heading('Overview', level=2)
    doc.add_paragraph(
        'Phase 6 focuses on comprehensive documentation, deployment infrastructure, and production '
        'launch preparation. This ensures the platform is fully documented, deployable, and '
        'maintainable for the long term.'
    )

    doc.add_heading('Feature 1: API Documentation', level=2)
    doc.add_paragraph('Estimated Time: 1 day')
    doc.add_paragraph()
    doc.add_paragraph('Documentation Deliverables:')
    api_docs = [
        'OpenAPI/Swagger Specification: Complete API schema',
        'Interactive API Docs: FastAPI auto-generated docs (/docs)',
        'ReDoc Documentation: Alternative API documentation (/redoc)',
        'Authentication Guide: How to obtain and use API tokens',
        'Endpoint Reference: All endpoints with request/response examples',
        'Error Codes: Complete list of error codes and meanings',
        'Rate Limiting: API rate limits and throttling rules',
        'Webhook Documentation: Event triggers and payloads (if applicable)'
    ]
    for doc_item in api_docs:
        doc.add_paragraph(doc_item, style='List Bullet')

    doc.add_heading('Feature 2: User Documentation', level=2)
    doc.add_paragraph('Estimated Time: 1 day')
    doc.add_paragraph()
    doc.add_paragraph('User Manual Sections:')
    user_manual = [
        'Getting Started: Registration and login',
        'Taking an Assessment: Step-by-step guide',
        'Understanding Results: Score interpretation',
        'Certificate Information: How to download and verify',
        'FAQ: Common questions and answers',
        'Troubleshooting: Audio issues, browser compatibility',
        'Support: How to contact support',
        'Privacy Policy: Data handling and GDPR compliance'
    ]
    for section in user_manual:
        doc.add_paragraph(section, style='List Bullet')

    doc.add_heading('Feature 3: Admin Documentation', level=2)
    doc.add_paragraph('Estimated Time: 1 day')
    doc.add_paragraph()
    doc.add_paragraph('Admin Manual Sections:')
    admin_manual = [
        'Admin Dashboard Overview: Main features and navigation',
        'User Management: Creating, editing, deactivating users',
        'Assessment Monitoring: Tracking in-progress and completed assessments',
        'Question Bank Management: Adding, editing, bulk importing questions',
        'Anti-Cheating Dashboard: Reviewing flagged assessments',
        'Analytics: Understanding reports and metrics',
        'Email Configuration: Setting up SMTP or SendGrid',
        'Certificate Management: Generating and revoking certificates',
        'Export Functionality: Bulk data exports and reports',
        'System Configuration: Environment variables and settings',
        'Backup and Restore: Database backup procedures',
        'Monitoring: Health checks and system status'
    ]
    for section in admin_manual:
        doc.add_paragraph(section, style='List Bullet')

    doc.add_heading('Feature 4: Deployment Documentation', level=2)
    doc.add_paragraph('Estimated Time: 1 day')
    doc.add_paragraph()
    doc.add_paragraph('Deployment Guide Sections:')
    deploy_guide = [
        'System Requirements: Hardware, OS, dependencies',
        'Installation Guide: Step-by-step setup instructions',
        'Environment Configuration: .env file setup',
        'Database Setup: PostgreSQL installation and migration',
        'Redis Setup: For Celery and caching',
        'SSL/TLS Configuration: HTTPS setup',
        'Docker Deployment: Dockerfile and docker-compose.yml',
        'Production Configuration: Gunicorn, Nginx setup',
        'Monitoring Setup: Logging, metrics, alerts',
        'Backup Strategy: Automated database backups',
        'Rollback Procedures: How to revert deployments',
        'Scaling Guide: Horizontal and vertical scaling'
    ]
    for section in deploy_guide:
        doc.add_paragraph(section, style='List Bullet')

    doc.add_heading('Feature 5: CI/CD Pipeline (Optional)', level=2)
    doc.add_paragraph('Estimated Time: 1 day (optional)')
    doc.add_paragraph()
    doc.add_paragraph('CI/CD Features:')
    cicd_features = [
        'GitHub Actions: Automated testing on push',
        'Automated Testing: Run pytest on every commit',
        'Code Coverage: Upload coverage reports to Codecov',
        'Linting: Run flake8, black, mypy',
        'Security Scanning: Bandit, Safety checks',
        'Docker Build: Automated Docker image creation',
        'Deployment: Automated deployment to staging',
        'Notifications: Slack/email notifications on build status'
    ]
    for feature in cicd_features:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Example GitHub Actions Workflow:')
    workflow_steps = [
        'Trigger: On push to main branch',
        'Step 1: Checkout code',
        'Step 2: Set up Python 3.11',
        'Step 3: Install dependencies',
        'Step 4: Run linting (flake8, black)',
        'Step 5: Run tests (pytest with coverage)',
        'Step 6: Upload coverage to Codecov',
        'Step 7: Build Docker image',
        'Step 8: Deploy to staging (if tests pass)',
        'Step 9: Notify team (Slack or email)'
    ]
    for step in workflow_steps:
        doc.add_paragraph(step, style='List Bullet')

    doc.add_heading('Feature 6: Production Launch Preparation', level=2)
    doc.add_paragraph('Estimated Time: 0.5 days')
    doc.add_paragraph()
    doc.add_paragraph('Pre-Launch Checklist:')
    prelaunch = [
        'Database: Production PostgreSQL setup and optimized',
        'Environment: All environment variables configured',
        'SSL Certificate: Valid SSL certificate installed',
        'Email: SMTP or SendGrid configured and tested',
        'Monitoring: Logging and alerting configured',
        'Backup: Automated daily database backups',
        'Security: All security scans passed',
        'Performance: Load testing passed',
        'Documentation: All documentation complete',
        'Training: Admin team trained on dashboard',
        'Support: Support contact information published',
        'Legal: Privacy policy and terms of service published'
    ]
    for item in prelaunch:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Launch Day Tasks:')
    launch_tasks = [
        'Database Migration: Run all pending migrations',
        'Load Initial Data: Import question bank, departments',
        'DNS Configuration: Point domain to production server',
        'SSL Verification: Verify HTTPS working correctly',
        'Smoke Testing: Test all critical user flows',
        'Monitoring: Monitor logs and metrics closely',
        'Communication: Notify users of launch',
        'On-Call: Team available for immediate issues'
    ]
    for task in launch_tasks:
        doc.add_paragraph(task, style='List Bullet')

    doc.add_heading('Phase 6 Technical Stack', level=2)
    tech_stack_p6 = [
        'Documentation: MkDocs or Sphinx',
        'API Docs: FastAPI built-in Swagger UI',
        'Deployment: Docker, docker-compose',
        'Web Server: Gunicorn + Nginx',
        'CI/CD: GitHub Actions (optional)',
        'Monitoring: Prometheus + Grafana (optional)',
        'Logging: Python logging + ELK stack (optional)'
    ]
    for tech in tech_stack_p6:
        doc.add_paragraph(tech, style='List Bullet')

    doc.add_heading('Phase 6 Deliverables', level=2)
    deliverables_p6 = [
        'Complete API documentation (OpenAPI spec)',
        'User manual (PDF or web-based)',
        'Admin manual (PDF or web-based)',
        'Deployment guide (step-by-step)',
        'Troubleshooting guide',
        'Docker configuration (Dockerfile, docker-compose.yml)',
        'CI/CD pipeline (GitHub Actions workflow)',
        'Production deployment scripts',
        'Monitoring and logging setup',
        'Backup and restore procedures',
        'Launch checklist and procedures',
        'Post-launch support plan'
    ]
    for deliverable in deliverables_p6:
        doc.add_paragraph(deliverable, style='List Bullet')

    doc.add_heading('Phase 6 Success Criteria', level=2)
    success_p6 = [
        'All documentation complete and reviewed',
        'API documentation accessible at /docs',
        'User and admin manuals available',
        'Deployment guide tested with fresh installation',
        'Docker deployment successful',
        'CI/CD pipeline running (if implemented)',
        'Production environment fully configured',
        'Pre-launch checklist 100% complete',
        'Successful test deployment to staging',
        'Team trained and ready for launch',
        'Support channels established and tested',
        'Production launch successful with zero downtime'
    ]
    for criteria in success_p6:
        doc.add_paragraph(criteria, style='List Bullet')

    doc.add_page_break()

    # ============================================
    # SUMMARY & COST ESTIMATES
    # ============================================
    add_heading(doc, 'Summary & Cost Estimates', level=1)

    doc.add_heading('Development Timeline Summary', level=2)

    summary_table = [
        ['Phase 0', 'Critical Fixes', '5 days', 'COMPLETE', '$0'],
        ['Phase 1', 'Admin Dashboard', '10 days', 'PENDING', '$0'],
        ['Phase 2', 'Analytics', '5 days', 'PENDING', '$0-50/mo'],
        ['Phase 3', 'Email Notifications', '5 days', 'PENDING', '$15-100/mo'],
        ['Phase 4', 'Certificates', '5 days', 'PENDING', '$0'],
        ['Phase 5', 'Testing & QA', '5 days', 'PENDING', '$0'],
        ['Phase 6', 'Documentation', '5 days', 'PENDING', '$0'],
        ['TOTAL', '8 Weeks', '40 days', '1/7 Complete', '$15-150/mo']
    ]
    add_table(doc, ['Phase', 'Name', 'Duration', 'Status', 'Monthly Cost'], summary_table)

    doc.add_heading('Cost Breakdown (Monthly Operating Costs)', level=2)

    cost_table = [
        ['Database', 'PostgreSQL on existing infrastructure', '$0'],
        ['Email Service', 'SendGrid (optional, Free tier: 100/day)', '$0-15/mo'],
        ['Cloud Storage', 'AWS S3 for audio/PDFs (optional)', '$0-20/mo'],
        ['Monitoring', 'Self-hosted or free tier (DataDog/New Relic)', '$0-50/mo'],
        ['SSL Certificate', 'Let\'s Encrypt (free)', '$0'],
        ['Compute', 'Existing infrastructure or VPS', '$0-100/mo'],
        ['TOTAL ESTIMATED', 'Low-cost setup', '$15-150/mo']
    ]
    add_table(doc, ['Service', 'Description', 'Cost'], cost_table)

    doc.add_paragraph()
    doc.add_paragraph('Note: All costs are estimates and can be minimized by using free tiers and existing infrastructure.')

    doc.add_heading('Next Steps', level=2)
    next_steps = [
        'Review and approve Phase 1 plan',
        'Allocate development resources for Phase 1',
        'Confirm email service choice (SMTP vs SendGrid)',
        'Prepare question bank data for import',
        'Schedule Phase 1 kickoff meeting',
        'Begin Phase 1: Admin Dashboard development'
    ]
    for step in next_steps:
        doc.add_paragraph(step, style='List Bullet')

    doc.add_page_break()

    # ============================================
    # APPENDIX: TECHNICAL DETAILS
    # ============================================
    add_heading(doc, 'Appendix: Technical Details', level=1)

    doc.add_heading('A. Technology Stack Summary', level=2)

    tech_categories = [
        ('Backend Framework', 'FastAPI (Python 3.11+)'),
        ('Database', 'PostgreSQL 14+ (or SQLite for dev)'),
        ('ORM', 'SQLAlchemy (async)'),
        ('Caching', 'Redis 7+'),
        ('Background Jobs', 'Celery + Redis'),
        ('Web Server', 'Gunicorn + Nginx'),
        ('Frontend', 'HTML/CSS/JavaScript, Bootstrap 5'),
        ('Charts', 'Chart.js, Plotly (optional)'),
        ('PDF Generation', 'ReportLab or WeasyPrint'),
        ('Email', 'aiosmtplib or SendGrid'),
        ('Testing', 'pytest, Selenium/Playwright'),
        ('CI/CD', 'GitHub Actions'),
        ('Containerization', 'Docker, docker-compose'),
        ('Monitoring', 'Prometheus + Grafana (optional)')
    ]

    tech_table = [[cat, tech] for cat, tech in tech_categories]
    add_table(doc, ['Category', 'Technology'], tech_table)

    doc.add_heading('B. Database Schema Overview', level=2)

    db_tables = [
        ('users', 'User accounts (name, email, division, department)'),
        ('assessments', 'Assessment instances with scores and analytics'),
        ('questions', 'Question bank (all modules, types, answers)'),
        ('assessment_responses', 'Individual question responses'),
        ('division_departments', 'Division and department mapping'),
        ('assessment_config', 'Configuration settings'),
        ('email_logs', 'Email sending history (Phase 3)'),
        ('certificates', 'Generated certificates metadata (Phase 4)')
    ]

    db_table_data = [[table, desc] for table, desc in db_tables]
    add_table(doc, ['Table', 'Description'], db_table_data)

    doc.add_heading('C. Key Design Decisions', level=2)

    design_decisions = [
        'Async Architecture: FastAPI + SQLAlchemy async for high performance',
        'Browser-Native Audio: MediaRecorder API to avoid external audio services',
        'JSON Analytics: analytics_data JSON field for flexible anti-cheating tracking',
        'Celery Background Jobs: For email sending, large exports, metric calculation',
        'Redis Caching: For computed metrics, session data, Celery broker',
        'Modular Design: Clear separation of API, business logic, database layers',
        'Test-Driven: High test coverage from Phase 0, maintained through all phases',
        'Docker-First: Containerized deployment for consistency and scalability'
    ]
    for decision in design_decisions:
        doc.add_paragraph(decision, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    footer = doc.add_paragraph('END OF DOCUMENT')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.bold = True

    # Save document
    filename = 'Complete_Development_Phases_Documentation.docx'
    doc.save(filename)
    print(f"Document created successfully: {filename}")
    return filename

if __name__ == "__main__":
    create_phases_document()
